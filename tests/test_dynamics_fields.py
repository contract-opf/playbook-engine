"""Tests for the negotiation-dynamics fields (issue #177, OPF §3.5.3).

Covers the five behaviors the issue names:
  1. proposed_by derived from a DOCX tracked insertion's author, mapped
     through our_party_aliases/our_authors — "us" on a match, "unknown"
     otherwise (never "counterparty": absence of a match against our own
     side is not evidence of the other side, issue #119).
  2. proposed_by "unknown" / observed_at omitted on a PDF-only trail —
     dynamics are never fabricated.
  3. stance_detail consistent with the counts feeding historical_stance;
     validator rejects held > of.
  4. negotiation_trail built from per-round diffs: a clause changed in
     rounds 1 and 3 yields exactly two entries with correct ordinals and
     resolvable refs.
  5. A dangling trail ref fails validation.
"""

from __future__ import annotations

import copy
import json
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from playbook_engine.clause_differ import ClauseDiff, DocumentDiff, TextHunk, VersionDiff
from playbook_engine.clause_position_compiler import compile_clause_positions
from playbook_engine.deviation_classifier import DeviationResult, RiskDelta
from playbook_engine.docx_ingester import TrackedChange, TrackedChanges, ingest_docx
from playbook_engine.observation_builder import (
    Observation,
    ObservationCitation,
    build_observations,
    build_round_moves,
)
from playbook_engine.tracked_changes_overlay import enrich_clause_diff
from playbook_engine.validator import validate_document

FIXTURES = Path(__file__).parent.parent / "examples" / "fixtures"

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _load_minimal() -> dict[str, Any]:
    with (FIXTURES / "valid_v0_2_minimal.json").open() as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# 1. proposed_by from tracked changes
# ---------------------------------------------------------------------------


def _counterparty_tracked_docx(tmp_path: Path) -> Path:
    """DOCX with a tracked insertion authored by a counterparty reviewer."""
    doc = Document()
    doc.add_heading("Indemnification", level=1)
    p = doc.add_paragraph()
    p.add_run("Each party shall ")
    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "University Counsel")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "solely "
    p.add_run("indemnify the other.")
    path = tmp_path / "counterparty-redline.docx"
    doc.save(str(path))
    return path


def test_proposed_by_from_tracked_changes(tmp_path: Path) -> None:
    """Tracked-change attribution (author + date) flows into proposed_by/
    observed_at. The fixture's author ("University Counsel") matches
    neither the configured our_party_aliases nor any our_authors, so —
    per issue #119 — proposed_by is "unknown", not a guessed
    "counterparty": the mechanism under test is that real tracked-change
    metadata reaches these fields at all, not that an unrecognized name is
    assumed to be the other side."""
    result = ingest_docx(_counterparty_tracked_docx(tmp_path), "deal-1", "v2")
    tracked = result.tracked
    assert tracked.changes, "fixture must carry tracked changes"

    diff = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified",
        hunks=(TextHunk(kind="insert", old_text="", new_text="solely"),),
        text_before="Each party shall indemnify the other.",
        text_after="Each party shall solely indemnify the other.",
        clause_version_before="v1",
        clause_version_after="v2",
        char_span_before=(0, 38),
        char_span_after=(0, 45),
    )
    enrichment = next(
        (eh.enrichment for eh in enrich_clause_diff(diff, tracked) if eh.enrichment is not None),
        None,
    )
    assert enrichment is not None, "tracked insertion must match the diff hunk"
    assert enrichment.author == "University Counsel"

    observations = build_observations(
        "deal-1",
        2,
        "our_paper",
        [(diff, DeviationResult("substantive", RiskDelta("worse", "minor"), basis="judge"))],
        reversals=[],
        attributions=[enrichment],
        our_party_aliases=["Alpha Corp"],
    )
    assert len(observations) == 1
    obs = observations[0]
    assert obs.proposed_by == "unknown"
    # The tracked-change date is still the observation's embedded-metadata
    # date, even though the side is unattributed — date and side are
    # derived independently.
    assert obs.observed_at == "2024-03-15"
    d = obs.to_dict()
    assert d["proposed_by"] == "unknown"
    assert d["observed_at"] == "2024-03-15"


def test_proposed_by_us_when_author_matches_alias(tmp_path: Path) -> None:
    """An author matching our_party_aliases maps to 'us', not 'counterparty'."""
    diff = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified",
        hunks=(TextHunk(kind="insert", old_text="", new_text="mutually"),),
        text_before="Each party shall indemnify the other.",
        text_after="Each party shall mutually indemnify the other.",
        clause_version_before="v1",
        clause_version_after="v2",
    )
    from playbook_engine.tracked_changes_overlay import HunkEnrichment

    observations = build_observations(
        "deal-1",
        2,
        "our_paper",
        [
            (
                diff,
                DeviationResult("reworded_equivalent", RiskDelta("neutral", "none"), basis="judge"),
            )
        ],
        reversals=[],
        attributions=[
            HunkEnrichment(
                author="Alpha Corp Legal", date="2024-03-15T10:00:00Z", tracked_type="insertion"
            )
        ],
        our_party_aliases=["Alpha Corp"],
    )
    assert observations[0].proposed_by == "us"


# ---------------------------------------------------------------------------
# 2. PDF-only trail: unknown / omitted — never fabricated
# ---------------------------------------------------------------------------


def test_proposed_by_unknown_on_pdf_only() -> None:
    """No tracked-changes side-channel (PDF-only trail): a changed clause is
    proposed_by='unknown' and observed_at is OMITTED — no fabrication."""
    diff = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified",
        hunks=(TextHunk(kind="replace", old_text="negligence", new_text="gross negligence"),),
        text_before="Indemnification for negligence.",
        text_after="Indemnification for gross negligence.",
        clause_version_before="v1",
        clause_version_after="v2",
    )
    observations = build_observations(
        "pdf-deal",
        2,
        "our_paper",
        [(diff, DeviationResult("substantive", RiskDelta("worse", "minor"), basis="judge"))],
        reversals=[],
        attributions=None,
        our_party_aliases=["Alpha Corp"],
    )
    obs = observations[0]
    assert obs.proposed_by == "unknown"
    assert obs.observed_at is None
    d = obs.to_dict()
    assert "observed_at" not in d
    assert "counterparty_ref" not in d


def test_unchanged_clause_carries_no_proposed_by() -> None:
    """deviation='none' records absence of change — nothing was proposed."""
    diff = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="unchanged",
        hunks=(),
        text_before="Same text.",
        text_after="Same text.",
        clause_version_before="v1",
        clause_version_after="v2",
    )
    observations = build_observations(
        "deal-1",
        2,
        "our_paper",
        [(diff, DeviationResult("none", RiskDelta("neutral", "none"), basis="deterministic"))],
        reversals=[],
        our_party_aliases=["Alpha Corp"],
    )
    assert observations[0].proposed_by is None
    assert "proposed_by" not in observations[0].to_dict()


# ---------------------------------------------------------------------------
# 3. stance_detail
# ---------------------------------------------------------------------------


def _obs(
    obs_id: str,
    outcome: str = "signed",
    direction: str = "neutral",
    deviation: str = "none",
) -> Observation:
    return Observation(
        observation_id=obs_id,
        taxonomy_id="indemnification",
        text_summary=f"Observation {obs_id}.",
        citation=ObservationCitation(
            document_id=f"deal-{obs_id}", version=2, clause_path="1", char_span=(0, 10)
        ),
        deviation=deviation,
        risk_delta={
            "direction": direction,
            "magnitude": "minor" if direction == "worse" else "none",
        },
        provenance="our_paper",
        outcome=outcome,
        basis="deterministic",
    )


def test_stance_detail_consistent() -> None:
    """stance_detail matches the counts feeding historical_stance: held =
    refusals + non-worse signings, of = OPF-outcome our-paper observations."""
    observations = [
        _obs("a", outcome="signed", direction="neutral"),  # held
        _obs("b", outcome="signed", direction="worse", deviation="substantive"),  # conceded
        _obs(
            "c", outcome="proposed_then_reversed", direction="worse", deviation="substantive"
        ),  # held
    ]
    positions, _, _ = compile_clause_positions(observations, [])
    assert len(positions) == 1
    detail = positions[0].rollup.stance_detail
    assert detail == {"held": 2, "of": 3, "basis": "our_paper"}
    emitted = positions[0].to_dict()["summary"]["stance_detail"]
    assert emitted == {"held": 2, "of": 3, "basis": "our_paper"}


def test_stance_detail_held_exceeding_of_rejected() -> None:
    """Validator rejects held > of (blocking)."""
    doc = _load_minimal()
    doc["evidence"]["clauses"][0]["summary"]["stance_detail"] = {
        "held": 3,
        "of": 2,
        "basis": "our_paper",
    }
    result = validate_document(doc)
    assert not result.ok
    assert any("held" in str(e) for e in result.errors)


# ---------------------------------------------------------------------------
# 4./5. negotiation_trail
# ---------------------------------------------------------------------------


def _clause_diff_for_round(v_before: str, v_after: str, changed: bool) -> ClauseDiff:
    return ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified" if changed else "unchanged",
        hunks=(TextHunk(kind="replace", old_text="a", new_text="b"),) if changed else (),
        text_before=f"Text as of {v_before}.",
        text_after=f"Text as of {v_after}." if changed else f"Text as of {v_before}.",
        clause_version_before=v_before,
        clause_version_after=v_after,
        char_span_before=(0, 20),
        char_span_after=(0, 20),
    )


def test_negotiation_trail_from_rounds() -> None:
    """Clause changed in rounds 1 and 3 of a 4-version deal → exactly two
    trail entries with the right ordinals, and refs that resolve."""
    order = ["v1", "v2", "v3", "v4"]
    doc_diff = DocumentDiff(
        consecutive=(
            VersionDiff("v1", "v2", (_clause_diff_for_round("v1", "v2", changed=True),)),
            VersionDiff("v2", "v3", (_clause_diff_for_round("v2", "v3", changed=False),)),
            VersionDiff("v3", "v4", (_clause_diff_for_round("v3", "v4", changed=True),)),
        ),
        net=VersionDiff("v1", "v4", (_clause_diff_for_round("v1", "v4", changed=True),)),
        version_order=tuple(order),
    )
    moves = build_round_moves("university-of-example", doc_diff)
    assert [m.round for m in moves] == [1, 3]
    # Post-move refs: round 1 cites version ordinal 2, round 3 cites 4.
    assert [m.citation.version for m in moves] == [2, 4]
    assert all(m.moved_by == "unknown" for m in moves)  # no side-channel given

    positions, _, _ = compile_clause_positions([_obs("a"), _obs("b")], [], round_moves=moves)
    clause_dict = positions[0].to_dict()
    trail = clause_dict["negotiation_trail"]
    assert len(trail) == 2
    assert trail[0]["round"] == 1 and trail[1]["round"] == 3
    for entry in trail:
        assert entry["ref"]["document_id"] == "university-of-example"
        assert "taxonomy_id" not in entry  # grouping key, not OPF surface

    # Refs must resolve against corpus.documents (§4).
    doc = _load_minimal()
    doc["evidence"]["clauses"][0]["negotiation_trail"] = [m.to_opf_dict() for m in moves]
    doc["corpus"]["documents"][0]["versions"] = 4
    result = validate_document(doc)
    assert result.ok, [str(e) for e in result.errors]


def test_build_round_moves_round_level_fallback_fires_with_single_author() -> None:
    """issue #118 fix round 2, finding 2 regression guard: build_round_moves
    must consult tracked_changes_overlay.round_level_fallback_attribution as
    a last resort once enrich_clause_diff finds no per-hunk match — a
    round's own side channel with exactly one distinct author must still
    recover moved_by, mapped through our_authors exactly like the direct
    per-hunk-match path already is."""
    diff = _clause_diff_for_round("v1", "v2", changed=True)
    doc_diff = DocumentDiff(
        consecutive=(VersionDiff("v1", "v2", (diff,)),),
        net=VersionDiff("v1", "v2", (diff,)),
        version_order=("v1", "v2"),
    )
    # clause_path "99" / no char_span never overlaps diff's clause_path "1"
    # or char_span (0, 20) — enrich_clause_diff finds no candidate at all,
    # so only the round-level fallback tier can recover an attribution here.
    tc = TrackedChange(
        change_type="insertion",
        author="Alice",
        date="2024-01-01",
        text="unrelated redline text elsewhere in the document",
        clause_path="99",
        char_span=None,
    )
    tracked_by_vid = {"v2": TrackedChanges(document_id="deal-1", version="v2", changes=[tc])}

    moves = build_round_moves(
        "deal-1", doc_diff, tracked_by_vid=tracked_by_vid, our_authors=["Alice"]
    )

    assert len(moves) == 1
    assert moves[0].moved_by == "us"


def test_build_round_moves_round_level_fallback_refuses_with_two_authors() -> None:
    """Two distinct authors in the round's own side channel — both parties'
    marks in one round — must refuse to guess between them: moved_by stays
    'unknown', mirroring round_level_fallback_attribution's own refusal
    rule."""
    diff = _clause_diff_for_round("v1", "v2", changed=True)
    doc_diff = DocumentDiff(
        consecutive=(VersionDiff("v1", "v2", (diff,)),),
        net=VersionDiff("v1", "v2", (diff,)),
        version_order=("v1", "v2"),
    )
    tc1 = TrackedChange(
        change_type="insertion",
        author="Alice",
        date="2024-01-01",
        text="unrelated redline text elsewhere",
        clause_path="99",
        char_span=None,
    )
    tc2 = TrackedChange(
        change_type="insertion",
        author="Bob",
        date="2024-01-02",
        text="yet more unrelated text",
        clause_path="98",
        char_span=None,
    )
    tracked_by_vid = {"v2": TrackedChanges(document_id="deal-1", version="v2", changes=[tc1, tc2])}

    moves = build_round_moves(
        "deal-1", doc_diff, tracked_by_vid=tracked_by_vid, our_authors=["Alice", "Bob"]
    )

    assert len(moves) == 1
    assert moves[0].moved_by == "unknown"


def test_trail_ref_dangling_fails() -> None:
    """A trail ref citing a version beyond the corpus record must fail."""
    doc = _load_minimal()
    entry = {
        "document_id": "university-of-example",
        "round": 1,
        "moved_by": "unknown",
        "change_summary": "Clause modified.",
        "ref": {
            "document_id": "university-of-example",
            "version": 99,
            "clause_path": "1",
        },
    }
    doc["evidence"]["clauses"][0]["negotiation_trail"] = [entry]
    result = validate_document(doc)
    assert not result.ok
    assert any("negotiation_trail" in (e.path or "") for e in result.errors)


def test_removed_clause_cites_last_existing_state() -> None:
    """A removed clause's trail entry cites the before side — the last state
    where the clause existed — because a post-move ref would dangle."""
    removed = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after=None,
        kind="removed",
        hunks=(),
        text_before="Clause that was dropped.",
        text_after="",
        clause_version_before="v1",
        clause_version_after=None,
        char_span_before=(0, 24),
        char_span_after=None,
    )
    doc_diff = DocumentDiff(
        consecutive=(VersionDiff("v1", "v2", (removed,)),),
        net=VersionDiff("v1", "v2", (removed,)),
        version_order=("v1", "v2"),
    )
    moves = build_round_moves("deal-1", doc_diff)
    assert len(moves) == 1
    assert moves[0].citation.version == 1
    assert moves[0].citation.version_id == "v1"
    assert "removed" in moves[0].change_summary


# ---------------------------------------------------------------------------
# content-hash participation (same discipline as x_* extensions)
# ---------------------------------------------------------------------------


def test_dynamics_fields_participate_in_content_hash() -> None:
    from playbook_engine.canonicalize import content_hash

    doc = _load_minimal()
    extended = copy.deepcopy(doc)
    extended["evidence"]["clauses"][0]["summary"]["stance_detail"] = {
        "held": 1,
        "of": 1,
        "basis": "our_paper",
    }
    assert content_hash(doc) != content_hash(extended)


# ---------------------------------------------------------------------------
# Review-pass regressions (2026-07-13 grind, Phase 2)
# ---------------------------------------------------------------------------


def test_no_aliases_never_fabricates_a_side() -> None:
    """With NO our_party_aliases/our_authors configured (the config default
    is [] for both), an attributed author must map to 'unknown' —
    publishing our own attorneys' edits as counterparty asks is exactly the
    guessed attribution §3.5.3 forbids."""
    from playbook_engine.observation_builder import party_side_for_author

    assert party_side_for_author("Jane Attorney", []) == "unknown"
    assert party_side_for_author("Jane Attorney", [""]) == "unknown"
    assert party_side_for_author("Jane Attorney", [], []) == "unknown"


def test_short_author_initials_do_not_match_alias_substring() -> None:
    """Word author strings are often initials; 'IT' ⊂ 'Summit Health' must
    not read as 'us' via reverse containment."""
    from playbook_engine.observation_builder import party_side_for_author

    # Neither matches, so — per issue #119 — these are "unknown", not a
    # confidently wrong "counterparty": absence of a match against our own
    # aliases is not evidence the author is on the other side.
    assert party_side_for_author("IT", ["Summit Health"]) == "unknown"
    assert party_side_for_author("Al", ["Alpha Corporation"]) == "unknown"
    # The forward direction and reasonable-length reverse still work.
    assert party_side_for_author("Alpha Corporation Legal", ["Alpha Corporation"]) == "us"
    assert party_side_for_author("alpha corp", ["Alpha Corporation Holdings"]) == "us"


def test_author_matching_neither_list_is_unknown_not_counterparty() -> None:
    """Issue #119 regression: with our_party_aliases NON-EMPTY, an author
    matching neither our_party_aliases nor our_authors must be 'unknown'.
    Before the fix this fell through to 'counterparty' — the exact guess
    the function's own docstring says never happens, and (per the real
    production corpus measurement in #119) a systematic one-directional
    bias for every unrecognized author, since w:ins/w:del authors are
    personal names/initials, a namespace our_party_aliases (entity/org
    names) was never going to match."""
    from playbook_engine.observation_builder import party_side_for_author

    assert (
        party_side_for_author("Pat Counterparty", ["FixtureCorp"], ["Jane Attorney", "J. Attorney"])
        == "unknown"
    )


def test_author_matching_our_authors_is_us() -> None:
    """Issue #119: our_authors is the people-namespace counterpart to
    our_party_aliases — an author matching it (personal name/initials/
    email, not an org alias) must map to 'us', the same as an
    our_party_aliases match does."""
    from playbook_engine.observation_builder import party_side_for_author

    assert party_side_for_author("Jane Attorney", ["FixtureCorp"], ["Jane Attorney"]) == "us"
    # our_party_aliases match still works unaffected by our_authors being set.
    assert party_side_for_author("FixtureCorp Legal", ["FixtureCorp"], ["Jane Attorney"]) == "us"
    # Short-initials guard applies identically to our_authors' reverse direction.
    assert party_side_for_author("Al", [], ["Alpha Attorney"]) == "unknown"


def test_move_summaries_truncate_after_pseudonymization_boundary() -> None:
    """_summarize_move keeps summaries untruncated (a raw-text slice can cut
    an entity name mid-word and defeat whole-word aliasing); the cap is
    applied by truncate_move_summaries, which the pipeline runs after the
    aliasing pass."""
    from playbook_engine.observation_builder import (
        _summarize_move,
        truncate_move_summaries,
    )

    long_name = "Alpha Corporation Holdings, LLC"
    added = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before=None,
        clause_path_after="1",
        kind="added",
        hunks=(),
        text_before="",
        text_after=("indemnify " + long_name + " against everything ") * 20,
        clause_version_before=None,
        clause_version_after="v2",
    )
    summary = _summarize_move(added)
    assert long_name in summary  # untruncated: the full name survives to the aliasing pass

    doc_diff = DocumentDiff(
        consecutive=(VersionDiff("v1", "v2", (added,)),),
        net=VersionDiff("v1", "v2", (added,)),
        version_order=("v1", "v2"),
    )
    moves = build_round_moves("deal-1", doc_diff)
    truncated = truncate_move_summaries(moves)
    assert all(len(m.change_summary) <= 200 for m in truncated)


def test_search_snippets_truncate_after_pseudonymization_boundary() -> None:
    """Mirrors test_move_summaries_truncate_after_pseudonymization_boundary
    for the sibling search_snippet field (issue #95): build_observations
    keeps search_snippet untruncated (== full_text) so a subsequent
    pipeline pseudonymization pass can whole-word-match a known entity name
    in full before anything is sliced; observation_builder.
    truncate_search_snippets applies the short-phrase cap only afterward —
    never the other way around (see that function's docstring, and
    Observation.search_snippet's)."""
    from playbook_engine.observation_builder import truncate_search_snippets

    long_name = "Alpha Corporation Holdings, LLC"
    diff = ClauseDiff(
        taxonomy_id="indemnification",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified",
        hunks=(),
        text_before="original",
        text_after=("indemnify " + long_name + " against everything ") * 20,
        clause_version_before="v1",
        clause_version_after="v2",
    )
    observations = build_observations(
        "deal-1",
        2,
        "our_paper",
        [(diff, DeviationResult("substantive", RiskDelta("worse", "minor"), basis="judge"))],
        reversals=[],
    )
    assert long_name in observations[0].search_snippet  # untruncated: survives to the aliasing pass
    assert observations[0].search_snippet == observations[0].full_text

    truncated = truncate_search_snippets(observations)
    assert all(len(o.search_snippet) <= 100 for o in truncated)


def test_export_profile_covers_negotiation_trail() -> None:
    """negotiation_trail.change_summary quotes raw clause text and must be
    in the export profile's judged free-text surface (issue #146 contract)."""
    from playbook_engine.export_profile import _extract_text_samples

    doc = _load_minimal()
    doc["evidence"]["clauses"][0]["negotiation_trail"] = [
        {
            "document_id": "university-of-example",
            "round": 1,
            "moved_by": "counterparty",
            "change_summary": "Cap raised from 1x fees to 2x fees.",
            "ref": {"document_id": "university-of-example", "version": 2, "clause_path": "1"},
        }
    ]
    samples, locations = _extract_text_samples(doc)
    trail_paths = [s.path for s in samples if "negotiation_trail" in s.path]
    assert trail_paths, "trail change_summary missing from the judged surface"
    assert locations[trail_paths[0]][1] == "negotiation_trail"


def test_export_profile_covers_search_snippet() -> None:
    """observed_positions[].x_search_snippet quotes raw clause text near a
    citation's location, exactly as residue-prone as text_summary/full_text
    (issue #95's MANDATORY residue-judgment wiring) — must be in the export
    profile's judged free-text surface alongside its two siblings."""
    from playbook_engine.export_profile import _FREE_TEXT_FIELDS, _extract_text_samples

    assert "x_search_snippet" in _FREE_TEXT_FIELDS

    doc = _load_minimal()
    doc["evidence"]["clauses"][0]["observed_positions"][0]["x_search_snippet"] = (
        "Mutual indemnification, negligence-based."
    )
    samples, locations = _extract_text_samples(doc)
    snippet_paths = [s.path for s in samples if s.path.endswith("x_search_snippet")]
    assert snippet_paths, "x_search_snippet missing from the judged free-text surface"
    loc = locations[snippet_paths[0]]
    assert loc[1] == "observed_positions"
    assert loc[3] == "x_search_snippet"
