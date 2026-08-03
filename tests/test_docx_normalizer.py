"""Tests for the tracked-changes/comment DOCX normalizer (issue #84).

SECURITY NOTE: All fixtures are synthetic and built programmatically, reusing
the tracked-changes DOCX builders already established in
tests/test_docx_ingester.py and tests/test_extraction.py (python-docx + lxml
XML injection). No real agreement files are ever committed or referenced
from tests — not even redlined drafts. Party/author names are fictitious
("Alice", "Bob", "Carol", "FixtureCorp") only.
"""

from __future__ import annotations

import zipfile
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from playbook_engine.docx_ingester import _extract_para_text
from playbook_engine.docx_normalizer import normalize_tracked_docx
from tests.test_docx_ingester import (
    _empty_docx,
    _hyperlink_tracked_docx,
    _multi_para_clause_docx,
    _simple_docx,
    _tracked_docx,
)
from tests.test_extraction import _tracked_and_table_docx

# Word namespace, matching the fixture convention already established in
# tests/test_docx_ingester.py and tests/test_extraction.py.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


# ---------------------------------------------------------------------------
# Additional fixtures not already covered by the reused builders above:
# a real Word comment (python-docx 1.2's add_comment API), and a document
# combining a tracked insertion/deletion WITH a comment — the realistic
# redline shape docling actually crashes on.
# ---------------------------------------------------------------------------


def _commented_docx(tmp_path: Path) -> Path:
    """Document with a real Word comment anchored to a run.

    Produces commentRangeStart/commentRangeEnd markers, a dedicated
    commentReference run, and a word/comments.xml part + relationship.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Party A shall perform its obligations diligently.")
    doc.add_comment([run], text="Please clarify scope.", author="Carol", initials="CC")
    path = tmp_path / "commented.docx"
    doc.save(str(path))
    return path


def _tracked_and_commented_docx(tmp_path: Path) -> Path:
    """Tracked insertion + deletion AND a comment in one paragraph."""
    doc = Document()
    doc.add_heading("Obligations", level=1)
    p = doc.add_paragraph()
    p.add_run("Party A shall ")

    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "Alice")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "promptly "

    tail_run = p.add_run("provide services")

    del_elem = etree.SubElement(p._p, _w("del"))
    del_elem.set(_w("id"), "2")
    del_elem.set(_w("author"), "Bob")
    del_elem.set(_w("date"), "2024-03-16T09:00:00Z")
    r_del = etree.SubElement(del_elem, _w("r"))
    dt_del = etree.SubElement(r_del, _w("delText"))
    dt_del.text = " to client"

    doc.add_comment([tail_run], text="Confirm timing.", author="Carol", initials="CC")

    path = tmp_path / "tracked_and_commented.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _count_tag(document: Document, tag: str) -> int:
    return len(list(document.element.body.iter(qn(tag))))


def _paragraph_texts(document: Document) -> list[str]:
    return [p.text for p in document.paragraphs]


# ---------------------------------------------------------------------------
# Structural: no tracked-change or comment-marker elements survive
# ---------------------------------------------------------------------------


def test_normalized_output_has_no_ins_or_del_elements(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert _count_tag(doc, "w:ins") == 0
        assert _count_tag(doc, "w:del") == 0
    finally:
        normalized.unlink(missing_ok=True)


def test_normalized_output_has_no_comment_marker_elements(tmp_path: Path) -> None:
    path = _commented_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert _count_tag(doc, "w:commentReference") == 0
        assert _count_tag(doc, "w:commentRangeStart") == 0
        assert _count_tag(doc, "w:commentRangeEnd") == 0
    finally:
        normalized.unlink(missing_ok=True)


def test_original_fixture_actually_has_a_comments_part(tmp_path: Path) -> None:
    """Sanity check on the fixture itself, so the drop assertions below are
    proven to exercise something real rather than passing on an
    already-commentless document."""
    path = _commented_docx(tmp_path)
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
    assert any("comments" in n.lower() for n in names), names


def test_normalized_output_drops_comments_part_from_the_saved_zip(tmp_path: Path) -> None:
    path = _commented_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        with zipfile.ZipFile(normalized) as zf:
            names = zf.namelist()
        assert not any("comments" in n.lower() for n in names), names
    finally:
        normalized.unlink(missing_ok=True)


def test_normalized_output_has_no_comment_relationships(tmp_path: Path) -> None:
    path = _commented_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert all("comment" not in rel.reltype.lower() for rel in doc.part.rels.values())
    finally:
        normalized.unlink(missing_ok=True)


def test_normalized_file_reopens_cleanly(tmp_path: Path) -> None:
    """Verifier correction: dropping the comments part must not leave a
    dangling relationship — the saved file must reopen via docx.Document
    without raising."""
    path = _tracked_and_commented_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        reopened = Document(str(normalized))  # must not raise
        assert reopened is not None
    finally:
        normalized.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Content preservation
# ---------------------------------------------------------------------------


def test_insertion_text_is_preserved(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert any("promptly" in t for t in _paragraph_texts(doc))
    finally:
        normalized.unlink(missing_ok=True)


def test_deletion_text_is_removed(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert not any("to client" in t for t in _paragraph_texts(doc))
    finally:
        normalized.unlink(missing_ok=True)


def test_insertion_nested_in_hyperlink_is_unwrapped_and_preserved(tmp_path: Path) -> None:
    path = _hyperlink_tracked_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert _count_tag(doc, "w:ins") == 0
        assert any("Schedule A" in t for t in _paragraph_texts(doc))
    finally:
        normalized.unlink(missing_ok=True)


def test_tracked_change_inside_table_cell_is_normalized(tmp_path: Path) -> None:
    path = _tracked_and_table_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert _count_tag(doc, "w:ins") == 0
        assert _count_tag(doc, "w:del") == 0
        all_text = " ".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "Training" in all_text
        assert "$100/hour" in all_text
    finally:
        normalized.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Semantic contract (core acceptance criterion, issue #84 Scope): the
# normalized copy's plain text must equal the accepted-changes view
# _extract_para_text already produces for the ORIGINAL document — property
# style, checked across every kind of fixture this module builds or reuses.
# ---------------------------------------------------------------------------

_FIXTURE_BUILDERS: list[Callable[[Path], Path]] = [
    _simple_docx,
    _tracked_docx,
    _hyperlink_tracked_docx,
    _multi_para_clause_docx,
    _tracked_and_table_docx,
    _commented_docx,
    _tracked_and_commented_docx,
]


@pytest.mark.parametrize("build_fixture", _FIXTURE_BUILDERS, ids=lambda fn: fn.__name__)
def test_normalized_text_matches_accepted_changes_view(
    tmp_path: Path, build_fixture: Callable[[Path], Path]
) -> None:
    path = build_fixture(tmp_path)

    original = Document(str(path))
    original_paragraphs = list(original.element.body.iter(qn("w:p")))
    expected = [_extract_para_text(p)[0] for p in original_paragraphs]

    normalized_path = normalize_tracked_docx(path)
    try:
        normalized = Document(str(normalized_path))
        normalized_paragraphs = list(normalized.element.body.iter(qn("w:p")))
        actual = [_extract_para_text(p)[0] for p in normalized_paragraphs]
    finally:
        normalized_path.unlink(missing_ok=True)

    assert actual == expected


# ---------------------------------------------------------------------------
# Robustness
# ---------------------------------------------------------------------------


def test_normalize_does_not_mutate_the_original_file(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    original_bytes = path.read_bytes()
    normalized = normalize_tracked_docx(path)
    try:
        assert path.read_bytes() == original_bytes
    finally:
        normalized.unlink(missing_ok=True)


def test_normalize_returns_a_distinct_temp_path(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        assert normalized != path
        assert normalized.exists()
    finally:
        normalized.unlink(missing_ok=True)


def test_normalize_plain_docx_with_no_tracked_changes_preserves_content(tmp_path: Path) -> None:
    """A document with nothing to normalize must come through unchanged."""
    path = _simple_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        doc = Document(str(normalized))
        assert _paragraph_texts(doc) == [
            "Definitions",
            '"Facility" means any FixtureCorp-operated training centre.',
            "Training Services",
            "FixtureCorp shall provide training per Schedule A.",
            "Obligations",
            "Each party shall perform its obligations diligently.",
        ]
    finally:
        normalized.unlink(missing_ok=True)


def test_normalize_empty_docx_does_not_raise(tmp_path: Path) -> None:
    path = _empty_docx(tmp_path)
    normalized = normalize_tracked_docx(path)
    try:
        Document(str(normalized))  # must not raise
    finally:
        normalized.unlink(missing_ok=True)
