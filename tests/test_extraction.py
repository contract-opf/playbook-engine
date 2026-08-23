"""Tests for the block-stream extractor.

SECURITY NOTE: All fixtures are built programmatically at test runtime with
python-docx / fpdf2 / raw RTF string literals.  No real agreement files are
ever committed to the repository or referenced from tests.  Party names use
fictitious identifiers ("Alpha Corp", "Beta Ltd") only.
"""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from playbook_engine import extraction
from playbook_engine.docx_ingester import ingest_docx
from playbook_engine.extraction import ExtractionCache, ExtractionError, extract_blocks
from playbook_engine.segmentation_grounding import Block

# fpdf2 is a dev-only fixture dependency — import lazily so CI doesn't fail
# when only pdfplumber is installed without fpdf2 (mirrors test_pdf_ingester.py).
try:
    from fpdf import FPDF  # type: ignore[import-untyped]

    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False

_PANDOC_AVAILABLE = shutil.which("pandoc") is not None


# ---------------------------------------------------------------------------
# Shared round-trip assertion
# ---------------------------------------------------------------------------


def _assert_round_trips(canonical_text: str, blocks: list[Block]) -> None:
    """Assert the extraction invariant required by issue #71:

    - block_ids are "b0", "b1", … in reading order
    - "\\n".join(block texts) == canonical_text
    - every block.text == canonical_text[slice(*block.char_span)]
    """
    assert [b.block_id for b in blocks] == [f"b{i}" for i in range(len(blocks))]
    assert "\n".join(b.text for b in blocks) == canonical_text
    for b in blocks:
        assert b.text == canonical_text[slice(*b.char_span)]


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _simple_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_heading("Indemnification", level=1)
    doc.add_paragraph("Alpha Corp shall indemnify Beta Ltd for direct damages.")
    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph("This agreement is governed by the laws of New York.")
    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


def test_extract_docx_round_trips(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == [
        "Indemnification",
        "Alpha Corp shall indemnify Beta Ltd for direct damages.",
        "Governing Law",
        "This agreement is governed by the laws of New York.",
    ]
    # DOCX is not paginated — page is always 0.
    assert all(b.page == 0 for b in blocks)


def test_extract_docx_skips_empty_paragraphs(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")  # blank paragraph — must not become a block
    doc.add_paragraph("Second paragraph.")
    path = tmp_path / "with_blanks.docx"
    doc.save(str(path))

    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == ["First paragraph.", "Second paragraph."]


def test_extract_docx_returns_shared_block_type(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    _, blocks, _ = extract_blocks(path)
    for b in blocks:
        assert isinstance(b, Block)


# Word namespace, matching docx_ingester's tracked-change fixture convention
# (tests/test_docx_ingester.py) — used to inject a raw w:ins element that
# python-docx's high-level API cannot produce.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _tracked_and_table_docx(tmp_path: Path) -> Path:
    """DOCX with a tracked-change insertion (``w:ins``) and a table.

    Regression fixture for issue #85: ``paragraph.text`` only concatenates
    runs that are direct children of ``w:p`` — runs inside ``w:ins`` are
    excluded — and ``doc.paragraphs`` skips table content entirely.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)

    p = doc.add_paragraph()
    p.add_run("Party A shall ")
    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "Alice")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "promptly "
    p.add_run("provide services.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Service"
    tbl.rows[0].cells[1].text = "Fee"
    tbl.rows[1].cells[0].text = "Training"
    tbl.rows[1].cells[1].text = "$100/hour"

    path = tmp_path / "tracked_and_table.docx"
    doc.save(str(path))
    return path


def test_legacy_docx_captures_tracked_insertions_and_tables(tmp_path: Path) -> None:
    """Regression for issue #85.

    Without docling on PATH, the legacy DOCX fallback must not silently drop
    tracked-change insertions or table content — the very text being
    negotiated (counterparty-inserted language) and table-borne terms must
    survive into ``canonical_text``, matching what ``docx_ingester`` (the
    deterministic path) already captures.
    """
    path = _tracked_and_table_docx(tmp_path)
    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)

    assert "Party A shall promptly provide services." in canonical_text
    assert "Service" in canonical_text
    assert "Fee" in canonical_text
    assert "Training" in canonical_text
    assert "$100/hour" in canonical_text


# ---------------------------------------------------------------------------
# extract_tracked_changes — the side-channel itself (issue #85)
#
# Distinct from the regression above: that one guards canonical_text/blocks
# (the extract_blocks legacy adapter must not silently DROP tracked-insertion
# TEXT). This one guards the separate author/date METADATA side-channel that
# extract_blocks's three call sites have no seam for — extract_tracked_changes
# is the new, independent helper the pipeline's LLM-segmentation branches call
# to get it (see pipeline._llm_tracked_changes).
# ---------------------------------------------------------------------------


def test_extract_tracked_changes_matches_docx_ingester(tmp_path: Path) -> None:
    """extract_tracked_changes must match ingest_docx's own parse of the
    identical fixture — same authors, dates, inserted/deleted text,
    clause_path, and char_span (issue #85's acceptance criterion) — even
    though it re-parses the file independently via a separate call, not by
    sharing ingest_docx's result object.
    """
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)
    expected = ingest_docx(path, document_id="", version="").tracked

    actual = extraction.extract_tracked_changes(path)

    assert actual is not None
    assert actual.to_dict() == expected.to_dict()
    # The acceptance criterion, spelled out explicitly: authors + inserted text.
    insertions = [c for c in actual.changes if c.change_type == "insertion"]
    assert any(c.author == "Alice" and "promptly" in c.text for c in insertions), (
        "Alice's inserted text must be captured"
    )
    deletions = [c for c in actual.changes if c.change_type == "deletion"]
    assert any(c.author == "Bob" for c in deletions), "Bob's deletion must be captured"


def test_extract_tracked_changes_none_for_non_docx(tmp_path: Path) -> None:
    """RTF/PDF have no tracked-changes concept — must return None, not raise
    (mirrors pipeline._ingest_file_tracked's convention for those formats)."""
    path = tmp_path / "agreement.rtf"
    path.write_text(r"{\rtf1\ansi\deff0 Hello\par}", encoding="utf-8")
    assert extraction.extract_tracked_changes(path) is None


def test_extract_tracked_changes_empty_for_clean_docx(tmp_path: Path) -> None:
    """A DOCX with no w:ins/w:del elements returns an EMPTY TrackedChanges,
    not None — same convention as ingest_docx's own DocxIngestResult.tracked.
    """
    path = _simple_docx(tmp_path)
    result = extraction.extract_tracked_changes(path)
    assert result is not None
    assert result.changes == []


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

pytestmark_pdf = pytest.mark.skipif(
    not _FPDF_AVAILABLE, reason="fpdf2 not installed; run: pip install fpdf2>=2.7"
)


def _make_pdf(*, pages: list[list[str]], tmp_path: Path, name: str = "doc.pdf") -> Path:
    """Build a PDF with one or more pages, each a list of line strings."""
    pdf = FPDF()
    for lines in pages:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        for line in lines:
            pdf.multi_cell(0, 8, line)
            pdf.ln(2)  # blank line separator — without it multi_cell exhausts cursor width
    dest = tmp_path / name
    pdf.output(str(dest))
    return dest


@pytestmark_pdf
def test_extract_pdf_round_trips(tmp_path: Path) -> None:
    path = _make_pdf(
        pages=[["Indemnification clause line one.", "Second line same page."]],
        tmp_path=tmp_path,
    )
    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == [
        "Indemnification clause line one.",
        "Second line same page.",
    ]
    assert all(b.page == 1 for b in blocks)


@pytestmark_pdf
def test_extract_pdf_pages_are_1_based(tmp_path: Path) -> None:
    path = _make_pdf(
        pages=[["Page one line."], ["Page two line."]],
        tmp_path=tmp_path,
    )
    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.page for b in blocks] == [1, 2]


@pytestmark_pdf
def test_extract_pdf_empty_raises_extraction_error(tmp_path: Path) -> None:
    pdf = FPDF()
    pdf.add_page()  # no text at all
    path = tmp_path / "empty.pdf"
    pdf.output(str(path))

    with pytest.raises(ExtractionError, match="no text"):
        extract_blocks(path)


# ---------------------------------------------------------------------------
# RTF (via pandoc subprocess)
# ---------------------------------------------------------------------------

pytestmark_rtf = pytest.mark.skipif(
    not _PANDOC_AVAILABLE, reason="pandoc not found on PATH; install it to run RTF extraction tests"
)


def _simple_rtf(tmp_path: Path, name: str = "doc.rtf") -> Path:
    # Trailing space after each \par is required: RTF control words otherwise
    # consume the following characters as part of the control word (matches
    # the convention in tests/test_rtf_ingester.py's _simple_rtf fixture).
    body = (
        r"Indemnification\par "
        r"Alpha Corp shall indemnify Beta Ltd for direct damages.\par "
        r"Governing Law\par "
        r"This agreement is governed by the laws of New York.\par "
    )
    content = r"{\rtf1\ansi\deff0" r"\f0\fs24 " + body + r"}"
    path = tmp_path / name
    path.write_text(content, encoding="utf-8")
    return path


@pytestmark_rtf
def test_extract_rtf_round_trips(tmp_path: Path) -> None:
    path = _simple_rtf(tmp_path)
    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == [
        "Indemnification",
        "Alpha Corp shall indemnify Beta Ltd for direct damages.",
        "Governing Law",
        "This agreement is governed by the laws of New York.",
    ]
    assert all(b.page == 0 for b in blocks)


def test_extract_rtf_missing_pandoc_raises_extraction_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """When pandoc is absent from PATH, RTF extraction must fail loud."""
    path = _simple_rtf(tmp_path)
    monkeypatch.setattr(shutil, "which", lambda _cmd: None)

    with pytest.raises(ExtractionError, match="pandoc"):
        extract_blocks(path)


# ---------------------------------------------------------------------------
# docling (preferred path, via mocked CLI subprocess)
# ---------------------------------------------------------------------------

_CANNED_DOCLING_MARKDOWN = (
    "# Indemnification\n"
    "\n"
    "Alpha Corp shall indemnify **Beta Ltd** for direct damages.\n"
    "\n"
    "## Remedies\n"
    "\n"
    "- Injunctive relief\n"
    "- Monetary damages\n"
    "\n"
    "| Party | Cap |\n"
    "| --- | --- |\n"
    "| Alpha Corp | $1,000,000 |\n"
)


def _mock_docling_subprocess(
    monkeypatch: pytest.MonkeyPatch, markdown: str, *, stem: str = "doc"
) -> list[list[str]]:
    """Mock ``shutil.which("docling")`` and ``subprocess.run`` so
    ``extract_blocks`` takes the docling path and writes ``markdown`` to the
    ``<stem>.md`` file docling would have produced in ``--output``.

    Returns a list that records each ``subprocess.run`` command, so tests can
    assert the real docling invocation shape (subcommand, flags).
    """
    calls: list[list[str]] = []

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        outdir = Path(cmd[cmd.index("--output") + 1])
        (outdir / f"{stem}.md").write_text(markdown, encoding="utf-8")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)
    return calls


def test_extract_docling_round_trips_and_strips_decoration(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _mock_docling_subprocess(monkeypatch, _CANNED_DOCLING_MARKDOWN)
    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")

    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == [
        "Indemnification",
        "Alpha Corp shall indemnify Beta Ltd for direct damages.",
        "Remedies",
        "Injunctive relief",
        "Monetary damages",
        "Party | Cap",
        "Alpha Corp | $1,000,000",
    ]
    # Headings are their own blocks (not merged with surrounding text).
    assert blocks[0].text == "Indemnification"
    assert blocks[2].text == "Remedies"
    # No Markdown decoration survives in the citable text.
    for b in blocks:
        assert "#" not in b.text
        assert "*" not in b.text
        assert not b.text.startswith("-")
        assert "|" not in b.text or b.text.count("|") == 1  # table cells joined, not raw pipes
    # docling output is not paginated.
    assert all(b.page == 0 for b in blocks)


def test_extract_docling_invocation_uses_convert_and_placeholder_images(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """docling >=2.x requires the `convert` subcommand, images must be
    exported as placeholders (not embedded base64), and OCR must be pinned to
    English (docling's RapidOCR default is Chinese, which garbles Latin scans).
    Regression guard for all three — mocked-subprocess tests can't catch a wrong
    CLI shape any other way.
    """
    calls = _mock_docling_subprocess(monkeypatch, _CANNED_DOCLING_MARKDOWN)
    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")

    extract_blocks(path)

    assert len(calls) == 1
    cmd = calls[0]
    assert cmd[0] == "docling"
    assert cmd[1] == "convert"  # subcommand is mandatory in docling >=2.x
    assert cmd[cmd.index("--to") + 1] == "md"
    assert cmd[cmd.index("--image-export-mode") + 1] == "placeholder"
    assert cmd[cmd.index("--ocr-lang") + 1] == "eng"  # not docling's Chinese default


def test_extract_docling_drops_image_and_comment_lines(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Embedded images (`![...](data:...base64)`) and docling's HTML image
    placeholder (`<!-- image -->`) are never citable text — they must be
    dropped, not emitted as blocks (else base64 blobs poison the LLM input).
    """
    markdown = (
        "# Title\n"
        "\n"
        "![Image](data:image/png;base64,iVBORw0KGgoAAAANSUhEUg)\n"
        "\n"
        "<!-- image -->\n"
        "\n"
        "Real clause text.\n"
    )
    _mock_docling_subprocess(monkeypatch, markdown)
    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")

    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == ["Title", "Real clause text."]
    assert "base64" not in canonical_text
    assert "<!--" not in canonical_text


def test_extract_docling_preferred_over_legacy_adapter(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Even a .docx path uses docling (uniformly) when it's on PATH."""
    _mock_docling_subprocess(monkeypatch, "# Title\n\nBody text.\n", stem="doc")
    path = tmp_path / "doc.docx"
    # Deliberately not a real DOCX — proves the legacy python-docx adapter
    # was never invoked (it would fail to open this file).
    path.write_bytes(b"not a real docx")

    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == ["Title", "Body text."]


def test_extract_docling_absent_falls_back_to_legacy_adapter(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """When docling is not on PATH, extraction is unchanged (regression)."""
    monkeypatch.setattr(extraction.shutil, "which", lambda _cmd: None)
    path = _simple_docx(tmp_path)

    canonical_text, blocks, _ = extract_blocks(path)
    _assert_round_trips(canonical_text, blocks)
    assert [b.text for b in blocks] == [
        "Indemnification",
        "Alpha Corp shall indemnify Beta Ltd for direct damages.",
        "Governing Law",
        "This agreement is governed by the laws of New York.",
    ]


def test_extract_blocks_reports_extractor(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """``extract_blocks`` surfaces which extractor ran (issue #129): this was
    previously only visible via a ``logging.info`` line suppressed by default
    Python logging config, so a host install falling back to the legacy
    adapter (no docling -> no OCR on scanned PDFs) was invisible to the
    operator. ``detect_extractor`` (the same check ``extract_blocks`` makes
    internally) must agree with what actually ran, in both directions."""
    path = _simple_docx(tmp_path)

    monkeypatch.setattr(extraction.shutil, "which", lambda _cmd: None)
    _, _, extractor = extract_blocks(path)
    assert extractor == "legacy"
    assert extraction.detect_extractor(path) == "legacy"

    _mock_docling_subprocess(monkeypatch, "# Title\n\nBody text.\n", stem=path.stem)
    _, _, extractor = extract_blocks(path)
    assert extractor == "docling"
    assert extraction.detect_extractor(path) == "docling"


# docling failure → legacy fallback (per-file, not a whole-doc skip).
#
# When docling is on PATH but fails on a specific file, ``extract_blocks``
# falls back to the legacy per-format adapter for that one file rather than
# skipping it — otherwise redline drafts (which docling 2.x's DOCX backend
# raises on, via ``etree.QName`` on comment nodes) silently drop out of the
# negotiation trail. The three failure shapes below (empty output, non-zero
# exit, timeout) each recover via the fallback; a file the legacy adapter
# *also* cannot parse still raises, preserving the skip-on-unrecoverable
# contract the caller relies on. The returned ``extractor`` label reports the
# fallback so it stays visible in reporting (issue #129).


def test_extract_docling_empty_output_falls_back_to_legacy(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = _simple_docx(tmp_path)
    # docling "succeeds" but emits only whitespace → treated as failure.
    _mock_docling_subprocess(monkeypatch, "   \n\n  ", stem=path.stem)

    canonical_text, blocks, extractor = extract_blocks(path)

    assert extractor == "legacy"  # fell back, and honestly reports it
    _assert_round_trips(canonical_text, blocks)
    assert any("indemnify" in b.text.lower() for b in blocks)


def test_extract_docling_subprocess_failure_falls_back_to_legacy(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = _simple_docx(tmp_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    canonical_text, blocks, extractor = extract_blocks(path)

    assert extractor == "legacy"
    assert any("indemnify" in b.text.lower() for b in blocks)


def test_docling_timeout_falls_back_and_enforces_timeout(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A hung docling subprocess must not block the run forever: the ``timeout=``
    cap is still passed (issue #98), and on ``TimeoutExpired`` the file falls
    back to the legacy adapter rather than being skipped outright.
    """
    path = _simple_docx(tmp_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        assert kwargs.get("timeout"), "docling subprocess must be called with a timeout"
        raise subprocess.TimeoutExpired(cmd, kwargs["timeout"])

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    canonical_text, blocks, extractor = extract_blocks(path)

    assert extractor == "legacy"
    assert any("indemnify" in b.text.lower() for b in blocks)


def test_extract_docling_failure_unparseable_file_still_raises(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The fallback does not weaken the skip-on-unrecoverable contract: when
    docling fails *and* the legacy adapter also cannot parse the file (here a
    stub that is not a real PDF), ``extract_blocks`` still raises
    ``ExtractionError`` so the caller skips that one version.
    """

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")  # not a real PDF; legacy fails too

    with pytest.raises(ExtractionError):
        extract_blocks(path)


# ---------------------------------------------------------------------------
# docling DOCX crash -> normalize-and-retry before legacy (issue #84)
#
# docling 2.x's DOCX backend crashes on tracked-changes/comment nodes
# (``etree.QName`` on a comment factory) — exactly what redline drafts
# contain, the highest-value documents in a negotiation corpus. A DOCX
# docling failure must now be retried once on a pre-normalized copy (see
# playbook_engine.docx_normalizer) BEFORE falling back to the legacy adapter
# — recovering real docling structure (headings as blocks) for redlines
# instead of degrading to the legacy adapter, which has no heading detection
# at all.
# ---------------------------------------------------------------------------


def test_docling_failure_on_tracked_docx_recovers_via_normalized_retry(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The acceptance criterion: a tracked-changes DOCX that makes docling
    raise is extracted via docling (on the normalized copy), not legacy —
    the block stream must show real docling structure (a heading as its own
    block), and the returned extractor label must show NO degradation at
    all (this is not counted as a ``backend-error`` fallback)."""
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    calls: list[list[str]] = []

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        target = Path(cmd[2])
        outdir = Path(cmd[cmd.index("--output") + 1])
        if len(calls) == 1:
            # First attempt targets the ORIGINAL tracked-changes file and
            # crashes, mirroring docling 2.x's real etree.QName-on-comment
            # failure on redlines.
            raise subprocess.CalledProcessError(1, cmd, stderr="etree.QName crash on w:ins")
        # Second attempt targets the normalized copy and succeeds, with real
        # docling structure (a heading) the legacy adapter cannot produce.
        (outdir / f"{target.stem}.md").write_text(
            "# Obligations\n\nParty A shall promptly provide services.\n", encoding="utf-8"
        )
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    canonical_text, blocks, extractor = extract_blocks(path)

    assert len(calls) == 2, "docling must be retried once on the normalized copy"
    assert calls[0][2] == str(path), "first attempt must target the original file"
    assert calls[1][2] != str(path), (
        "retry must target the normalized temp copy, not the original file"
    )
    assert not Path(calls[1][2]).exists(), "the normalized temp copy must be cleaned up"

    assert extractor == "docling", "must recover via docling structure, not degrade to legacy"
    assert extractor.reason is None
    assert extractor.fallback_from is None
    assert extractor.detail is None

    assert "Obligations" in canonical_text
    assert blocks[0].text == "Obligations"  # docling heading -> its own block


def test_docling_failure_on_docx_falls_back_to_legacy_when_retry_also_fails(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """When docling fails on BOTH the original and the normalized copy, the
    original per-file legacy fallback still runs unchanged (regression
    guard: the new retry must never swallow a genuine double failure)."""
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    calls: list[list[str]] = []

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    canonical_text, blocks, extractor = extract_blocks(path)

    assert len(calls) == 2, "both the original and normalized-copy attempts must be made"
    assert calls[0][2] == str(path)
    assert calls[1][2] != str(path)

    assert extractor == "legacy"
    assert extractor.reason == "backend-error"
    assert extractor.fallback_from == "docling"
    # Legacy DOCX adapter still captures the accepted-changes text (issue #85).
    assert "Party A shall promptly provide services" in canonical_text
    assert "to client" not in canonical_text


def test_docling_failure_on_pdf_does_not_attempt_normalized_retry(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The normalize-and-retry recovery is DOCX-specific: a PDF docling
    failure must fall back to legacy in exactly one subprocess call, exactly
    as before this ticket."""

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    calls: list[list[str]] = []

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-1.4 fake content")  # not a real PDF; legacy fails too

    with pytest.raises(ExtractionError):
        extract_blocks(path)

    assert len(calls) == 1, "non-DOCX docling failures must not trigger the normalize retry"


def test_docling_failure_on_unparseable_docx_still_raises(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The new retry must not mask the final raise: a ``.docx`` that is not
    actually a valid DOCX fails docling, fails the normalize attempt
    (python-docx cannot open it either — the failure is caught broadly, not
    just ``ExtractionError``), and then fails the legacy adapter too —
    ``extract_blocks`` must still raise, preserving the
    skip-on-unrecoverable contract."""

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    calls: list[list[str]] = []

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    path = tmp_path / "not_a_real.docx"
    path.write_bytes(b"this is not a docx file at all")

    with pytest.raises(ExtractionError):
        extract_blocks(path)

    # Only the ORIGINAL attempt ever reaches docling: normalize_tracked_docx
    # itself raises trying to open the garbage bytes, before a second
    # docling invocation could happen.
    assert len(calls) == 1


# ---------------------------------------------------------------------------
# Unsupported types / general errors
# ---------------------------------------------------------------------------


def test_unsupported_extension_raises_extraction_error(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("plain text file", encoding="utf-8")

    with pytest.raises(ExtractionError, match="unsupported"):
        extract_blocks(path)


def test_missing_file_raises_extraction_error(tmp_path: Path) -> None:
    path = tmp_path / "does_not_exist.docx"

    with pytest.raises(ExtractionError, match="not found"):
        extract_blocks(path)


# ---------------------------------------------------------------------------
# ExtractionCache (issue #132) — a repeat extract_blocks() call over
# unchanged file content must skip extraction entirely.
# ---------------------------------------------------------------------------


def test_extraction_cache_second_call_skips_extraction(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A second ``extract_blocks(path, cache=...)`` call for unchanged content
    never re-invokes the underlying (legacy DOCX) extraction helper."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")

    calls: list[Path] = []
    real_extract_docx_lines = extraction._extract_docx_lines

    def _counting_extract_docx_lines(p: Path) -> list[tuple[str, int]]:
        calls.append(p)
        return real_extract_docx_lines(p)

    monkeypatch.setattr(extraction, "_extract_docx_lines", _counting_extract_docx_lines)

    first = extract_blocks(path, cache=cache)
    assert len(calls) == 1, "first call must extract (cache miss)"

    second = extract_blocks(path, cache=cache)
    assert len(calls) == 1, "second call over unchanged content must hit the cache"

    assert second == first


def test_extraction_cache_persists_across_instances(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A fresh ``ExtractionCache`` pointed at the same file on disk still hits
    (load-on-init, same contract as ``VerdictStore``/``SegmentationVerdictCache``)
    — for the SAME extractor environment. The cache key is now content hash
    PLUS extractor environment (issue #77), so this pins ``detect_extractor``
    rather than relying on whatever happens to be on the host's PATH, to test
    same-environment persistence specifically (a cross-environment miss is
    covered separately by ``test_extraction_cache_success_retried_under_better_extractor``)."""
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    monkeypatch.setattr(extraction, "detect_extractor", lambda p: "legacy")

    canonical_text, blocks, extractor = extract_blocks(path, cache=ExtractionCache(cache_path))

    # New instance, same on-disk file — must load the entry written above.
    reloaded = ExtractionCache(cache_path)
    cached = reloaded.get(path)
    assert cached is not None
    cached_text, cached_blocks, cached_extractor = cached
    assert cached_text == canonical_text
    assert cached_extractor == extractor
    assert [b.text for b in cached_blocks] == [b.text for b in blocks]
    assert [b.char_span for b in cached_blocks] == [b.char_span for b in blocks]


def test_extraction_cache_miss_for_changed_content(tmp_path: Path) -> None:
    """Editing the file's content after a cache hit is recorded busts the cache
    (key is the file's content hash, not its path)."""
    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Original text.")
    doc.save(str(path))

    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    first_text, _, _ = extract_blocks(path, cache=cache)
    assert cache.get(path) is not None

    doc2 = Document()
    doc2.add_paragraph("Changed text.")
    doc2.save(str(path))

    assert cache.get(path) is None, "changed content must not replay the old entry"
    second_text, _, _ = extract_blocks(path, cache=cache)
    assert second_text != first_text


def test_extraction_cache_put_omits_per_block_text(tmp_path: Path) -> None:
    """``put()`` must not persist per-block ``"text"`` (issue #67): it is fully
    determined by ``canonical_text[char_span]``, so storing it duplicated every
    document's text in the cache entry. ``get()`` must still round-trip the
    original block text, reconstructed from canonical_text/char_span."""
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    cache = ExtractionCache(cache_path)

    canonical_text, blocks, extractor = extract_blocks(path, cache=cache)

    record = json.loads(cache_path.read_text().splitlines()[0])
    stored_blocks = record["verdict"]["blocks"]
    assert stored_blocks, "fixture must produce at least one block"
    for b in stored_blocks:
        assert "text" not in b, "per-block text must not be serialized (issue #67)"

    cached = cache.get(path)
    assert cached is not None
    cached_text, cached_blocks, cached_extractor = cached
    assert cached_text == canonical_text
    assert cached_extractor == extractor
    assert [b.text for b in cached_blocks] == [b.text for b in blocks]
    assert [b.char_span for b in cached_blocks] == [b.char_span for b in blocks]


def test_extraction_cache_get_honors_stored_text_for_pre_fix_entries(tmp_path: Path) -> None:
    """A cache entry written by the pre-#67 ``put()`` (per-block ``"text"``
    present) must still load via the ``b.get("text")`` fallback.

    The stored ``"text"`` is deliberately made to differ from what
    ``canonical_text[char_span]`` would produce, so this test fails if a
    future change drops the fallback and always reconstructs from the span
    instead of preferring an already-stored value.
    """
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    cache = ExtractionCache(cache_path)

    canonical_text = "Alpha Corp shall indemnify Beta Ltd for direct damages."
    legacy_value = {
        "canonical_text": canonical_text,
        "blocks": [
            {
                "block_id": "b0",
                "page": 0,
                "char_span": [0, len(canonical_text)],
                "text": "STORED LEGACY TEXT",  # deliberately != canonical_text[0:len]
            }
        ],
        "extractor": "legacy",
    }
    cache._store.put(extraction._extraction_cache_payload(path), legacy_value)

    cached = cache.get(path)
    assert cached is not None
    _, cached_blocks, _ = cached
    assert cached_blocks[0].text == "STORED LEGACY TEXT"


# ---------------------------------------------------------------------------
# Negative caching of extraction failures (pre-derivation QA): a scanned PDF
# that yields no text must not be re-OCR'd on every pipeline command — but a
# better extractor environment (legacy -> docling) must retry it.
# ---------------------------------------------------------------------------


def test_extraction_failure_is_negative_cached(tmp_path: Path, monkeypatch) -> None:
    from playbook_engine import extraction as ext

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    calls = {"n": 0}

    def _no_text(path, suffix):
        calls["n"] += 1
        return []

    monkeypatch.setattr(ext, "_extract_legacy_lines", _no_text)
    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")

    with pytest.raises(ext.ExtractionError):
        ext.extract_blocks(pdf, cache=cache)
    assert calls["n"] == 1

    # Second call: served from the negative cache, extractor never invoked.
    with pytest.raises(ext.ExtractionError, match="cached failure"):
        ext.extract_blocks(pdf, cache=cache)
    assert calls["n"] == 1


def test_extraction_failure_retried_under_better_extractor(tmp_path: Path, monkeypatch) -> None:
    from playbook_engine import extraction as ext

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [])
    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")
    with pytest.raises(ext.ExtractionError):
        ext.extract_blocks(pdf, cache=cache)

    # docling appears on PATH: the legacy-era failure must NOT block a retry.
    monkeypatch.setattr(ext, "detect_extractor", lambda p: "docling")
    monkeypatch.setattr(ext, "_extract_docling_lines", lambda p: [("OCR recovered text.", 0)])
    canonical, blocks, extractor = ext.extract_blocks(pdf, cache=cache)
    assert "OCR recovered text." in canonical
    assert extractor == "docling"

    # And the success overwrites the failure marker.
    hit = cache.get(pdf)
    assert hit is not None and hit[2] == "docling"


def test_extraction_cache_success_retried_under_better_extractor(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Mirror of ``test_extraction_failure_retried_under_better_extractor`` for
    the SUCCESS case (issue #77): ``get_failure()`` was already scoped to the
    current extractor environment, but ``get()`` was not — a legacy-era
    SUCCESS was strictly stickier than a legacy-era failure. A born-digital
    file that legacy extracted badly-but-non-emptily (no OCR, garbled
    columns) must not be frozen at that bad extraction once docling becomes
    available: it must be re-extracted, and the docling result must win.

    This must fail against pre-fix code: the pre-fix cache key is content
    hash only, so the second ``extract_blocks`` call below would hit the
    legacy-era entry and never invoke docling at all.
    """
    from playbook_engine import extraction as ext

    path = _simple_docx(tmp_path)
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")
    legacy_canonical, _, legacy_extractor = ext.extract_blocks(path, cache=cache)
    assert legacy_extractor == "legacy"

    # docling appears on PATH: the legacy-era SUCCESS must NOT be replayed.
    monkeypatch.setattr(ext, "detect_extractor", lambda p: "docling")
    docling_calls = {"n": 0}

    def _fake_docling_lines(p: Path) -> list[tuple[str, int]]:
        docling_calls["n"] += 1
        return [("Docling recovered text.", 0)]

    monkeypatch.setattr(ext, "_extract_docling_lines", _fake_docling_lines)

    docling_canonical, _, docling_extractor = ext.extract_blocks(path, cache=cache)
    assert docling_calls["n"] == 1, "legacy-era success must not short-circuit docling extraction"
    assert docling_extractor == "docling"
    assert docling_canonical != legacy_canonical
    assert "Docling recovered text." in docling_canonical

    # And the docling success is what a same-environment lookup now returns.
    hit = cache.get(path)
    assert hit is not None and hit[2] == "docling"

    # The stale legacy-era entry is still there but no longer reachable: a
    # same-environment (legacy) lookup would still hit it too, unchanged.
    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")
    legacy_hit = cache.get(path)
    assert legacy_hit is not None and legacy_hit[2] == "legacy"
    assert legacy_hit[0] == legacy_canonical


def test_docling_timeout_failure_negative_caches_under_docling_env(
    tmp_path: Path, monkeypatch
) -> None:
    """A docling-environment failure must be cached AS a docling failure.

    The docling→legacy fallback relabels the attempt "legacy" before the
    no-text raise; storing that label would make every docling-environment
    lookup miss and re-burn the full OCR timeout each pipeline round.
    """
    from playbook_engine import extraction as ext

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    docling_calls = {"n": 0}

    def _docling_times_out(path):
        docling_calls["n"] += 1
        raise ext.ExtractionError("docling timed out after 600s")

    monkeypatch.setattr(ext, "detect_extractor", lambda p: "docling")
    monkeypatch.setattr(ext, "_extract_docling_lines", _docling_times_out)
    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [])

    with pytest.raises(ext.ExtractionError):
        ext.extract_blocks(pdf, cache=cache)
    assert docling_calls["n"] == 1

    # Second call in the SAME docling environment: fail fast from the cache.
    with pytest.raises(ext.ExtractionError, match="cached failure"):
        ext.extract_blocks(pdf, cache=cache)
    assert docling_calls["n"] == 1


# ---------------------------------------------------------------------------
# extract_blocks(refresh=True) (issue #78) — the primitive an operator-
# invoked --no-cache relies on to force a real recompute of a suspect
# extraction: cache READS are bypassed but WRITES still happen, so the
# refreshed entry serves the next plain (non-refresh) call.
# ---------------------------------------------------------------------------


def test_extract_blocks_refresh_bypasses_read_but_still_writes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """``refresh=True`` must skip the cache read (always re-extract from
    source) while still refreshing the cache entry — "reads bypassed, writes
    still happen". Before this parameter existed, ``extract_blocks`` had no
    way at all to bypass a cache hit, so an operator's ``--no-cache`` could
    not force re-extraction of a suspect cached result.
    """
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")

    calls: list[Path] = []
    real_extract_docx_lines = extraction._extract_docx_lines

    def _counting_extract_docx_lines(p: Path) -> list[tuple[str, int]]:
        calls.append(p)
        return real_extract_docx_lines(p)

    monkeypatch.setattr(extraction, "_extract_docx_lines", _counting_extract_docx_lines)

    first = extract_blocks(path, cache=cache)
    assert len(calls) == 1, "first call must extract (cache miss)"

    second = extract_blocks(path, cache=cache)
    assert len(calls) == 1, "plain repeat call must hit the cache, not re-extract"
    assert second == first

    third = extract_blocks(path, cache=cache, refresh=True)
    assert len(calls) == 2, "refresh=True must bypass the cache read and re-extract"
    assert third == first  # unchanged source content -> identical result

    fourth = extract_blocks(path, cache=cache)
    assert len(calls) == 2, (
        "the refresh must have written a fresh entry — the next plain call must hit it"
    )
    assert fourth == first


def test_extract_blocks_refresh_ignores_cached_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """``refresh=True`` must also bypass a cached FAILURE (``get_failure``),
    not just a cached success — an operator forcing ``--no-cache`` to retry a
    suspect extraction must not be blocked by its own prior negative-cache
    entry in the SAME extractor environment.
    """
    from playbook_engine import extraction as ext

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")
    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [])

    with pytest.raises(ext.ExtractionError):
        ext.extract_blocks(pdf, cache=cache)

    # Plain retry (no refresh): served from the negative cache, fails fast.
    with pytest.raises(ext.ExtractionError, match="cached failure"):
        ext.extract_blocks(pdf, cache=cache)

    # refresh=True, same environment: the operator forces a real retry — this
    # time extraction succeeds. The cached failure must not block it.
    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [("Recovered text.", 0)])
    canonical, _, extractor = ext.extract_blocks(pdf, cache=cache, refresh=True)
    assert "Recovered text." in canonical
    assert extractor == "legacy"

    # The success overwrites the failure marker for future plain calls.
    hit = cache.get(pdf)
    assert hit is not None and hit[2] == "legacy"


def test_extract_blocks_refresh_failure_does_not_clobber_cached_success(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A ``refresh=True`` call that transiently yields no text must NOT
    overwrite a pre-existing cached SUCCESS for the same key (issue #78
    round 2).

    ``_extraction_cache_payload`` keys ``put`` and ``put_failure`` identically
    (path content hash + format version + extractor environment — no
    dependency on ``refresh``), so before this guard, a refresh attempt that
    transiently yielded no lines (the documented docling-OCR-timeout mode)
    would negative-cache directly over a previously-good entry. Every later
    PLAIN call would then raise "cached failure" forever, even once the
    extractor is healthy again — recovery was hand-deleting
    extraction_cache.jsonl by hand, exactly the state issue #78 exists to
    eliminate, and newly reachable only via the very flag meant to fix it.
    """
    from playbook_engine import extraction as ext

    path = _simple_docx(tmp_path)
    cache = ext.ExtractionCache(tmp_path / "extraction_cache.jsonl")

    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")

    # Cold run: succeeds and caches the real extracted content.
    good = ext.extract_blocks(path, cache=cache)
    assert good[2] == "legacy"

    # Refresh run: extraction transiently yields nothing (e.g. an OCR
    # timeout) even though the source content is unchanged and genuinely
    # extractable (proven by the successful cold run above under the
    # IDENTICAL cache key — same bytes, same extractor environment).
    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [])
    with pytest.raises(ext.ExtractionError, match="yielded no text"):
        ext.extract_blocks(path, cache=cache, refresh=True)

    # The prior success must survive untouched...
    hit = cache.get(path)
    assert hit is not None, "a transient refresh failure must not clobber the cached success"
    assert hit == good

    # ...and a subsequent PLAIN call must replay it, not raise a cached
    # failure (the extractor being "healthy again" in the wording above is
    # irrelevant here — the cache read never re-invokes the extractor at
    # all on a plain hit).
    replay = ext.extract_blocks(path, cache=cache)
    assert replay == good


def test_extract_blocks_refresh_failure_still_negative_caches_without_prior_success(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The new refresh-vs-clobber guard (issue #78 round 2) must not disable
    negative-caching outright — only the "overwrite an existing success"
    case is special-cased. A refresh with NO pre-existing cached success
    (first attempt under this key, or a prior same-key failure) still
    negative-caches as before, so repeated plain calls keep failing fast
    instead of re-attempting a full docling OCR timeout every round.
    """
    from playbook_engine import extraction as ext

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ext.ExtractionCache(tmp_path / "cache.jsonl")

    monkeypatch.setattr(ext, "detect_extractor", lambda p: "legacy")
    monkeypatch.setattr(ext, "_extract_legacy_lines", lambda p, s: [])

    # refresh=True as the very first call for this key: no cached success
    # exists yet, so the failure must still be negative-cached.
    with pytest.raises(ext.ExtractionError, match="yielded no text"):
        ext.extract_blocks(pdf, cache=cache, refresh=True)

    assert cache.get(pdf) is None
    assert cache.get_failure(pdf) is not None, (
        "a refresh failure with no pre-existing success must still negative-cache"
    )

    # A subsequent PLAIN call must fail fast from the negative cache, not
    # re-invoke the extractor.
    calls = {"n": 0}
    real = ext._extract_legacy_lines

    def _counting(p: Path, s: str) -> list[tuple[str, int]]:
        calls["n"] += 1
        return real(p, s)

    monkeypatch.setattr(ext, "_extract_legacy_lines", _counting)
    with pytest.raises(ext.ExtractionError, match="cached failure"):
        ext.extract_blocks(pdf, cache=cache)
    assert calls["n"] == 0


# ---------------------------------------------------------------------------
# Declared extractor (issue #80): extract_blocks(extractor="docling" |
# "legacy" | "auto") — config-declarable environment with fail-loud
# semantics. "auto" (the default) is exactly today's behavior; the tests
# above (which never pass `extractor=`) already pin that. These tests cover
# the two NEW declared values.
# ---------------------------------------------------------------------------


def test_declared_docling_unavailable_raises_before_any_file_io(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """``extractor="docling"`` on a docling-less host must raise
    ``ExtractionError`` immediately — before any file I/O — so even a path
    that does not exist on disk fails with the docling-missing message, not
    a "file not found" message (issue #80's "raised before any file I/O"
    acceptance criterion)."""
    monkeypatch.setattr(extraction.shutil, "which", lambda _cmd: None)
    missing_path = tmp_path / "does-not-exist.docx"
    assert not missing_path.exists()

    with pytest.raises(ExtractionError, match="docling") as exc_info:
        extract_blocks(missing_path, extractor="docling")
    assert "file not found" not in str(exc_info.value)


def test_declared_docling_unavailable_message_names_remedy(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The fail-loud error must name both the config line and the remedy
    (issue #80's acceptance criterion), not just say "unavailable"."""
    path = _simple_docx(tmp_path)
    monkeypatch.setattr(extraction.shutil, "which", lambda _cmd: None)

    with pytest.raises(ExtractionError) as exc_info:
        extract_blocks(path, extractor="docling")
    message = str(exc_info.value)
    assert "extraction.extractor" in message
    assert "docling" in message
    assert "PATH" in message
    assert "legacy" in message or "auto" in message  # the remedy is named


def test_declared_docling_available_uses_docling(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = _simple_docx(tmp_path)
    _mock_docling_subprocess(monkeypatch, "# Title\n\nBody text.\n", stem=path.stem)

    canonical_text, _blocks, extractor = extract_blocks(path, extractor="docling")
    assert extractor == "docling"
    assert "Title" in canonical_text


def test_declared_legacy_forces_legacy_even_when_docling_available(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """``extractor="legacy"`` must skip docling entirely even when it IS on
    PATH — deterministic, container-free runs (issue #80's acceptance
    criterion)."""
    path = _simple_docx(tmp_path)
    calls = _mock_docling_subprocess(monkeypatch, "# SHOULD NOT BE USED\n", stem=path.stem)

    canonical_text, _blocks, extractor = extract_blocks(path, extractor="legacy")
    assert extractor == "legacy"
    assert calls == [], "declared legacy must never invoke the docling subprocess"
    assert "Indemnification" in canonical_text  # real legacy-adapter DOCX heading


def test_invalid_extractor_value_raises(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    with pytest.raises(ExtractionError, match="invalid extractor"):
        extract_blocks(path, extractor="turbo")


def test_auto_extractor_default_matches_no_argument(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """``extractor="auto"`` (the new parameter's default) must be
    byte-identical to calling ``extract_blocks`` with no ``extractor``
    argument at all — "absent section and auto are behaviorally identical
    to today" (issue #80's acceptance criterion)."""
    path = _simple_docx(tmp_path)
    monkeypatch.setattr(extraction, "detect_extractor", lambda p: "legacy")

    explicit_auto = extract_blocks(path, extractor="auto")
    no_arg = extract_blocks(path)
    assert explicit_auto == no_arg


def test_declared_extractor_cache_keys_do_not_collide_across_environments(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A run declaring ``extractor="legacy"`` on a docling-equipped host must
    not share a cache key with a declared ``extractor="docling"`` extraction
    of the SAME file — else a later docling-declared (or auto, which also
    resolves to docling here) run would silently replay the lower-quality
    legacy content under the docling label. This is the same class of bug
    issue #77 closed for auto-detected environments, now for the DECLARED
    override (issue #80) — a bare ``detect_extractor(path)`` PATH check
    would wrongly key both entries under "docling" since docling is on PATH
    throughout this test.
    """
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "cache.jsonl")
    calls = _mock_docling_subprocess(
        monkeypatch, "# DOCLING HEADING\n\nDocling body.\n", stem=path.stem
    )

    # Force legacy, even though docling is on PATH throughout this test.
    legacy_text, _, legacy_extractor = extract_blocks(path, cache=cache, extractor="legacy")
    assert legacy_extractor == "legacy"
    assert calls == [], "declared legacy must never invoke docling"
    assert "DOCLING HEADING" not in legacy_text

    # Now declare docling explicitly for the SAME file: must actually run
    # docling, not replay the legacy-declared cache entry.
    docling_text, _, docling_extractor = extract_blocks(path, cache=cache, extractor="docling")
    assert docling_extractor == "docling"
    assert len(calls) == 1, "the docling-declared run must actually invoke docling"
    assert "DOCLING HEADING" in docling_text
    assert docling_text != legacy_text

    # Both entries persist independently, each reachable under its own
    # declared environment.
    legacy_hit = cache.get(path, extractor="legacy")
    assert legacy_hit is not None and legacy_hit[2] == "legacy"
    assert legacy_hit[0] == legacy_text
    docling_hit = cache.get(path, extractor="docling")
    assert docling_hit is not None and docling_hit[2] == "docling"
    assert docling_hit[0] == docling_text


def test_declared_extractor_failure_scoped_to_declared_environment(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A negative-cached failure under a declared ``"legacy"`` run must not
    block a subsequent declared ``"docling"`` run for the same file — mirrors
    ``test_extraction_failure_retried_under_better_extractor`` for the
    DECLARED-extractor dimension rather than PATH auto-detection (issue
    #80).
    """
    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake scanned pdf")
    cache = ExtractionCache(tmp_path / "cache.jsonl")

    def _fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    monkeypatch.setattr(extraction.shutil, "which", _fake_which)
    monkeypatch.setattr(extraction, "_extract_legacy_lines", lambda p, s: [])

    with pytest.raises(ExtractionError):
        extract_blocks(pdf, cache=cache, extractor="legacy")

    # Same declared environment: fails fast from the negative cache.
    with pytest.raises(ExtractionError, match="cached failure"):
        extract_blocks(pdf, cache=cache, extractor="legacy")

    # Declared docling for the SAME file: the legacy-declared failure must
    # NOT block this — docling actually runs and succeeds.
    monkeypatch.setattr(extraction, "_extract_docling_lines", lambda p: [("Recovered.", 0)])
    canonical_text, _, extractor = extract_blocks(pdf, cache=cache, extractor="docling")
    assert extractor == "docling"
    assert "Recovered." in canonical_text


# ---------------------------------------------------------------------------
# ExtractorLabel — structured reason for why LEGACY ran (issue #81).
#
# The three cases: docling absent from PATH under "auto" -> "env-missing";
# docling raised on this file -> "backend-error"; config declared legacy ->
# "declared". reason is None whenever extractor == "docling" (no fallback).
# ---------------------------------------------------------------------------


def test_extractor_label_reason_env_missing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """auto mode, docling never on PATH -> reason="env-missing", no fallback_from
    (docling was never attempted, so there's nothing it fell back FROM)."""
    monkeypatch.setattr(extraction.shutil, "which", lambda _cmd: None)
    path = _simple_docx(tmp_path)

    _, _, label = extract_blocks(path)

    assert isinstance(label, extraction.ExtractorLabel)
    assert label == "legacy"
    assert label.extractor == "legacy"
    assert label.reason == "env-missing"
    assert label.fallback_from is None
    assert label.detail is None


def test_extractor_label_reason_declared(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """extractor="legacy" (config-declared) -> reason="declared", even when
    docling IS available — a deliberate choice, not a degradation, and
    distinct from "env-missing" even though both resolve to "legacy"."""
    path = _simple_docx(tmp_path)
    calls = _mock_docling_subprocess(monkeypatch, "# SHOULD NOT RUN\n", stem=path.stem)

    _, _, label = extract_blocks(path, extractor="legacy")

    assert label == "legacy"
    assert label.reason == "declared"
    assert label.fallback_from is None
    assert calls == [], "declared legacy must never invoke docling"


def test_extractor_label_reason_backend_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """docling on PATH but raises on this file -> reason="backend-error",
    fallback_from="docling", and detail carries the underlying exception
    text (in-memory only — see the privacy-focused tests below)."""
    path = _simple_docx(tmp_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    _, _, label = extract_blocks(path)

    assert label == "legacy"
    assert label.reason == "backend-error"
    assert label.fallback_from == "docling"
    assert label.detail is not None
    assert "docling" in label.detail.lower()


def test_extractor_label_reason_none_when_docling_succeeds(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """docling runs clean, no fallback -> reason/fallback_from/detail all None."""
    path = _simple_docx(tmp_path)
    _mock_docling_subprocess(monkeypatch, "# Title\n\nBody text.\n", stem=path.stem)

    _, _, label = extract_blocks(path)

    assert label == "docling"
    assert label.reason is None
    assert label.fallback_from is None
    assert label.detail is None


def test_extractor_label_is_str_subclass_backward_compatible(tmp_path: Path) -> None:
    """ExtractorLabel must be usable exactly like the old plain string
    everywhere a caller might store/format/compare/serialize it — this is
    what lets every pre-existing ``extractor == "legacy"``/f-string/
    ``json.dumps`` call site in this codebase keep working unchanged even
    though extract_blocks now returns a richer object (issue #81)."""
    path = _simple_docx(tmp_path)
    _, _, label = extract_blocks(path)

    assert isinstance(label, str)
    assert f"extractor={label}" == "extractor=legacy"
    assert json.dumps({"extractor": label}) == '{"extractor": "legacy"}'
    assert {"legacy": 1}.get(label) == 1  # hashable/usable as a dict key like any str


# ---------------------------------------------------------------------------
# ExtractionCache round-tripping of reason/fallback_from (issue #81) —
# "detail" is deliberately NEVER round-tripped (in-memory-only, see
# ExtractorLabel's docstring: it embeds the absolute source path).
# ---------------------------------------------------------------------------


def test_extraction_cache_roundtrips_reason_and_fallback_from(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    _, _, written = extract_blocks(path, cache=cache)
    assert written.reason == "backend-error"
    assert written.fallback_from == "docling"
    assert written.detail is not None

    # A fresh ExtractionCache instance (load-on-init from disk) must still
    # report the reason/fallback_from correctly — but never the detail.
    reloaded = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _, _, cached = reloaded.get(path)
    assert cached == "legacy"
    assert cached.reason == "backend-error"
    assert cached.fallback_from == "docling"
    assert cached.detail is None, "detail must never be round-tripped through the cache"


def test_extraction_cache_does_not_persist_detail_on_disk(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The raw exception text (which embeds the absolute source path, and
    therefore the counterparty/entity name baked into the corpus folder
    structure) must never reach the on-disk extraction_cache.jsonl at all —
    not just be unreadable via get() (issue #81)."""
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    cache = ExtractionCache(cache_path)

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    extract_blocks(path, cache=cache)

    record = json.loads(cache_path.read_text().splitlines()[0])
    assert "detail" not in record["verdict"], "detail must never be persisted to disk"
    assert record["verdict"]["reason"] == "backend-error"
    assert record["verdict"]["fallback_from"] == "docling"


def test_extraction_cache_get_tolerates_pre_81_entry_without_reason_key(
    tmp_path: Path,
) -> None:
    """A cache VALUE written before issue #81 (no "reason"/"fallback_from"
    keys at all) must still load without raising KeyError — get() falls back
    to None for both via ``dict.get``. This is a shape-completeness
    guarantee only; it does NOT assert reason=None is the correct real-world
    contract for a genuinely warm pre-#81 cache. A REAL pre-#81 entry is
    filed under format_version "1" and never reaches this code path
    unlabeled: the #81 ladder rung either reconstructs its reason exactly or
    discards it (see the scoped-invalidation section at the end of this
    file)."""
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    cache = ExtractionCache(cache_path)

    canonical_text = "Alpha Corp shall indemnify Beta Ltd for direct damages."
    pre_81_value = {
        "canonical_text": canonical_text,
        "blocks": [
            {"block_id": "b0", "page": 0, "char_span": [0, len(canonical_text)]},
        ],
        "extractor": "legacy",
        # No "reason"/"fallback_from" keys — exactly the pre-#81 shape.
    }
    cache._store.put(extraction._extraction_cache_payload(path), pre_81_value)

    cached = cache.get(path)
    assert cached is not None, "loading a pre-#81-shaped entry must not raise KeyError"
    _, _, label = cached
    assert label == "legacy"
    assert label.reason is None
    assert label.fallback_from is None


# ---------------------------------------------------------------------------
# Scoped cache-format invalidation — the format ladder
# (playbook_engine.cache_format; see _EXTRACTION_CACHE_FORMAT_LADDER).
#
# A format bump used to be a bare number in the cache KEY, so bumping it
# discarded the whole corpus cache — 1h45m-5h17m of re-extraction on the
# reference corpus (44 documents / 161 versions) — regardless of how few
# entries the change actually broke. The tests below pin the replacement:
# entries are migrated (#81, representational) or invalidated by the specific
# rung that broke them (#84, redline-DOCX-only), with the blunt full discard
# kept as a declared, deliberate escape hatch.
# ---------------------------------------------------------------------------


def _fake_pdf(tmp_path: Path, name: str = "draft.pdf") -> Path:
    """A file with a .pdf extension whose CONTENT is never read.

    Every test below either hits the cache (extraction never runs) or asserts
    that extraction was attempted; none needs a real PDF, and using a stub
    keeps these tests free of the optional fpdf2 fixture dependency.
    """
    path = tmp_path / name
    path.write_bytes(b"%PDF-1.4 stub")
    return path


def _plant(
    cache: ExtractionCache,
    path: Path,
    *,
    format_version: str,
    extractor_env: str,
    value: dict,
) -> None:
    """File *value* under a hand-built cache key at an OLD *format_version*.

    Hardcodes the key shape (rather than calling _extraction_cache_payload,
    which would pick up today's version) specifically to simulate genuine
    on-disk state written by an older release.
    """
    cache._store.put(
        {
            "file_sha256": extraction._sha256_file(path),
            "format_version": format_version,
            "extractor_env": extractor_env,
        },
        value,
    )


def _success(text: str, extractor: str, **extra: object) -> dict:
    """A stored SUCCESS value in the on-disk shape (no per-block text — #67)."""
    return {
        "canonical_text": text,
        "blocks": [{"block_id": "b0", "page": 0, "char_span": [0, len(text)]}],
        "extractor": extractor,
        **extra,
    }


_PLANTED_TEXT = "Alpha Corp shall indemnify Beta Ltd for direct damages."


# --- #81 rung: representational, therefore migrated -------------------------


def test_pre_81_docling_entry_is_migrated_instead_of_re_extracted(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A pre-#81 entry whose stored extractor is "docling" carries exactly the
    blocks today's code would produce for the same bytes; only the label's
    SHAPE changed. It must be rewritten in place — no docling subprocess, no
    re-extraction (issue #81's bump used to force one for every file in the
    corpus, this entry included)."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "docling"),
    )

    monkeypatch.setattr(
        extraction.shutil, "which", lambda cmd: "/usr/bin/docling" if cmd == "docling" else None
    )

    def _never(path: Path) -> list[tuple[str, int]]:
        raise AssertionError("a migratable pre-#81 entry must not trigger re-extraction")

    monkeypatch.setattr(extraction, "_extract_docling_lines", _never)

    canonical_text, blocks, label = extract_blocks(path, cache=cache)

    assert canonical_text == _PLANTED_TEXT
    assert blocks[0].text == _PLANTED_TEXT
    assert label == "docling"
    assert label.reason is None, "docling ran clean — the one exactly-recoverable reason"
    assert label.fallback_from is None
    assert cache.migrated_count == 1
    assert cache.invalidated_count == 0


def test_migrated_entry_is_refiled_under_the_current_format(tmp_path: Path) -> None:
    """Migration is paid once per entry, not once per lookup: the rewritten
    value is stored under the CURRENT format key, so the next process (a fresh
    ExtractionCache loading the same JSONL) hits it directly."""
    path = _simple_docx(tmp_path)
    cache_path = tmp_path / "extraction_cache.jsonl"
    cache = ExtractionCache(cache_path)
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "docling"),
    )

    assert cache.get(path, extractor="docling") is not None
    assert cache.migrated_count == 1

    reloaded = ExtractionCache(cache_path)
    hit = reloaded.get(path, extractor="docling")
    assert hit is not None
    assert reloaded.migrated_count == 0, "the entry is already current — nothing left to migrate"

    stored = [json.loads(line) for line in cache_path.read_text().splitlines()]
    assert len(stored) == 2, "the migration appended one record; the stale one is superseded"
    assert stored[-1]["verdict"]["reason"] is None
    assert stored[-1]["verdict"]["fallback_from"] is None


def test_pre_81_legacy_entry_under_docling_env_migrates_to_backend_error(
    tmp_path: Path,
) -> None:
    """The reason for a pre-#81 LEGACY entry under a DOCLING environment is
    exactly recoverable: legacy output can only exist under a docling
    environment via a live per-file fallback. Migrating it to
    "backend-error"/"docling" is strictly better than the old ``dict.get``
    tolerance, which reloaded it as ``reason=None`` — the mislabel that made
    max_fallback/corpus_manifest blind and motivated the bump."""
    path = _fake_pdf(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "legacy"),
    )

    hit = cache.get(path, extractor="docling")

    assert hit is not None
    _, _, label = hit
    assert label == "legacy"
    assert label.reason == "backend-error"
    assert label.fallback_from == "docling"


def test_pre_81_legacy_env_reason_comes_from_the_declared_extractor(tmp_path: Path) -> None:
    """Under a LEGACY environment docling was never attempted, so the reason is
    "declared" or "env-missing" purely as a function of what the CURRENT run
    declared — the same expression extract_blocks uses for a fresh extraction.
    The caller supplies it, so the entry is migrated rather than discarded."""
    path = _fake_pdf(tmp_path)

    declared = ExtractionCache(tmp_path / "declared.jsonl")
    _plant(
        declared,
        path,
        format_version="1",
        extractor_env="legacy",
        value=_success(_PLANTED_TEXT, "legacy"),
    )
    hit = declared.get(path, extractor="legacy", declared_extractor="legacy")
    assert hit is not None
    assert hit[2].reason == "declared", "config asked for legacy — a choice, not a degradation"

    auto = ExtractionCache(tmp_path / "auto.jsonl")
    _plant(
        auto,
        path,
        format_version="1",
        extractor_env="legacy",
        value=_success(_PLANTED_TEXT, "legacy"),
    )
    hit = auto.get(path, extractor="legacy", declared_extractor="auto")
    assert hit is not None
    assert hit[2].reason == "env-missing", "auto found no docling — a real degradation"
    assert hit[2].fallback_from is None


def test_pre_81_legacy_env_entry_is_invalidated_when_the_reason_is_undecidable(
    tmp_path: Path,
) -> None:
    """Without a declared extractor to consult, a pre-#81 legacy-environment
    entry could be either "declared" (never counts against
    config.extraction.max_fallback) or "env-missing" (always does). That
    difference drives a fail-loud budget gate, so the entry is DISCARDED
    rather than labeled by guess — re-extraction costs minutes, a wrong label
    silently poisons every downstream artifact (replaces the old
    test_extraction_cache_format_version_bump_misses_pre_81_key, which asserted
    the same outcome via the blanket bump)."""
    path = _fake_pdf(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    assert extraction._EXTRACTION_CACHE_FORMAT_VERSION != "1", (
        "this test's premise requires the current format version to have moved on from '1'"
    )
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="legacy",
        value=_success(_PLANTED_TEXT, "legacy"),
    )

    assert cache.get(path, extractor="legacy") is None
    assert cache.invalidated_count == 1


def test_pre_81_entry_with_unrecognised_extractor_is_invalidated(tmp_path: Path) -> None:
    """A stored extractor that is neither "docling" nor "legacy" (a hand-edit,
    a foreign writer) has no defensible label to migrate to."""
    path = _fake_pdf(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "tesseract"),
    )

    assert cache.get(path, extractor="docling") is None
    assert cache.invalidated_count == 1


# --- #84 rung: semantic but narrow, therefore a predicate -------------------


def test_84_rung_invalidates_the_redline_docx_fallback_entry(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A pre-#84 entry for a tracked-changes DOCX that recorded a
    docling->legacy fallback must be invalidated: #84 made that exact case
    retry docling on a normalized copy, so the correct output for the same
    bytes under the same environment genuinely changed. This is the one class
    of entry the #84 bump HAD to discard — and now the only one it does.

    The planted entry mirrors real pre-#84 on-disk state: the KEY's
    ``extractor_env`` is "docling" (the environment the file was extracted
    UNDER, not the post-fallback adapter — see the pipeline.py comment added
    in d08b895), while the stored VALUE's ``extractor`` is "legacy" with
    ``reason="backend-error"``/``fallback_from="docling"``.
    """
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")

    assert extraction._EXTRACTION_CACHE_FORMAT_VERSION not in ("1", "2"), (
        "this test's premise requires the current format version to have moved on from '2'"
    )

    stale_text = "Party A shall provide services to client."
    _plant(
        cache,
        path,
        format_version="2",
        extractor_env="docling",
        value=_success(stale_text, "legacy", reason="backend-error", fallback_from="docling"),
    )

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else None

    calls: list[list[str]] = []

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        calls.append(cmd)
        target = Path(cmd[2])
        outdir = Path(cmd[cmd.index("--output") + 1])
        if len(calls) == 1:
            # First attempt targets the ORIGINAL tracked-changes file and
            # crashes, mirroring docling 2.x's real etree.QName-on-comment
            # failure on redlines.
            raise subprocess.CalledProcessError(1, cmd, stderr="etree.QName crash on w:ins")
        # Second attempt targets the normalized copy and succeeds, with real
        # docling structure (a heading) the legacy adapter cannot produce.
        (outdir / f"{target.stem}.md").write_text(
            "# Obligations\n\nParty A shall promptly provide services.\n", encoding="utf-8"
        )
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)

    canonical_text, blocks, extractor = extract_blocks(path, cache=cache)

    assert len(calls) == 2, (
        "a warm pre-#84 cache entry must not short-circuit extraction with the "
        "stale legacy/backend-error result — the normalize-and-retry must "
        "still actually run against docling"
    )
    assert extractor == "docling", "must recover via docling, not replay the cached legacy label"
    assert extractor.reason is None
    assert extractor.fallback_from is None
    assert "Obligations" in canonical_text
    assert blocks[0].text == "Obligations"
    assert cache.invalidated_count == 1


def test_84_rung_keeps_every_entry_the_normalized_retry_cannot_change(tmp_path: Path) -> None:
    """The entries the #84 bump discarded for nothing. #84's retry is DOCX-only
    and only fires when docling FAILED, so none of these could have changed —
    each must survive the bump and hit."""
    docx = _simple_docx(tmp_path)
    pdf = _fake_pdf(tmp_path)

    # A DOCX docling success: docling never failed, so the retry never ran.
    docling_ok = ExtractionCache(tmp_path / "a.jsonl")
    _plant(
        docling_ok,
        docx,
        format_version="2",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "docling", reason=None, fallback_from=None),
    )
    hit = docling_ok.get(docx, extractor="docling")
    assert hit is not None and hit[2].reason is None

    # A DOCX under a declared-legacy environment: docling was never attempted.
    declared = ExtractionCache(tmp_path / "b.jsonl")
    _plant(
        declared,
        docx,
        format_version="2",
        extractor_env="legacy",
        value=_success(_PLANTED_TEXT, "legacy", reason="declared", fallback_from=None),
    )
    hit = declared.get(docx, extractor="legacy")
    assert hit is not None and hit[2].reason == "declared"

    # A PDF that DID fall back on a docling failure: the normalize-and-retry
    # is DOCX-only, so this entry's output is unchanged by #84.
    pdf_fallback = ExtractionCache(tmp_path / "c.jsonl")
    _plant(
        pdf_fallback,
        pdf,
        format_version="2",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "legacy", reason="backend-error", fallback_from="docling"),
    )
    hit = pdf_fallback.get(pdf, extractor="docling")
    assert hit is not None and hit[2].reason == "backend-error"

    for cache in (docling_ok, declared, pdf_fallback):
        assert cache.invalidated_count == 0
        assert cache.migrated_count == 1, "carried forward from format 2 without re-extraction"


def test_84_rung_invalidates_a_docx_legacy_entry_with_no_recorded_reason(tmp_path: Path) -> None:
    """Undecidable is treated as affected. A format-2 DOCX entry labeled
    "legacy" with no ``reason`` key cannot occur through put() or through the
    #81 migration (both always record one), so it is a hand-edited or foreign
    entry — exactly when guessing is least defensible."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _plant(
        cache,
        path,
        format_version="2",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "legacy"),
    )

    assert cache.get(path, extractor="docling") is None
    assert cache.invalidated_count == 1


# --- Failure (negative-cache) entries follow the same ladder ---------------


def test_pre_84_docx_failure_under_docling_is_retried_but_pdf_failure_is_kept(
    tmp_path: Path,
) -> None:
    """A DOCX that negative-cached "no text" under a docling environment before
    the normalize-and-retry existed must be retried (the retry may now yield
    text). A PDF failure cannot be changed by a DOCX-only retry, so it keeps
    its negative cache instead of re-burning the full docling/OCR timeout."""
    docx = _simple_docx(tmp_path)
    pdf = _fake_pdf(tmp_path)

    docx_cache = ExtractionCache(tmp_path / "docx.jsonl")
    _plant(
        docx_cache,
        docx,
        format_version="2",
        extractor_env="docling",
        value={"error": "extraction yielded no text", "extractor": "docling"},
    )
    assert docx_cache.get_failure(docx, extractor="docling") is None
    assert docx_cache.invalidated_count == 1

    pdf_cache = ExtractionCache(tmp_path / "pdf.jsonl")
    _plant(
        pdf_cache,
        pdf,
        format_version="2",
        extractor_env="docling",
        value={"error": "extraction yielded no text", "extractor": "docling"},
    )
    assert pdf_cache.get_failure(pdf, extractor="docling") == "extraction yielded no text"
    assert pdf_cache.invalidated_count == 0


def test_pre_81_failure_entry_migrates_unchanged(tmp_path: Path) -> None:
    """Failure entries never carried reason/fallback_from in either format, so
    the #81 rung passes them through — a pre-#81 negative cache for a PDF is
    still honored instead of costing a fresh (and still fruitless) OCR pass."""
    pdf = _fake_pdf(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    _plant(
        cache,
        pdf,
        format_version="1",
        extractor_env="docling",
        value={"error": "extraction yielded no text", "extractor": "docling"},
    )

    assert cache.get_failure(pdf, extractor="docling") == "extraction yielded no text"
    assert cache.get(pdf, extractor="docling") is None, "a failure entry is not a success hit"


# --- Mechanism-level guarantees --------------------------------------------


def test_warm_current_format_hit_never_probes_stale_versions(tmp_path: Path) -> None:
    """The hot path pays nothing for the ladder: an entry already at the
    current format costs exactly one store lookup, as before."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    cache.put(path, _PLANTED_TEXT, [], "docling", environment="docling")

    lookups: list[dict] = []
    real_get = cache._store.get

    def counting_get(payload: dict) -> dict | None:
        lookups.append(payload)
        return real_get(payload)

    cache._store.get = counting_get  # type: ignore[method-assign]
    assert cache.get(path, extractor="docling") is not None
    assert len(lookups) == 1
    assert lookups[0]["format_version"] == extraction._EXTRACTION_CACHE_FORMAT_VERSION


def test_migration_can_be_disabled(tmp_path: Path) -> None:
    """``migrate=False`` restores the pre-ladder behavior — an older-format
    entry is simply invisible — for a caller that wants a provably cold read."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl", migrate=False)
    _plant(
        cache,
        path,
        format_version="1",
        extractor_env="docling",
        value=_success(_PLANTED_TEXT, "docling"),
    )

    assert cache.get(path, extractor="docling") is None
    assert cache.migrated_count == 0
    assert cache.invalidated_count == 0


def test_a_discard_all_rung_still_invalidates_everything(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The blunt option survives as a declared escape hatch: a hypothetical
    future rung that genuinely invalidates every entry discards even a
    perfectly-shaped current-format entry."""
    path = _simple_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    cache.put(path, _PLANTED_TEXT, [], "docling", environment="docling")
    assert cache.get(path, extractor="docling") is not None

    superseded = extraction._EXTRACTION_CACHE_FORMAT_VERSION
    monkeypatch.setattr(
        extraction,
        "_EXTRACTION_CACHE_FORMAT_LADDER",
        (
            *extraction._EXTRACTION_CACHE_FORMAT_LADDER,
            extraction.CacheFormatStep(
                version="4",
                issue="#999",
                summary="hypothetical change that no predicate can scope",
                discard_all=True,
            ),
        ),
    )
    monkeypatch.setattr(extraction, "_EXTRACTION_CACHE_FORMAT_VERSION", "4")
    monkeypatch.setattr(extraction, "_EXTRACTION_CACHE_STALE_VERSIONS", (superseded, "2", "1"))

    fresh = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    assert fresh.get(path, extractor="docling") is None
    assert fresh.invalidated_count == 1


def test_format_version_constant_tracks_the_ladder_head() -> None:
    """The constant is derived, not maintained — a new rung IS the bump."""
    assert (
        extraction._EXTRACTION_CACHE_FORMAT_LADDER[-1].version
        == extraction._EXTRACTION_CACHE_FORMAT_VERSION
    )
    assert extraction._EXTRACTION_CACHE_STALE_VERSIONS == ("2", "1")
