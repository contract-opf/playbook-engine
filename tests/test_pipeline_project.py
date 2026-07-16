"""Pipeline-level tests for the mine/project split (issue #63).

Verifies two acceptance criteria:
  AC-1: project_playbook performs NO _ingest_file and NO classify_tree/judge
        calls even when a template is configured.
  AC-2: Re-running project_playbook does not rewrite observations.jsonl
        (no re-mining side-effect).

SECURITY NOTE: All fixtures use programmatically constructed RTF text with
synthetic, fictional content.  No real agreement files are referenced.
Fictional party names only (e.g. "Alpha Corp", "Beta University").
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import yaml
from docx import Document
from lxml import etree

from playbook_engine.canonicalize import compute_section_digests, content_hash
from playbook_engine.config import load_config
from playbook_engine.deviation_classifier import DeviationResult, RiskDelta
from playbook_engine.observation_builder import Observation, ObservationCitation
from playbook_engine.pipeline import mine_corpus, project_playbook
from playbook_engine.taxonomy import load_taxonomy

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


# ---------------------------------------------------------------------------
# RTF fixture helpers
# ---------------------------------------------------------------------------

_RTF_PROLOGUE = (
    r"{\rtf1\ansi\deff0"
    r"{\fonttbl{\f0\froman\fcharset0 Times New Roman;}}"
    r"\f0\fs24 "
)
_RTF_EPILOGUE = r"}"


def _rtf(body: str) -> str:
    return _RTF_PROLOGUE + body + _RTF_EPILOGUE


def _write_rtf(path: Path, body: str) -> None:
    path.write_text(_rtf(body), encoding="utf-8")


_TAXONOMY_PATH = Path(__file__).parent.parent / "spec" / "taxonomy" / "affiliation-agreement.yaml"

_CORPUS_BODY = (
    r"1. Indemnification\par "
    r"Alpha Corp shall indemnify Beta University against third-party claims "
    r"arising from the placement programme.\par "
    r"2. Governing Law\par "
    r"This agreement is governed by the laws of the State of California.\par "
    r"3. Term\par "
    r"This agreement commences on the date of execution and continues for one year.\par "
)

_TEMPLATE_BODY = (
    r"1. Indemnification\par "
    r"The service provider shall indemnify the institution against third-party claims.\par "
    r"2. Governing Law\par "
    r"This agreement is governed by the laws of the State of New York.\par "
    r"3. Term\par "
    r"Initial term of one year with automatic renewal.\par "
)


# ---------------------------------------------------------------------------
# Corpus + config factory
# ---------------------------------------------------------------------------


def _make_corpus_with_template(tmp_path: Path) -> tuple[Path, Path, Path, Path]:
    """Build a synthetic corpus + config WITH a template; return
    (corpus_dir, config_path, out_dir, template_path).
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "v1.rtf", _CORPUS_BODY)

    template_dir = tmp_path / "template"
    template_dir.mkdir()
    template_path = template_dir / "template.rtf"
    _write_rtf(template_path, _TEMPLATE_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
            "aliases": ["eiaa"],
        },
        "baseline": {"template": str(template_path)},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    out_dir = tmp_path / "out"
    return corpus_dir, config_path, out_dir, template_path


class _FakeDeviationJudge:
    """Deterministic fake ``DeviationJudge`` — always returns a fixed
    substantive/worse/material verdict and records every batch it is asked
    to assess, so a test can assert it was actually invoked (as opposed to
    a clause being silently hardcoded deviation="none" without ever reaching
    the judge — issue #103)."""

    def __init__(self) -> None:
        self.received_items: list[dict[str, str]] = []

    def assess_batch(self, items: list[dict[str, str]], our_standard: str) -> list[DeviationResult]:
        self.received_items.extend(items)
        return [
            DeviationResult(
                deviation="substantive",
                risk_delta=RiskDelta(direction="worse", magnitude="material"),
                basis="judge",
                rationale="Fake judge: clause differs materially from the template.",
            )
            for _ in items
        ]


# ---------------------------------------------------------------------------
# AC-1: project_playbook makes no _ingest_file and no classify_tree calls
# ---------------------------------------------------------------------------


def test_project_playbook_no_ingest_no_judge_with_template(tmp_path: Path) -> None:
    """project_playbook must not call _ingest_file or classify_tree, even
    when a template is configured in the engine config.

    Regression guard for issue #63: previously project_playbook re-ingested
    the template via _ingest_file and re-classified it via classify_tree,
    violating the 'zero ingest / LLM work' criterion.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    # Step 1: mine — real run, writes template_observations.jsonl to the store.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )
    assert (out_dir / "template_observations.jsonl").exists(), (
        "mine_corpus must write template_observations.jsonl"
    )

    # Step 2: project — spy on _ingest_file and classify_tree; assert neither is called.
    ingest_spy = MagicMock(side_effect=AssertionError("_ingest_file must not be called by project"))
    classify_spy = MagicMock(
        side_effect=AssertionError("classify_tree must not be called by project")
    )

    with (
        patch("playbook_engine.pipeline._ingest_file", ingest_spy),
        patch("playbook_engine.pipeline.classify_tree", classify_spy),
    ):
        playbook = project_playbook(
            out_dir=out_dir,
            config=cfg,
            taxonomy=taxonomy,
        )

    # Spies must not have been called at all.
    ingest_spy.assert_not_called()
    classify_spy.assert_not_called()

    # And the playbook must still be produced.
    assert isinstance(playbook, dict)
    assert (out_dir / "playbook.opf.json").exists()


def test_compiled_agreement_type_carries_id_and_aliases(tmp_path: Path) -> None:
    """Issue #142: agreement_type is the shared cross-tool key. The compiler
    must pass ``config.agreement_type.aliases`` straight through into the
    assembled playbook alongside ``id``, so a consumer (whose own registry
    key is 'eiaa') can match on either without a hand-joined
    mapping table."""
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    assert cfg.agreement_type.aliases == ["eiaa"]

    mine_corpus(corpus_dir=corpus_dir, config=cfg, taxonomy=taxonomy, out_dir=out_dir)
    playbook = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    assert playbook["agreement_type"]["id"] == "educational-affiliation"
    assert playbook["agreement_type"]["aliases"] == ["eiaa"]

    on_disk = json.loads((out_dir / "playbook.opf.json").read_text(encoding="utf-8"))
    assert on_disk["agreement_type"]["aliases"] == ["eiaa"]


# ---------------------------------------------------------------------------
# AC-2: Re-running project does not rewrite observations.jsonl
# ---------------------------------------------------------------------------


def test_project_playbook_does_not_rewrite_observations(tmp_path: Path) -> None:
    """Re-running project_playbook must not touch observations.jsonl.

    Ensures that the project step is purely read-only with respect to the
    observation store written by mine_corpus.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    # Mine first.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    obs_path = out_dir / "observations.jsonl"
    mtime_after_mine = obs_path.stat().st_mtime

    # Project — must not overwrite observations.jsonl.
    project_playbook(
        out_dir=out_dir,
        config=cfg,
        taxonomy=taxonomy,
    )

    mtime_after_project = obs_path.stat().st_mtime
    assert mtime_after_project == mtime_after_mine, (
        "project_playbook must not rewrite observations.jsonl"
    )


# ---------------------------------------------------------------------------
# Issue #101: a default (fully zero-LLM) compile must watermark its output.
# ---------------------------------------------------------------------------

_CORPUS_BODY_V2 = (
    r"1. Indemnification\par "
    r"Alpha Corp shall indemnify Beta University against third-party claims "
    r"arising from the placement programme.\par "
    r"2. Governing Law\par "
    r"This agreement is governed by the laws of the State of Delaware.\par "
    r"3. Term\par "
    r"This agreement commences on the date of execution and continues for one year.\par "
)


def test_default_stub_judges_compile_watermarks_playbook(tmp_path: Path) -> None:
    """End-to-end regression guard for issue #101.

    A real ``playbook compile`` with NO judges configured (the CLI default)
    uses ``_AllInScopeJudge`` (scope) and ``_NullDeviationJudge`` (deviation)
    — both stubs, neither backed by an LLM. Neither stub ever emits an
    Observation with basis="stub" literally: ``_AllInScopeJudge``'s
    basis="stub" lands on the ``ScopeDecision`` (scope.json), not an
    Observation, and ``_NullDeviationJudge`` emits basis="needs_review" (a
    judge protocol IS wired for deviation, it's just the stub default).
    Before the fix, neither signal reached ``assemble_playbook``, so a
    default compile's ``compiler.stub_basis_present`` was always False — the
    exact liability scenario issue #101 exists to catch (a stub-derived
    playbook a consumer cannot tell apart from a real one). This test
    drives the actual ``mine_corpus`` -> ``project_playbook`` path with no
    judges passed at all and asserts the watermark now fires.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)
    # Two versions with a real clause-text change (governing law: California
    # -> Delaware) so the deviation classifier actually dispatches to
    # _NullDeviationJudge for a changed clause — a single-version document
    # never calls the deviation judge at all (see
    # _observations_from_single_version), so this must be multi-version.
    _write_rtf(deal_dir / "v1.rtf", _CORPUS_BODY)
    _write_rtf(deal_dir / "v2.rtf", _CORPUS_BODY_V2)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    # No scope_judge / deviation_judge passed -> the pipeline's zero-LLM
    # defaults apply (_AllInScopeJudge, _NullDeviationJudge).
    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    # Sanity check the premise: the stub scope decision really did land on
    # scope.json (basis="stub"), not on any Observation.
    scope = json.loads((out_dir / "scope.json").read_text(encoding="utf-8"))
    assert scope["documents"][0]["basis"] == "stub"
    obs_lines = (out_dir / "observations.jsonl").read_text(encoding="utf-8").splitlines()
    observations = [json.loads(line) for line in obs_lines if line.strip()]
    assert all(obs.get("basis") != "stub" for obs in observations), (
        "no Observation should ever carry basis='stub' literally — the stub "
        "signal from _NullDeviationJudge is basis='needs_review'"
    )

    playbook = project_playbook(
        out_dir=out_dir,
        config=config,
        taxonomy=taxonomy,
    )

    assert playbook["compiler"]["stub_basis_present"] is True, (
        "a default (zero-LLM) compile must watermark its output so a "
        "consuming review application can refuse to run against it"
    )


# ---------------------------------------------------------------------------
# Issue #82: an empty ClauseTree from a non-empty file must be recorded as a
# per-version ingest failure, never fed to the scope gate as first_tree.
# ---------------------------------------------------------------------------


def test_empty_tree_ingest_is_recorded_failure_not_out_of_scope(tmp_path: Path) -> None:
    """A version whose ingest yields an EMPTY ClauseTree from a non-empty
    source file (e.g. a scanned/image PDF with no OCR wired on the
    deterministic path) must be treated as an ingest failure, not silently
    fed to the scope gate as ``first_tree``.

    Regression guard for issue #82: previously the empty tree entered
    ``version_trees`` as if ingestion had succeeded. If it was the
    alphabetically-first version, ``scope_gate`` saw zero clause nodes and
    classified the *entire agreement* out-of-scope with
    ``basis="deterministic_empty"`` — one unreadable version taking down an
    otherwise-valid negotiation trail.

    "v1.rtf" here is a well-formed but body-less RTF file (76 bytes of RTF
    control words, zero paragraph text) — deterministically produces an
    empty ClauseTree without any OCR/LLM mocking. "v2.rtf" carries real
    clause content and is alphabetically second, so before the fix it would
    never be reached: ``first_tree`` would already have been the empty v1.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)

    # v1: well-formed RTF, but no body text at all -> ingests to an empty tree.
    _write_rtf(deal_dir / "v1.rtf", "")
    assert (deal_dir / "v1.rtf").stat().st_size > 0, "fixture file must be non-empty"

    # v2: real clause content -> ingests to a valid, in-scope-able tree.
    _write_rtf(deal_dir / "v2.rtf", _CORPUS_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    progress_lines: list[str] = []
    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
        progress=progress_lines.append,
    )

    # The empty version must have produced a per-version WARNING, not a silent skip.
    assert any("v1.rtf" in line and "WARNING" in line for line in progress_lines), (
        f"expected a per-version WARNING for the empty-tree version v1.rtf; got: {progress_lines}"
    )

    # normalized/deal-001/v1.clauses.json must NOT have been written (ingest failure,
    # not success) -- only v2 (the valid version) reaches the normalized-tree write.
    normalized_dir = out_dir / "normalized" / "deal-001"
    assert not (normalized_dir / "v1.clauses.json").exists(), (
        "an empty-tree ingest failure must not be written as a normalized tree"
    )
    assert (normalized_dir / "v2.clauses.json").exists(), (
        "the valid version v2 must still be ingested and normalized"
    )

    # The scope gate must have evaluated v2 (real content), not v1 (empty) --
    # so the document must NOT be marked out-of-scope as "deterministic_empty".
    scope = json.loads((out_dir / "scope.json").read_text(encoding="utf-8"))
    doc_entries = {d["document_id"]: d for d in scope["documents"]}
    assert "deal-001" in doc_entries, "deal-001 must have a scope decision recorded"
    decision = doc_entries["deal-001"]
    assert decision["basis"] != "deterministic_empty", (
        "the empty-tree version must not have reached the scope gate as first_tree: "
        f"got basis={decision['basis']!r}"
    )
    assert decision["in_scope"] is True, (
        f"deal-001 has a valid version (v2) with real clause content and should remain "
        f"in-scope; got decision={decision}"
    )


# ---------------------------------------------------------------------------
# Issue #89: a failed per-version ingest must be a durable manifest record,
# not just a scrolled-past progress-line WARNING, and must surface as a
# Needs-Attention item in the after-action report.
# ---------------------------------------------------------------------------


def test_failed_version_ingest_recorded_in_manifest_and_needs_attention(
    tmp_path: Path,
) -> None:
    """A document with one failing version file and one good version file must:

    - record ``versions_mined`` < ``versions_found`` in corpus_manifest.json
      (not the old ``versions == len(version_files)`` that counted files
      found rather than versions actually mined).
    - carry a per-version ``version_ingest`` entry for the failed version
      with ``status="failed"`` and a non-empty ``error`` string.
    - surface that failure as a Needs Attention item in the after-action
      report, not just a console WARNING that a cache hit wouldn't re-print.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)

    # v1: well-formed RTF, no body text -> ingest failure (empty clause tree
    # from a non-empty source file; same fixture shape as issue #82's test).
    _write_rtf(deal_dir / "v1.rtf", "")
    assert (deal_dir / "v1.rtf").stat().st_size > 0, "fixture file must be non-empty"

    # v2: real clause content -> ingests successfully.
    _write_rtf(deal_dir / "v2.rtf", _CORPUS_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
        progress=lambda _: None,
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    doc_entries = {d["document_id"]: d for d in manifest}
    assert "deal-001" in doc_entries
    entry = doc_entries["deal-001"]

    assert entry["versions_found"] == 2, (
        f"expected 2 version files found (v1.rtf, v2.rtf); got {entry['versions_found']}"
    )
    assert entry["versions_mined"] == 1, (
        f"expected only v2 to have been mined (v1 failed); got {entry['versions_mined']}"
    )
    assert entry["versions_mined"] < entry["versions_found"], (
        "versions_mined must be strictly less than versions_found when a version fails"
    )
    # Back-compat alias: "versions" now means versions MINED, not files found.
    assert entry["versions"] == entry["versions_mined"]

    ingest_by_version = {v["version"]: v for v in entry["version_ingest"]}
    assert set(ingest_by_version) == {"v1", "v2"}
    assert ingest_by_version["v1"]["status"] == "failed"
    assert ingest_by_version["v1"]["error"], "failed version must carry a non-empty error string"
    assert ingest_by_version["v1"]["extractor"] == "rtf"
    assert ingest_by_version["v2"]["status"] == "ok"

    # The failure must surface in the after-action report's Needs Attention
    # section, not just as a console WARNING.
    from playbook_engine.aar import build_after_action_data

    aar_data = build_after_action_data(out_dir)
    needs_attention = aar_data["needs_attention"]
    matches = [
        item
        for item in needs_attention
        if item["document_id"] == "deal-001" and item["version"] == "v1"
    ]
    assert matches, f"expected a Needs-Attention item for deal-001/v1; got {needs_attention}"
    assert any("version ingest failed" in r for r in matches[0]["reasons"])


# ---------------------------------------------------------------------------
# Issue #83: a trail with no detected signed copy must not fabricate
# signed_copy_confidence or outcome="signed".
# ---------------------------------------------------------------------------


def test_unsigned_trail_no_fabricated_confidence_or_outcome(tmp_path: Path) -> None:
    """A document where no version is detected as signed must record an
    honest 'no signed copy' trail state, not a fabricated one.

    Regression guard for issue #83: previously
    ``version_order.signed_id or ordered_ids[-1]`` silently treated the
    last-in-chain draft as the signed copy, reporting a signed=False
    determination's confidence (e.g. 0.85 for basis="no_signature_section")
    as if it were confidence in the signed copy, and every non-reversed
    observation was stamped ``outcome="signed"`` regardless of whether any
    execution evidence existed.

    Neither ``_CORPUS_BODY`` version below contains a signature section, so
    ``detect_signed`` deterministically returns ``signed=False`` for both
    (no judge/LLM required) and ``order_versions`` never anchors a signed_id.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "v1.rtf", _CORPUS_BODY)
    _write_rtf(deal_dir / "v2.rtf", _CORPUS_BODY.replace("one year", "two years"))

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    trail = json.loads((out_dir / "trail" / "deal-001.json").read_text(encoding="utf-8"))
    assert trail["signed_version"] is None, (
        f"no version should be identified as signed; got {trail['signed_version']!r}"
    )
    assert trail["signed_copy_confidence"] is None, (
        "signed_copy_confidence must be None when no signed copy was detected "
        f"(fabricated from a fallback version otherwise); got "
        f"{trail['signed_copy_confidence']!r}"
    )

    obs_lines = (out_dir / "observations.jsonl").read_text(encoding="utf-8").splitlines()
    observations = [json.loads(line) for line in obs_lines if line.strip()]
    deal_obs = [o for o in observations if o["citation"]["document_id"] == "deal-001"]
    assert deal_obs, "deal-001 must have observations"
    assert all(o["outcome"] != "signed" for o in deal_obs), (
        "no observation may carry outcome='signed' when no signed copy was detected: "
        f"{[o['outcome'] for o in deal_obs]}"
    )


# ---------------------------------------------------------------------------
# Issue #84: hints.yaml written per docs/CORPUS-LAYOUT.md's documented
# example (entries WITH file extensions) must still anchor the intended
# version, whose version_id is a bare file STEM.
# ---------------------------------------------------------------------------


def test_hints_signed_version_with_extension_anchors_stem(tmp_path: Path) -> None:
    """A hints.yaml naming ``fully-executed.pdf`` (extension included, as
    docs/CORPUS-LAYOUT.md documents) must anchor the version whose
    version_id (file stem) is ``fully-executed``.

    Regression guard for issue #84: version ids are file stems (``vf.stem``
    — pipeline.py's per-version loop), so a hint value carrying an
    extension previously never matched any real version_id and the
    signed_version override silently had no effect. Neither version's body
    here contains a signature section, so — mirroring
    ``test_unsigned_trail_no_fabricated_confidence_or_outcome`` above —
    ``detect_signed`` deterministically returns ``signed=False`` for both;
    the ONLY way ``trail["signed_version"]`` can end up set is the hint
    actually taking effect.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "draft-we-sent.rtf", _CORPUS_BODY)
    _write_rtf(deal_dir / "fully-executed.rtf", _CORPUS_BODY.replace("one year", "two years"))

    # As documented in CORPUS-LAYOUT.md's hints.yaml example, entries carry
    # extensions -- even though the file actually on disk here is .rtf.
    (deal_dir / "hints.yaml").write_text("signed_version: fully-executed.pdf\n", encoding="utf-8")

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    trail = json.loads((out_dir / "trail" / "deal-001.json").read_text(encoding="utf-8"))
    assert trail["signed_version"] == "fully-executed", (
        "hints.yaml's signed_version (given WITH an extension, as documented) "
        f"must anchor the version whose stem is 'fully-executed'; got {trail['signed_version']!r}"
    )
    assert trail["signed_copy_confidence"] == 1.0, (
        "a hint-driven signed_version override must report full confidence; "
        f"got {trail['signed_copy_confidence']!r}"
    )


# ---------------------------------------------------------------------------
# Issue #88: DOCX tracked-changes author/date attribution must reach the
# observation store, not be discarded as an unconsumed side-channel.
# ---------------------------------------------------------------------------


def _docx_no_tracked_changes(tmp_path: Path, filename: str) -> Path:
    """Baseline DOCX version, no tracked changes.

    Two headings — scope_gate.MIN_CLAUSE_COUNT requires at least 2 clause
    nodes or the document is rejected as "too short" before L4 ever runs.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)
    doc.add_paragraph("Party A shall provide services to client.")
    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph("This agreement is governed by the laws of the State of California.")
    path = tmp_path / filename
    doc.save(str(path))
    return path


def _docx_with_tracked_insertion(tmp_path: Path, filename: str) -> Path:
    """Redlined DOCX version: 'promptly' inserted by Alice via w:ins.

    SECURITY NOTE: synthetic text and a fictional author name only, matching
    the tracked-changes fixture convention in tests/test_docx_ingester.py.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)
    p = doc.add_paragraph()
    p.add_run("Party A shall ")

    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "Alice")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "promptly "

    p.add_run("provide services to client.")
    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph("This agreement is governed by the laws of the State of California.")

    path = tmp_path / filename
    doc.save(str(path))
    return path


def test_docx_redline_observation_carries_tracked_changes_attribution(tmp_path: Path) -> None:
    """A clause changed via a tracked ``w:ins`` must carry author/date attribution
    on its observation — regression guard for issue #88.

    Previously ``_ingest_file`` returned only ``DocxIngestResult.tree``,
    discarding ``.tracked`` entirely: no observation could ever carry
    tracked-changes attribution, and ``tracked_changes_overlay.py`` had no
    caller in the pipeline (dead code, per the audit finding). This wires
    ``DocxIngestResult.tracked`` through ``_compute_doc_result`` into
    ``build_observations`` via ``tracked_changes_overlay.enrich_clause_diff``.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-002"
    deal_dir.mkdir(parents=True)
    _docx_no_tracked_changes(deal_dir, "v1.docx")
    _docx_with_tracked_insertion(deal_dir, "v2.docx")

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    obs_lines = (out_dir / "observations.jsonl").read_text(encoding="utf-8").splitlines()
    observations = [json.loads(line) for line in obs_lines if line.strip()]
    deal_obs = [o for o in observations if o["citation"]["document_id"] == "deal-002"]
    assert deal_obs, "deal-002 must have observations"

    attributed = [o for o in deal_obs if o.get("attribution") is not None]
    assert attributed, (
        f"expected at least one observation with tracked-changes attribution; "
        f"got attributions={[o.get('attribution') for o in deal_obs]}"
    )
    assert attributed[0]["attribution"] == {
        "author": "Alice",
        "date": "2024-03-15T10:00:00Z",
        "tracked_type": "insertion",
    }


# ---------------------------------------------------------------------------
# Issue #103: a single-version document's clauses must be diffed against the
# canonical template, not hardcoded deviation="none".
# ---------------------------------------------------------------------------


def test_single_version_clause_diffed_against_template(tmp_path: Path) -> None:
    """A single-version document with a signed clause that differs from the
    canonical template must route that clause through the deviation judge,
    not hardcode deviation="none"/basis="deterministic".

    Regression guard for issue #103: previously
    ``_observations_from_single_version`` recorded EVERY clause of a
    single-version document as deviation="none" unconditionally — a document
    with no negotiation trail was never actually checked against the
    template at all. ``_make_corpus_with_template`` gives a single-version
    document (``deal-001/v1.rtf``) whose clause texts (``_CORPUS_BODY``)
    differ materially from the configured template's (``_TEMPLATE_BODY`` —
    different party names, different governing-law state, different term
    language), so every classified clause should be routed to the injected
    fake judge.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    fake_judge = _FakeDeviationJudge()

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        deviation_judge=fake_judge,
    )

    assert fake_judge.received_items, (
        "the deviation judge must be invoked for a single-version document "
        "whose clauses differ from the canonical template — it was never "
        "reached at all before issue #103"
    )

    obs_lines = (out_dir / "observations.jsonl").read_text(encoding="utf-8").splitlines()
    observations = [json.loads(line) for line in obs_lines if line.strip()]
    deal_obs = [o for o in observations if o["citation"]["document_id"] == "deal-001"]
    assert deal_obs, "deal-001 must have observations"

    judged = [o for o in deal_obs if o["basis"] == "judge"]
    assert judged, (
        f"expected at least one observation routed through the injected "
        f"deviation judge (basis='judge'); got bases={[o['basis'] for o in deal_obs]}"
    )
    assert all(o["deviation"] == "substantive" for o in judged), (
        "the injected fake judge's verdict must reach the observation, not a "
        "hardcoded deviation='none'"
    )


# ---------------------------------------------------------------------------
# Issue #147: embedded attorney-pinned positions (curation overlay) survive a
# recompile and are conflict-flagged when fresh evidence contradicts them.
# ---------------------------------------------------------------------------


def _inject_pin(out_dir: Path, clause_id: str, position: str, baseline_stance: str) -> None:
    """Simulate an attorney pin having been embedded in a PRIOR compile.

    Writes ``curation.pins`` directly into ``playbook.opf.json`` — mirroring
    what ``viewer.apply_feedback``'s ``override`` handling does — so the next
    ``project_playbook`` call has a prior curation section to merge against.
    """
    opf_path = out_dir / "playbook.opf.json"
    doc = json.loads(opf_path.read_text(encoding="utf-8"))
    doc["curation"] = {
        "pins": [
            {
                "clause_id": clause_id,
                "item_id": "C1",
                "position": position,
                "baseline_stance": baseline_stance,
                "pinned_at": "2026-01-01T00:00:00Z",
            }
        ]
    }
    opf_path.write_text(json.dumps(doc), encoding="utf-8")


def _add_our_paper_observations(out_dir: Path, taxonomy_id: str, n: int) -> None:
    """Append *n* synthetic our-paper, deviation=none, signed observations for
    *taxonomy_id* to the observation store, and register their (synthetic)
    document ids in ``corpus_manifest.json`` so citation validation resolves.

    This is how the tests below simulate "new evidence was mined" between two
    ``project_playbook`` calls, without invoking the full ingest/judge
    pipeline (project_playbook's own contract is "given the store, recompile"
    — manipulating the store directly is a legitimate way to exercise that).
    """
    manifest_path = out_dir / "corpus_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    obs_path = out_dir / "observations.jsonl"

    new_obs: list[Observation] = []
    with obs_path.open("a", encoding="utf-8") as f:
        for i in range(1, n + 1):
            doc_id = f"synthetic-our-paper-{i}"
            manifest.append(
                {
                    "document_id": doc_id,
                    "provenance": "our_paper",
                    "in_scope": True,
                    "versions": 1,
                    "versions_mined": 1,
                    "versions_found": 1,
                    "version_ingest": [
                        {"version": "v1", "status": "ok", "error": None, "extractor": "rtf"}
                    ],
                }
            )
            obs = Observation(
                observation_id=f"{doc_id}/1/2",
                taxonomy_id=taxonomy_id,
                text_summary="Synthetic matching clause text.",
                citation=ObservationCitation(
                    document_id=doc_id, version=1, clause_path="2", char_span=(0, 10)
                ),
                deviation="none",
                risk_delta={"direction": "neutral", "magnitude": "none"},
                provenance="our_paper",
                outcome="signed",
                confidence=0.9,
                basis="judge",
            )
            new_obs.append(obs)
            f.write(json.dumps(obs.to_dict()) + "\n")

    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")


def test_curation_pin_survives_recompile_with_unchanged_evidence(tmp_path: Path) -> None:
    """A pinned position survives a recompile when evidence is unchanged:
    the pin is preserved verbatim and no conflict is raised.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(corpus_dir=corpus_dir, config=cfg, taxonomy=taxonomy, out_dir=out_dir)
    playbook_v1 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    clause = next(
        c for c in playbook_v1["evidence"]["clauses"] if c["taxonomy_id"] == "governing_law"
    )
    baseline_stance = clause["summary"]["historical_stance"]
    _inject_pin(
        out_dir,
        clause_id=clause["id"],
        position="consistently_held",
        baseline_stance=baseline_stance,
    )

    # Recompile with NO change to the observation store.
    playbook_v2 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    pins = playbook_v2["curation"]["pins"]
    assert len(pins) == 1
    pin = pins[0]
    assert pin["clause_id"] == clause["id"]
    assert pin["position"] == "consistently_held", "the pinned position must be preserved verbatim"
    assert pin.get("conflict") is None, (
        "no conflict should be raised when evidence (and thus historical_stance) is unchanged"
    )

    clause_v2 = next(
        c for c in playbook_v2["evidence"]["clauses"] if c["taxonomy_id"] == "governing_law"
    )
    assert clause_v2["summary"]["historical_stance"] == baseline_stance, (
        "historical_stance itself must remain descriptive/unchanged — the pin does not overwrite it"
    )


def test_curation_pin_flags_conflict_on_contradicting_evidence(tmp_path: Path) -> None:
    """New evidence that changes a clause's historical_stance since a pin was
    made raises a conflict flag on that pin — the pin is NOT silently
    overridden or dropped.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(corpus_dir=corpus_dir, config=cfg, taxonomy=taxonomy, out_dir=out_dir)
    playbook_v1 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    clause = next(
        c for c in playbook_v1["evidence"]["clauses"] if c["taxonomy_id"] == "governing_law"
    )
    baseline_stance = clause["summary"]["historical_stance"]
    assert baseline_stance == "no_signal", (
        "premise: a single-document corpus has insufficient our-paper evidence "
        f"(see clause_position_compiler.MIN_EVIDENCE_N); got {baseline_stance!r}"
    )
    _inject_pin(
        out_dir,
        clause_id=clause["id"],
        position="consistently_held",
        baseline_stance=baseline_stance,
    )

    # Simulate new evidence: enough our-paper, deviation=none, signed
    # observations to push governing_law's historical_stance away from
    # "no_signal" (evidence_sufficient flips True; see
    # clause_position_compiler._historical_stance).
    _add_our_paper_observations(out_dir, taxonomy_id="governing_law", n=2)

    playbook_v2 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    clause_v2 = next(
        c for c in playbook_v2["evidence"]["clauses"] if c["taxonomy_id"] == "governing_law"
    )
    recomputed_stance = clause_v2["summary"]["historical_stance"]
    assert recomputed_stance != baseline_stance, (
        "premise: the injected evidence must actually move historical_stance"
    )

    pin = playbook_v2["curation"]["pins"][0]
    assert pin["position"] == "consistently_held", (
        "the pin's asserted position must survive even when flagged — it is "
        "not silently overridden by the recomputed rollup"
    )
    conflict = pin.get("conflict")
    assert conflict is not None, "contradicting evidence must raise a conflict flag"
    assert conflict["recomputed_historical_stance"] == recomputed_stance
    assert baseline_stance in conflict["reason"]


def test_curation_pin_excluded_from_content_hash_but_digested_separately(tmp_path: Path) -> None:
    """Adding/updating a curation pin must not perturb identity.content_hash
    (curation is excluded — issue #147), but identity.section_digests.curation
    must change to track the pin content.
    """
    corpus_dir, config_path, out_dir, _template_path = _make_corpus_with_template(tmp_path)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(corpus_dir=corpus_dir, config=cfg, taxonomy=taxonomy, out_dir=out_dir)
    playbook_v1 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)
    hash_v1 = content_hash(playbook_v1)
    assert hash_v1 == playbook_v1["identity"]["content_hash"]
    curation_digest_v1 = playbook_v1["identity"]["section_digests"]["curation"]

    clause = next(
        c for c in playbook_v1["evidence"]["clauses"] if c["taxonomy_id"] == "governing_law"
    )
    _inject_pin(
        out_dir,
        clause_id=clause["id"],
        position="consistently_held",
        baseline_stance=clause["summary"]["historical_stance"],
    )

    playbook_v2 = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    assert playbook_v2["identity"]["content_hash"] == hash_v1, (
        "a curation-only change (no evidence change) must not perturb content_hash"
    )
    assert playbook_v2["identity"]["section_digests"]["curation"] != curation_digest_v1, (
        "the curation section digest must change once a pin is present"
    )
    # Sanity: the digest really is computed the same way canonicalize.py defines it.
    assert playbook_v2["identity"]["section_digests"] == compute_section_digests(playbook_v2)
