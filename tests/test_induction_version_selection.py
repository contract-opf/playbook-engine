"""Tests for content-based representative-version selection in induce-taxonomy.

SECURITY NOTE: All fixtures are synthetic DOCX/RTF files built programmatically
(python-docx / raw RTF markup). No real agreement files are referenced. Party
names use fictional identifiers only ("Alice Smith", "Bob Jones").

Issue #169: ``induce-taxonomy`` used to pick a document's representative
version by filename natural-sort, contradicting the engine's own
content-based "filenames don't matter" contract
(``docs/CORPUS-LAYOUT.md:23``). These tests exercise the CLI end-to-end to
confirm the fix: signed-copy detection wins first, the edit-distance chain
terminal is the fallback, and filename ordering is used only as a
last-resort tiebreak when content evidence is genuinely tied.
"""

from __future__ import annotations

from pathlib import Path

import yaml
from click.testing import CliRunner
from docx import Document

from playbook_engine.cli import cli

# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def _docx_doc(path: Path, sections: list[tuple[str, str]], *, signed: bool = False) -> None:
    """Write a synthetic DOCX with (heading, body) sections.

    When ``signed`` is True, appends a "Signatures" section with two filled
    "By:" lines — enough for ``signed_detector.detect_signed`` to return
    ``basis="dual_signatures"`` (confidence 0.90), unambiguously signed.
    """
    doc = Document()
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    if signed:
        doc.add_heading("Signatures", level=1)
        doc.add_paragraph("By: Alice Smith")
        doc.add_paragraph("By: Bob Jones")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _rtf_doc(path: Path, headings: list[str]) -> None:
    """Write a minimal RTF file with numbered headings (no signature cues)."""
    parts: list[str] = []
    for i, heading in enumerate(headings, start=1):
        parts.append(rf"{i}. {heading}\par Clause text for {heading}.\par ")
    content = r"{\rtf1\ansi " + "".join(parts) + r"}"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


# ---------------------------------------------------------------------------
# 1. Content-based signed detection beats filename sort
# ---------------------------------------------------------------------------


def test_content_beats_filename(tmp_path: Path) -> None:
    """A content-signed draft named to sort FIRST must still win over an
    unsigned draft whose filename would naturally sort LAST (issue #169)."""
    corpus = tmp_path / "corpus"
    doc_dir = corpus / "deal-alpha"

    # Filename-highest ("z-final-v9") is actually an unsigned draft.
    _docx_doc(
        doc_dir / "z-final-v9.docx",
        [("Governing Law", "This Agreement is governed by the laws of Delaware.")],
        signed=False,
    )
    # Filename-lowest ("a-first-draft") is actually the executed copy.
    _docx_doc(
        doc_dir / "a-first-draft.docx",
        [("Indemnification", "Each party shall indemnify the other against losses.")],
        signed=True,
    )

    out_yaml = tmp_path / "candidate.yaml"
    runner = CliRunner()
    result = runner.invoke(cli, ["induce-taxonomy", str(corpus), "--out", str(out_yaml)])
    assert result.exit_code == 0, result.output

    data = yaml.safe_load(out_yaml.read_text(encoding="utf-8"))
    ids = {e["id"] for e in data["entries"]}
    assert any("indemnif" in eid for eid in ids), (
        f"Expected the SIGNED a-first-draft.docx to be selected (indemnification), got: {ids}"
    )
    assert not any("governing" in eid for eid in ids), (
        "Unsigned z-final-v9.docx must not be picked over the signed a-first-draft.docx "
        "merely because its filename sorts higher"
    )
    assert "basis=signed" in result.output


# ---------------------------------------------------------------------------
# 2. Filename is a last-resort tiebreak only, when content is genuinely tied
# ---------------------------------------------------------------------------


def test_filename_tiebreak_only(tmp_path: Path) -> None:
    """Two versions with no signed cues and no content-derived terminal
    (edit distance is symmetric for exactly two versions) must fall back to
    filename natural-sort, and the CLI must log that basis."""
    corpus = tmp_path / "corpus"
    doc_dir = corpus / "deal-beta"

    # Neither file has any signature content. With only two versions the
    # edit-distance chain is symmetric (a->b costs the same as b->a), so
    # there is no content-derived terminal — this is a genuine tie.
    _rtf_doc(doc_dir / "a-first.rtf", ["Confidentiality"])
    _rtf_doc(doc_dir / "b-second.rtf", ["Term"])

    out_yaml = tmp_path / "candidate.yaml"
    runner = CliRunner()
    result = runner.invoke(cli, ["induce-taxonomy", str(corpus), "--out", str(out_yaml)])
    assert result.exit_code == 0, result.output

    data = yaml.safe_load(out_yaml.read_text(encoding="utf-8"))
    ids = {e["id"] for e in data["entries"]}
    # Natural sort: "b-second" > "a-first" -> b-second.rtf wins the tiebreak.
    assert any("term" in eid for eid in ids), f"Expected b-second.rtf (term) selected, got: {ids}"
    assert not any("confidential" in eid for eid in ids), (
        "a-first.rtf must not be selected once b-second.rtf wins the filename tiebreak"
    )
    assert "basis=filename_tiebreak" in result.output


# ---------------------------------------------------------------------------
# 3. Existing (single-version-per-agreement) induction output is unchanged
# ---------------------------------------------------------------------------


def test_existing_induction_output_stable(tmp_path: Path) -> None:
    """On a corpus where names and content already agree (one version per
    agreement, nothing to arbitrate), induction output is unchanged."""
    corpus = tmp_path / "corpus"
    _rtf_doc(
        corpus / "deal-alpha" / "v1.rtf",
        ["Indemnification", "Governing Law", "Term"],
    )
    _rtf_doc(
        corpus / "deal-beta" / "v1.rtf",
        ["Indemnification", "Governing Law", "Termination"],
    )

    out_yaml = tmp_path / "candidate.yaml"
    runner = CliRunner()
    result = runner.invoke(cli, ["induce-taxonomy", str(corpus), "--out", str(out_yaml)])
    assert result.exit_code == 0, result.output

    data = yaml.safe_load(out_yaml.read_text(encoding="utf-8"))
    ids = {e["id"] for e in data["entries"]}
    assert any("indemnif" in eid for eid in ids), f"indemnification missing from {ids}"
    assert any("governing" in eid for eid in ids), f"governing_law missing from {ids}"
    # Single-version agreements need no arbitration, so no basis line is logged.
    assert "basis=" not in result.output
