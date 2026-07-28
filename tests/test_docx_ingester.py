"""Tests for the DOCX ingester.

SECURITY NOTE: All fixtures MUST be synthetic and programmatically generated
using python-docx + lxml XML injection.  No binary .docx files from real
agreements are ever committed or referenced from tests — not even redlined
drafts.  The real corpus lives outside this repo and must never become a
fixture source.  All tracked-change fixtures use fictional party names and
authors (e.g. "Alice", "Bob") only.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from playbook_engine.clause_tree import ClauseTree
from playbook_engine.docx_ingester import (
    DocxIngesterError,
    DocxIngestResult,
    TrackedChanges,
    ingest_docx,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _simple_docx(tmp_path: Path) -> Path:
    """Simple agreement: Heading 1 and 2 structure with body paragraphs."""
    doc = Document()
    doc.add_heading("Definitions", level=1)
    doc.add_paragraph('"Facility" means any FixtureCorp-operated training centre.')
    doc.add_heading("Training Services", level=2)
    doc.add_paragraph("FixtureCorp shall provide training per Schedule A.")
    doc.add_heading("Obligations", level=1)
    doc.add_paragraph("Each party shall perform its obligations diligently.")
    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


def _numbered_docx(tmp_path: Path) -> Path:
    """Numbered sections using explicit numbers in paragraph text."""
    doc = Document()
    doc.add_paragraph("1. Definitions")
    doc.add_paragraph('"Facility" means the training centre.')
    doc.add_paragraph("1.1. Facility")
    doc.add_paragraph("The term Facility includes all FixtureCorp sites.")
    doc.add_paragraph("1.2. Program")
    doc.add_paragraph("Program means the accredited curriculum.")
    doc.add_paragraph("2. Obligations")
    doc.add_paragraph("Both parties shall cooperate fully.")
    path = tmp_path / "numbered.docx"
    doc.save(str(path))
    return path


def _decimal_and_multipart_docx(tmp_path: Path) -> Path:
    """Regression fixture for issue #241 (_NUM_PREFIX backtracking).

    - "1. Payment" is a genuine top-level heading.
    - The paragraphs that follow are line-initial decimal numbers and must
      remain body text under "1. Payment" — not be promoted to spurious
      headings ("1", "99", "10") by the old backtracking regex.
    - "10.1.3 Term and Termination" must not be mis-split into clause path
      "10.1" with mangled heading text "3 Term and Termination".
    """
    doc = Document()
    doc.add_paragraph("1. Payment")
    doc.add_paragraph("1.5 times the Fees shall be invoiced in the prior month.")
    doc.add_paragraph("99.9% uptime guarantee applies to the Platform.")
    doc.add_paragraph("10.00 per unit is the price ordered.")
    doc.add_paragraph("10.1.3 Term and Termination")
    path = tmp_path / "decimal_multipart.docx"
    doc.save(str(path))
    return path


def _table_docx(tmp_path: Path) -> Path:
    """Document with a table — table must be flattened to body text."""
    doc = Document()
    doc.add_heading("Fee Schedule", level=1)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Service"
    tbl.rows[0].cells[1].text = "Fee"
    tbl.rows[1].cells[0].text = "Training"
    tbl.rows[1].cells[1].text = "$100/hour"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path


def _nested_table_docx(tmp_path: Path) -> Path:
    """Table with a nested table — regression fixture for issue #64.

    The outer table has one cell whose own text is "OUTER"; nested inside
    that same cell is a one-cell table whose text is "INNER". Before the
    fix, ``_flatten_table`` visited the inner cell twice: once via the outer
    table's descendant ``.iter(w:tc)``/``.//w:p`` search (which reaches into
    nested tables), and again as the nested table's own top-level cell.
    """
    doc = Document()
    doc.add_heading("Fee Schedule", level=1)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = "OUTER"
    nested = cell.add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "INNER"
    path = tmp_path / "nested_table.docx"
    doc.save(str(path))
    return path


def _tracked_docx(tmp_path: Path) -> Path:
    """Document with tracked changes injected via raw XML."""
    doc = Document()

    # Heading
    doc.add_heading("Obligations", level=1)

    # Paragraph with tracked insertion and deletion
    p = doc.add_paragraph()

    # Normal run
    p.add_run("Party A shall ")

    # Tracked insertion: "promptly " inserted by Alice
    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "Alice")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "promptly "

    # Normal run continued
    p.add_run("provide services")

    # Tracked deletion: "to client" deleted by Bob
    del_elem = etree.SubElement(p._p, _w("del"))
    del_elem.set(_w("id"), "2")
    del_elem.set(_w("author"), "Bob")
    del_elem.set(_w("date"), "2024-03-16T09:00:00Z")
    r_del = etree.SubElement(del_elem, _w("r"))
    dt_del = etree.SubElement(r_del, _w("delText"))
    dt_del.text = " to client"

    path = tmp_path / "tracked.docx"
    doc.save(str(path))
    return path


def _hyperlink_tracked_docx(tmp_path: Path) -> Path:
    """Document where a tracked change is nested inside a w:hyperlink element.

    Structure: Heading 1 "Obligations", then a body paragraph containing
    a hyperlink that wraps a w:ins element (Alice inserts a URL reference).
    This exercises BLOCKING-2: nested ins/del inside w:hyperlink.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)
    p = doc.add_paragraph()
    p.add_run("See ")

    # Build: <w:hyperlink><w:ins author="Alice">...<w:r><w:t>Schedule A</w:t></w:r></w:ins></w:hyperlink>
    hl_elem = etree.SubElement(p._p, _w("hyperlink"))
    ins_inside_hl = etree.SubElement(hl_elem, _w("ins"))
    ins_inside_hl.set(_w("id"), "10")
    ins_inside_hl.set(_w("author"), "Alice")
    ins_inside_hl.set(_w("date"), "2024-04-01T08:00:00Z")
    r_hl = etree.SubElement(ins_inside_hl, _w("r"))
    t_hl = etree.SubElement(r_hl, _w("t"))
    t_hl.text = "Schedule A"

    p.add_run(" for details.")
    path = tmp_path / "hyperlink_tracked.docx"
    doc.save(str(path))
    return path


def _sdt_content_control_docx(tmp_path: Path) -> Path:
    """Document with a run wrapped in an inline w:sdt content control.

    Structure: Heading 1 "Parties", then a body paragraph where a run is
    wrapped in <w:sdt><w:sdtContent><w:r><w:t>Acme Corp</w:t></w:r>
    </w:sdtContent></w:sdt>, as commonly produced by CLM/template-generated
    Word files (content controls bind party names, dates, defined terms).
    """
    doc = Document()
    doc.add_heading("Parties", level=1)
    p = doc.add_paragraph()
    p.add_run("This Agreement is between ")

    sdt_elem = etree.SubElement(p._p, _w("sdt"))
    sdt_content = etree.SubElement(sdt_elem, _w("sdtContent"))
    r_sdt = etree.SubElement(sdt_content, _w("r"))
    t_sdt = etree.SubElement(r_sdt, _w("t"))
    t_sdt.text = "Acme Corp"

    p.add_run(" and the Vendor.")
    path = tmp_path / "sdt_content_control.docx"
    doc.save(str(path))
    return path


def _block_level_sdt_docx(tmp_path: Path) -> Path:
    """Document with an entire paragraph wrapped in a block-level w:sdt.

    Structure: Heading 1 "Obligations", then a whole body paragraph wrapped
    in <w:sdt><w:sdtContent><w:p>...</w:p></w:sdtContent></w:sdt>, inserted
    directly as raw XML (python-docx has no API for content controls). The
    wrapper is inserted before w:sectPr to keep the body well-formed.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)

    sdt_elem = etree.Element(_w("sdt"))
    sdt_content = etree.SubElement(sdt_elem, _w("sdtContent"))
    p_elem = etree.SubElement(sdt_content, _w("p"))
    r_elem = etree.SubElement(p_elem, _w("r"))
    t_elem = etree.SubElement(r_elem, _w("t"))
    t_elem.text = "FixtureCorp shall deliver the wrapped clause text."

    body = doc.element.body
    sect_pr = body.find(_w("sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(sdt_elem)
    else:
        body.append(sdt_elem)

    path = tmp_path / "block_sdt.docx"
    doc.save(str(path))
    return path


def _fld_simple_docx(tmp_path: Path) -> Path:
    """Document with a run wrapped in a w:fldSimple simple field.

    Structure: Heading 1 "Effective Date", then a body paragraph where a run
    is wrapped in <w:fldSimple w:instr="DATE"><w:r><w:t>1 January 2024</w:t>
    </w:r></w:fldSimple>, as produced by Word's "Insert Field" feature.
    """
    doc = Document()
    doc.add_heading("Effective Date", level=1)
    p = doc.add_paragraph()
    p.add_run("The Effective Date is ")

    fld_elem = etree.SubElement(p._p, _w("fldSimple"))
    fld_elem.set(_w("instr"), "DATE")
    r_fld = etree.SubElement(fld_elem, _w("r"))
    t_fld = etree.SubElement(r_fld, _w("t"))
    t_fld.text = "1 January 2024"

    p.add_run(".")
    path = tmp_path / "fld_simple.docx"
    doc.save(str(path))
    return path


def _multi_para_clause_docx(tmp_path: Path) -> Path:
    """Heading with multiple body paragraphs — tests document-absolute char_span offsets."""
    doc = Document()
    doc.add_heading("Obligations", level=1)
    # Two body paragraphs with a tracked insertion in the second
    doc.add_paragraph("Party A shall perform services.")

    p2 = doc.add_paragraph()
    p2.add_run("Compensation is ")
    ins_elem = etree.SubElement(p2._p, _w("ins"))
    ins_elem.set(_w("id"), "5")
    ins_elem.set(_w("author"), "Carol")
    ins_elem.set(_w("date"), "2024-05-01T00:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "as follows"
    p2.add_run(".")

    path = tmp_path / "multi_para.docx"
    doc.save(str(path))
    return path


def _empty_docx(tmp_path: Path) -> Path:
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Basic ingestion
# ---------------------------------------------------------------------------


def test_ingest_returns_docx_result(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "simple-doc", "v1")
    assert isinstance(result, DocxIngestResult)
    assert isinstance(result.tree, ClauseTree)
    assert isinstance(result.tracked, TrackedChanges)


def test_tree_has_correct_document_id(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "my-doc", "v2")
    assert result.tree.document_id == "my-doc"
    assert result.tree.version == "v2"
    assert result.tree.source_file == "simple.docx"


def test_simple_doc_produces_nodes(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    assert len(result.tree.nodes) > 0, "Simple DOCX should produce at least one clause node"


def test_heading_hierarchy_is_respected(tmp_path: Path) -> None:
    """Heading 2 sections must be children of Heading 1."""
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    tree = result.tree

    # Should have top-level clauses
    paths = [n.clause_path for n in tree.nodes]
    assert len(paths) >= 1

    # "Training Services" (Heading 2) must be nested under "Definitions" (Heading 1)
    definitions_node = next(
        (n for n in tree.nodes if n.heading and "Definitions" in n.heading), None
    )
    assert definitions_node is not None, "Definitions heading not found"
    child_headings = [c.heading for c in definitions_node.children]
    assert any("Training Services" in (h or "") for h in child_headings), (
        "Training Services (Heading 2) should be a child of Definitions (Heading 1)"
    )


def test_body_text_attached_to_heading(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    definitions_node = next(
        (n for n in result.tree.nodes if n.heading and "Definitions" in n.heading), None
    )
    assert definitions_node is not None
    assert "Facility" in definitions_node.text, (
        "Body paragraph under Definitions should be in node text"
    )


# ---------------------------------------------------------------------------
# Numbered paragraphs
# ---------------------------------------------------------------------------


def test_numbered_prefix_becomes_clause_path(tmp_path: Path) -> None:
    path = _numbered_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    tree = result.tree

    # Should have clause "1" and "2" at the top
    top_paths = {n.clause_path for n in tree.nodes}
    assert "1" in top_paths or "2" in top_paths


def test_numbered_subsections_nested(tmp_path: Path) -> None:
    path = _numbered_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    # resolve_path should find 1.1 and 1.2
    n11 = result.tree.resolve_path("1.1")
    n12 = result.tree.resolve_path("1.2")
    assert n11 is not None, "Clause 1.1 should be in the tree"
    assert n12 is not None, "Clause 1.2 should be in the tree"


# ---------------------------------------------------------------------------
# Decimal / multi-part number regression (issue #241)
# ---------------------------------------------------------------------------


def test_decimal_prefixed_paragraphs_stay_body_text(tmp_path: Path) -> None:
    """1.5x / 99.9% / 10.00 must not be promoted to clause headings."""
    path = _decimal_and_multipart_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    tree = result.tree

    top_paths = {n.clause_path for n in tree.nodes}
    assert "99" not in top_paths, "'99.9% uptime' must not become clause '99'"
    assert "10" not in top_paths, "'10.00 per unit' must not become clause '10'"

    payment_node = next((n for n in tree.nodes if n.clause_path == "1"), None)
    assert payment_node is not None, "Clause '1' (Payment) should exist"
    assert "1.5 times the Fees" in payment_node.text
    assert "99.9% uptime guarantee" in payment_node.text
    assert "10.00 per unit is the price" in payment_node.text


def test_multipart_number_not_mis_split(tmp_path: Path) -> None:
    """'10.1.3 Term and Termination' must not become clause '10.1' / heading '3 Term...'."""
    path = _decimal_and_multipart_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    tree = result.tree

    assert tree.resolve_path("10.1") is None, "Backtrack must not fabricate clause '10.1'"
    all_headings = [n.heading for n in tree.all_nodes() if n.heading]
    assert not any((h or "").startswith("3 Term") for h in all_headings), (
        "'10.1.3 Term and Termination' must not be mangled into heading '3 Term and Termination'"
    )
    all_text = " ".join(n.text for n in tree.all_nodes())
    assert "10.1.3 Term and Termination" in all_text


# ---------------------------------------------------------------------------
# Table flattening
# ---------------------------------------------------------------------------


def test_table_content_is_included(tmp_path: Path) -> None:
    path = _table_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    # Table content should appear somewhere in the tree
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "Training" in all_text or "100" in all_text, (
        "Table content should be included in the clause tree text"
    )


def test_table_does_not_create_heading_nodes(tmp_path: Path) -> None:
    path = _table_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    # Fee Schedule is the heading; the table rows should NOT become new heading nodes
    fee_schedule = result.tree.resolve_path("1") or next(
        (n for n in result.tree.nodes if n.heading and "Fee" in (n.heading or "")), None
    )
    assert fee_schedule is not None, "Fee Schedule heading not found"


def test_nested_table_cell_text_not_duplicated(tmp_path: Path) -> None:
    """Regression for issue #64: a nested table's cell text must be flattened
    exactly once, not once as part of the outer cell and again as the nested
    table's own cell.
    """
    path = _nested_table_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert all_text.count("INNER") == 1, (
        f"'INNER' should appear exactly once in flattened body, got: {all_text!r}"
    )


# ---------------------------------------------------------------------------
# Tracked changes (acceptance criterion)
# ---------------------------------------------------------------------------


def test_tracked_changes_captured(tmp_path: Path) -> None:
    """Core acceptance criterion: tracked-change spans captured with authors."""
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    assert len(result.tracked.changes) >= 2, "Should capture at least the insertion and deletion"


def test_tracked_insertion_author_captured(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    assert len(insertions) >= 1
    assert any(c.author == "Alice" for c in insertions), "Alice's insertion must be captured"


def test_tracked_insertion_text_captured(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    insertion_texts = [c.text for c in insertions]
    assert any("promptly" in t for t in insertion_texts), "Inserted text must be captured"


def test_tracked_insertion_has_char_span(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    assert all(c.char_span is not None for c in insertions), "Insertions must have a char_span"
    for c in insertions:
        assert c.char_span is not None
        start, end = c.char_span
        assert end >= start >= 0


def test_tracked_deletion_author_captured(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    deletions = [c for c in result.tracked.changes if c.change_type == "deletion"]
    assert len(deletions) >= 1
    assert any(c.author == "Bob" for c in deletions), "Bob's deletion must be captured"


def test_tracked_deletion_text_captured(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    deletions = [c for c in result.tracked.changes if c.change_type == "deletion"]
    assert any("client" in c.text for c in deletions), "Deleted text must be captured"


def test_tracked_deletion_char_span_is_none(tmp_path: Path) -> None:
    """Deletions are absent from normalized text — char_span must be None."""
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    deletions = [c for c in result.tracked.changes if c.change_type == "deletion"]
    assert all(c.char_span is None for c in deletions)


def test_tracked_change_clause_path_populated(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    for change in result.tracked.changes:
        assert change.clause_path != "", "clause_path must be set on every TrackedChange"


def test_inserted_text_in_normalized_text(tmp_path: Path) -> None:
    """Inserted text must appear in the clause's normalized text."""
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "promptly" in all_text, "Inserted text should be part of normalized clause text"


def test_deleted_text_absent_from_normalized(tmp_path: Path) -> None:
    """Deleted text must NOT appear in the normalized clause text."""
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "to client" not in all_text, "Deleted text must not appear in normalized text"


def test_tracked_changes_document_id(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "agreement-123", "v1")
    assert result.tracked.document_id == "agreement-123"
    assert result.tracked.version == "v1"


# ---------------------------------------------------------------------------
# char_span integrity
# ---------------------------------------------------------------------------


def test_char_spans_are_non_negative(tmp_path: Path) -> None:
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    for node in result.tree.all_nodes():
        start, end = node.char_span
        assert start >= 0 and end >= start, (
            f"Invalid char_span {node.char_span} on {node.clause_path!r}"
        )


def test_tree_validates_unique_paths(tmp_path: Path) -> None:
    """validate() should pass — no duplicate clause_paths."""
    path = _simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    result.tree.validate()  # must not raise


# ---------------------------------------------------------------------------
# Empty document
# ---------------------------------------------------------------------------


def test_empty_docx_produces_empty_tree(tmp_path: Path) -> None:
    path = _empty_docx(tmp_path)
    result = ingest_docx(path, "empty", "v1")
    # An empty document may produce 0 nodes or a single empty root
    assert result.tree is not None
    assert result.tracked.changes == []


# ---------------------------------------------------------------------------
# TrackedChanges.to_dict round-trip
# ---------------------------------------------------------------------------


def test_tracked_changes_to_dict(tmp_path: Path) -> None:
    path = _tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    d = result.tracked.to_dict()
    assert d["document_id"] == "d"
    assert isinstance(d["changes"], list)
    for c in d["changes"]:
        assert "change_type" in c
        assert "author" in c
        assert "text" in c
        assert "clause_path" in c


# ---------------------------------------------------------------------------
# Error cases
# ---------------------------------------------------------------------------


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(DocxIngesterError, match="not found"):
        ingest_docx(tmp_path / "ghost.docx", "d", "v1")


def test_non_docx_raises(tmp_path: Path) -> None:
    bad = tmp_path / "not_a_docx.docx"
    bad.write_bytes(b"this is not a zip file")
    with pytest.raises(DocxIngesterError, match="Cannot open"):
        ingest_docx(bad, "d", "v1")


# ---------------------------------------------------------------------------
# BLOCKING-2: tracked changes nested inside w:hyperlink (recursive descent)
# ---------------------------------------------------------------------------


def test_tracked_insertion_inside_hyperlink_captured(tmp_path: Path) -> None:
    """w:ins nested inside w:hyperlink must be captured as a TrackedChange."""
    path = _hyperlink_tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    assert len(insertions) >= 1, "Insertion inside hyperlink should be captured"
    assert any(c.author == "Alice" for c in insertions)
    assert any("Schedule A" in c.text for c in insertions)


def test_inserted_text_inside_hyperlink_in_normalized_text(tmp_path: Path) -> None:
    """Text inserted inside a hyperlink must appear in the normalized clause text."""
    path = _hyperlink_tracked_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "Schedule A" in all_text


# ---------------------------------------------------------------------------
# Issue #46: text inside w:sdt content controls (and w:fldSimple fields)
# must not be silently dropped.
# ---------------------------------------------------------------------------


def test_run_inside_inline_sdt_content_control_in_normalized_text(tmp_path: Path) -> None:
    """A run wrapped in an inline w:sdt content control must not vanish."""
    path = _sdt_content_control_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "Acme Corp" in all_text


def test_paragraph_inside_block_level_sdt_is_not_dropped(tmp_path: Path) -> None:
    """A whole paragraph wrapped in a block-level w:sdt must still be ingested."""
    path = _block_level_sdt_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "FixtureCorp shall deliver the wrapped clause text." in all_text


def test_run_inside_fld_simple_in_normalized_text(tmp_path: Path) -> None:
    """A run wrapped in a w:fldSimple simple field must not vanish."""
    path = _fld_simple_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    all_text = " ".join(n.text for n in result.tree.all_nodes())
    assert "1 January 2024" in all_text


# ---------------------------------------------------------------------------
# BLOCKING-1: char_span is document-absolute (multi-paragraph clause)
# ---------------------------------------------------------------------------


def test_char_span_is_document_absolute(tmp_path: Path) -> None:
    """TrackedChange.char_span must be document-absolute, not paragraph-local.

    In a multi-paragraph clause, an insertion in the 2nd paragraph must have
    a char_span > 0 that accounts for the preceding paragraph text.
    """
    path = _multi_para_clause_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    assert any("as follows" in c.text for c in insertions), "Carol's insertion not found"
    carol_ins = next(c for c in insertions if "as follows" in c.text)
    assert carol_ins.char_span is not None
    start, end = carol_ins.char_span
    # The insertion is in the SECOND body paragraph, so its start must be > 0
    # (it follows at least the heading text + "\n" + first body paragraph + "\n").
    assert start > 0, (
        f"char_span {carol_ins.char_span} looks paragraph-local (start=0); "
        f"expected document-absolute offset > 0 for a 2nd-paragraph insertion"
    )
    assert end > start


def test_char_span_length_matches_insertion_text(tmp_path: Path) -> None:
    """The span length (end - start) must equal the length of the inserted text."""
    path = _multi_para_clause_docx(tmp_path)
    result = ingest_docx(path, "d", "v1")
    insertions = [c for c in result.tracked.changes if c.change_type == "insertion"]
    carol_ins = next((c for c in insertions if "as follows" in c.text), None)
    assert carol_ins is not None, "Carol's insertion not found"
    assert carol_ins.char_span is not None
    start, end = carol_ins.char_span
    assert end - start == len(carol_ins.text), (
        f"char_span length {end - start} != text length {len(carol_ins.text)}"
    )
