"""Tests for the segmenter (post-ingestion sub-clause promotion).

SECURITY NOTE: All fixtures use programmatically constructed ClauseTree
objects with synthetic text.  No real agreement files are referenced.
Party names use fictional identifiers ("Alpha Corp", "Beta Ltd", "Party A",
"Party B") only.
"""

from __future__ import annotations

import json
from pathlib import Path

import jsonschema
import pytest
from docx import Document

from playbook_engine.clause_tree import ClauseNode, ClauseTree
from playbook_engine.docx_ingester import ingest_docx
from playbook_engine.segmenter import _split_lettered, _split_roman, segment

SCHEMA_PATH = Path(__file__).parent.parent / "spec" / "clause-tree.schema.json"


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def _leaf(path: str, text: str, *, start: int = 0, heading: str | None = None) -> ClauseNode:
    return ClauseNode(
        clause_path=path,
        heading=heading,
        text=text,
        char_span=(start, start + len(text)),
    )


def _tree_from_nodes(*nodes: ClauseNode) -> ClauseTree:
    return ClauseTree(document_id="test", version="v1", source_file="test.docx", nodes=list(nodes))


# ---------------------------------------------------------------------------
# Unit: _split_lettered
# ---------------------------------------------------------------------------


def test_split_lettered_no_items() -> None:
    result = _split_lettered("Plain body text with no sub-items.")
    assert len(result) == 1
    marker, text, offset = result[0]
    assert marker is None
    assert "Plain body text" in text


def test_split_lettered_basic() -> None:
    body = "(a) First item.\n(b) Second item.\n(c) Third item."
    result = _split_lettered(body)
    markers = [r[0] for r in result]
    assert markers == [None, "a", "b", "c"]


def test_split_lettered_preamble_preserved() -> None:
    body = "The parties agree to the following:\n(a) First item.\n(b) Second item."
    result = _split_lettered(body)
    assert result[0][0] is None
    assert "parties agree" in result[0][1]


def test_split_lettered_text_content() -> None:
    body = "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    result = _split_lettered(body)
    assert result[1][0] == "a"
    assert "Alpha Corp" in result[1][1]
    assert result[2][0] == "b"
    assert "Beta Ltd" in result[2][1]


def test_split_lettered_empty_text() -> None:
    result = _split_lettered("")
    assert len(result) == 1
    assert result[0][0] is None


def test_split_lettered_offset_increases() -> None:
    body = "(a) Item one.\n(b) Item two."
    result = _split_lettered(body)
    offsets = [r[2] for r in result if r[0] is not None]
    assert offsets[1] > offsets[0]


# ---------------------------------------------------------------------------
# Unit: _split_roman
# ---------------------------------------------------------------------------


def test_split_roman_no_items() -> None:
    result = _split_roman("Plain text with no roman items.")
    assert len(result) == 1
    assert result[0][0] is None


def test_split_roman_basic() -> None:
    body = "(i) First sub-item.\n(ii) Second sub-item.\n(iii) Third sub-item."
    result = _split_roman(body)
    markers = [r[0] for r in result]
    assert markers == [None, "i", "ii", "iii"]


def test_split_roman_unrecognised_ignored() -> None:
    # _ROMAN_NUMERALS only goes up to "viii"; neither "ix" nor "xii" is in the set.
    # Both should be absorbed as body text of the previous item.
    body = "(i) First.\n(xii) Not in the set.\n(ii) Second."
    result = _split_roman(body)
    markers = [r[0] for r in result if r[0] is not None]
    assert "xii" not in markers


# ---------------------------------------------------------------------------
# segment(): no-op on trees with no inline items
# ---------------------------------------------------------------------------


def test_segment_no_inline_items_unchanged_structure() -> None:
    node = _leaf("1", "Plain body text with no sub-items.", heading="Terms")
    tree = _tree_from_nodes(node)
    result = segment(tree)
    assert len(result.nodes) == 1
    assert result.nodes[0].clause_path == "1"
    assert result.nodes[0].text == node.text


def test_segment_preserves_document_metadata() -> None:
    tree = ClauseTree(document_id="my-doc", version="draft-2", source_file="f.docx")
    result = segment(tree)
    assert result.document_id == "my-doc"
    assert result.version == "draft-2"
    assert result.source_file == "f.docx"


def test_segment_empty_tree() -> None:
    tree = ClauseTree(document_id="d", version="v1", source_file="f")
    result = segment(tree)
    assert result.nodes == []


def test_segment_does_not_mutate_input() -> None:
    node = _leaf(
        "1",
        "(a) First item.\n(b) Second item.",
        start=0,
        heading="Terms",
    )
    tree = _tree_from_nodes(node)
    _ = segment(tree)
    assert tree.nodes[0].children == []  # original unchanged


# ---------------------------------------------------------------------------
# segment(): lettered items promoted
# ---------------------------------------------------------------------------


def test_segment_promotes_lettered_items() -> None:
    node = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text="(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay.",
        char_span=(0, 11),  # span of "1. Obligations" heading
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n1 = result.resolve_path("1")
    assert n1 is not None
    child_paths = [c.clause_path for c in n1.children]
    assert "1.a" in child_paths
    assert "1.b" in child_paths


def test_segment_promoted_child_text() -> None:
    node = ClauseNode(
        clause_path="2",
        heading="Representations",
        text="(a) Party A represents the following.\n(b) Party B agrees.",
        char_span=(0, 16),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n_a = result.resolve_path("2.a")
    assert n_a is not None
    assert "Party A" in n_a.text


def test_segment_preamble_stays_in_parent() -> None:
    node = ClauseNode(
        clause_path="3",
        heading="Covenants",
        text="The parties covenant as follows:\n(a) Covenant one.\n(b) Covenant two.",
        char_span=(0, 10),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n3 = result.resolve_path("3")
    assert n3 is not None
    assert "parties covenant" in n3.text


def test_segment_empty_preamble_ok() -> None:
    node = ClauseNode(
        clause_path="4",
        heading="Terms",
        text="(a) First.\n(b) Second.",
        char_span=(0, 6),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n4 = result.resolve_path("4")
    assert n4 is not None
    # Preamble is empty — parent text should be empty or whitespace
    assert n4.text.strip() == ""


def test_segment_all_letters_a_to_e() -> None:
    body = "(a) A.\n(b) B.\n(c) C.\n(d) D.\n(e) E."
    node = ClauseNode(clause_path="5", heading=None, text=body, char_span=(0, 1))
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n5 = result.resolve_path("5")
    assert n5 is not None
    child_paths = {c.clause_path for c in n5.children}
    assert {"5.a", "5.b", "5.c", "5.d", "5.e"} == child_paths


# ---------------------------------------------------------------------------
# segment(): roman-numeral sub-items
# ---------------------------------------------------------------------------


def test_segment_roman_items_under_lettered() -> None:
    body = "(a) Delivery terms:\n(i) On time.\n(ii) In full.\n(b) Payment terms."
    node = ClauseNode(clause_path="6", heading="Delivery", text=body, char_span=(0, 8))
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n6a = result.resolve_path("6.a")
    assert n6a is not None
    roman_paths = {c.clause_path for c in n6a.children}
    assert "6.a.i" in roman_paths
    assert "6.a.ii" in roman_paths


def test_segment_roman_not_promoted_at_top_level() -> None:
    """Roman items at top-level body text (not inside lettered item) are NOT split."""
    body = "(i) First standalone roman item.\n(ii) Second standalone roman item."
    node = ClauseNode(clause_path="7", heading=None, text=body, char_span=(0, 1))
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n7 = result.resolve_path("7")
    assert n7 is not None
    # No lettered split → no children for standalone roman items
    # (roman splitting only occurs WITHIN a lettered item)
    # The node should remain as-is or have no lettered children
    assert not any(c.clause_path == "7.i" for c in n7.children)


# ---------------------------------------------------------------------------
# segment(): char_span correctness
# ---------------------------------------------------------------------------


def test_segment_char_span_non_negative() -> None:
    body = "(a) Alpha Corp.\n(b) Beta Ltd."
    node = ClauseNode(clause_path="1", heading="H", text=body, char_span=(0, 3))
    tree = _tree_from_nodes(node)
    result = segment(tree)
    for n in result.all_nodes():
        assert n.char_span[0] >= 0
        assert n.char_span[1] >= n.char_span[0]


def test_segment_child_span_after_parent_heading() -> None:
    """Children's char_span must start at or after parent heading_end + 1."""
    heading_text = "1. Obligations"
    body = "(a) First item.\n(b) Second item."
    heading_end = len(heading_text)
    node = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text=body,
        char_span=(0, heading_end),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n1 = result.resolve_path("1")
    assert n1 is not None
    for child in n1.children:
        # Each child span must start after the parent heading
        assert child.char_span[0] >= heading_end + 1


def test_segment_child_a_span_length_matches_text() -> None:
    """The span width of child (a) must equal the length of its text."""
    body = "(a) Alpha Corp shall deliver goods.\n(b) Beta Ltd shall pay."
    node = ClauseNode(clause_path="1", heading="H", text=body, char_span=(0, 1))
    tree = _tree_from_nodes(node)
    result = segment(tree)
    n_a = result.resolve_path("1.a")
    assert n_a is not None
    span_len = n_a.char_span[1] - n_a.char_span[0]
    assert span_len == len(n_a.text)


# ---------------------------------------------------------------------------
# segment(): nested ClauseTree (existing children preserved)
# ---------------------------------------------------------------------------


def test_segment_preserves_existing_children() -> None:
    child = _leaf("1.1", "Sub-section text.", start=20, heading="Sub")
    parent = ClauseNode(
        clause_path="1",
        heading="Top",
        text="(a) Inline item.",
        char_span=(0, 9),
        children=[child],
    )
    tree = _tree_from_nodes(parent)
    result = segment(tree)
    n1 = result.resolve_path("1")
    assert n1 is not None
    child_paths = {c.clause_path for c in n1.children}
    # Existing child "1.1" plus new promoted "1.a"
    assert "1.1" in child_paths
    assert "1.a" in child_paths


def test_segment_recurses_into_children() -> None:
    child = ClauseNode(
        clause_path="1.1",
        heading="Sub",
        text="(a) Sub-item A.\n(b) Sub-item B.",
        char_span=(20, 25),
    )
    parent = ClauseNode(
        clause_path="1",
        heading="Top",
        text="Parent body.",
        char_span=(0, 9),
        children=[child],
    )
    tree = _tree_from_nodes(parent)
    result = segment(tree)
    n11 = result.resolve_path("1.1")
    assert n11 is not None
    assert result.resolve_path("1.1.a") is not None
    assert result.resolve_path("1.1.b") is not None


# ---------------------------------------------------------------------------
# segment(): idempotence
# ---------------------------------------------------------------------------


def test_segment_idempotent() -> None:
    node = ClauseNode(
        clause_path="1",
        heading="Terms",
        text="(a) First.\n(b) Second.\n(c) Third.",
        char_span=(0, 5),
    )
    tree = _tree_from_nodes(node)
    once = segment(tree)
    twice = segment(once)
    # Second pass should not create extra children
    n1_once = once.resolve_path("1")
    n1_twice = twice.resolve_path("1")
    assert n1_once is not None and n1_twice is not None
    assert len(n1_once.children) == len(n1_twice.children)


# ---------------------------------------------------------------------------
# ClauseTree contract: validate() and schema
# ---------------------------------------------------------------------------


def test_segment_validate_passes() -> None:
    node = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text="(a) Deliver.\n(b) Pay.\n(c) Notify.",
        char_span=(0, 11),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    result.validate()  # must not raise


def test_segment_validates_against_json_schema() -> None:
    schema = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
    node = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text="(a) Deliver.\n(b) Pay.",
        char_span=(0, 11),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    jsonschema.validate(instance=result.to_dict(), schema=schema)  # must not raise


def test_segment_round_trips_to_json() -> None:
    node = ClauseNode(
        clause_path="2",
        heading="Representations",
        text="(a) Party A.\n(b) Party B.",
        char_span=(0, 15),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)
    restored = ClauseTree.from_json(result.to_json())
    assert restored.resolve_path("2.a") is not None
    assert restored.resolve_path("2.b") is not None


# ---------------------------------------------------------------------------
# B1 regression: resolve_span round-trip (span resolves to child.text)
# ---------------------------------------------------------------------------


def test_child_span_resolves_to_child_text() -> None:
    """B1 regression: promoted child char_span must resolve to its text content.

    Before the fix, the span pointed at the line start (including the marker
    '(a) '), so ClauseTree.resolve_span() returned '(a) First i', not the text.
    """

    heading_text = "1. Obligations"
    body_text = "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    heading_end = len(heading_text)
    # Virtual normalized text: heading + "\n" + body
    full_text = heading_text + "\n" + body_text

    node = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text=body_text,
        char_span=(0, heading_end),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)

    n_a = result.resolve_path("1.a")
    assert n_a is not None
    resolved = ClauseTree.resolve_span(full_text, n_a.char_span)
    assert resolved == n_a.text, (
        f"resolve_span returned {resolved!r} but child.text is {n_a.text!r}"
    )


def test_roman_child_span_resolves_to_roman_text() -> None:
    """B1+B2 regression: roman sub-item span resolves to its text."""
    heading_text = "6. Delivery"
    # body starts after heading + "\n" (heading_end + 1 in virtual text)
    body_text = "(a) Delivery terms:\n(i) On time.\n(ii) In full.\n(b) Payment terms."
    heading_end = len(heading_text)
    full_text = heading_text + "\n" + body_text

    node = ClauseNode(
        clause_path="6",
        heading="Delivery",
        text=body_text,
        char_span=(0, heading_end),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)

    n_i = result.resolve_path("6.a.i")
    assert n_i is not None
    resolved = ClauseTree.resolve_span(full_text, n_i.char_span)
    assert resolved == n_i.text, (
        f"resolve_span returned {resolved!r} but child.text is {n_i.text!r}"
    )


# ---------------------------------------------------------------------------
# B3 regression: no duplicate clause_path when ingester child collides
# ---------------------------------------------------------------------------


def test_no_duplicate_when_ingester_child_collides() -> None:
    """B3 regression: if an ingester child already has path '1.a', the segmenter
    must not add another '1.a' node from the body text."""
    existing_child = ClauseNode(
        clause_path="1.a",
        heading=None,
        text="Ingester-produced sub-clause.",
        char_span=(20, 47),
    )
    parent = ClauseNode(
        clause_path="1",
        heading="Obligations",
        text="(a) Duplicate item from body.",
        char_span=(0, 14),
        children=[existing_child],
    )
    tree = _tree_from_nodes(parent)
    result = segment(tree)
    result.validate()  # must not raise (no duplicate clause_paths)
    n1 = result.resolve_path("1")
    assert n1 is not None
    paths_1a = [c for c in n1.children if c.clause_path == "1.a"]
    assert len(paths_1a) == 1, "Expected exactly one '1.a' child"


# ---------------------------------------------------------------------------
# Issue #39 regression: DOCX blank paragraph must not drift promoted spans
# ---------------------------------------------------------------------------


def test_docx_blank_paragraph_after_heading_does_not_drift_child_spans(
    tmp_path: Path,
) -> None:
    """A blank paragraph before a DOCX heading must not shift that heading's
    promoted (a)/(b) child char_spans.

    Before the fix, docx_ingester advanced doc_char_offset for every blank
    paragraph even though blank paragraphs are excluded from node.text, so a
    blank paragraph anywhere before a heading shifted that heading's own
    char_span right by one character per blank paragraph — and every
    promoted child computed from ``heading.char_span[1] + 1`` drifted with
    it. (A blank paragraph strictly between a heading and its own body text
    does not surface the bug, because add_body() ignores its doc_offset
    argument once a heading is already open — hence the second heading here,
    which receives a fresh char_span computed from the drifted offset.)
    """
    doc = Document()
    doc.add_heading("Term", level=1)
    doc.add_paragraph("")  # blank paragraph before the second heading
    doc.add_heading("Obligations", level=1)
    doc.add_paragraph("(a) Alpha Corp shall deliver.")
    doc.add_paragraph("(b) Beta Ltd shall pay.")
    path = tmp_path / "blank_para.docx"
    doc.save(str(path))

    ingested = ingest_docx(path, "d", "v1")
    segmented = segment(ingested.tree)

    obligations = next(n for n in segmented.nodes if n.heading == "Obligations")
    promoted = [c for c in obligations.children if c.clause_path.split(".")[-1] in ("a", "b")]
    assert len(promoted) == 2, "Expected both (a) and (b) to be promoted"

    # Virtual normalized text per the docx_ingester's char_span contract:
    # kept (non-blank, stripped) units joined by "\n". The blank paragraph
    # contributes nothing.
    virtual_text = (
        "Term"
        + "\n"
        + "Obligations"
        + "\n"
        + "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    )

    for child in promoted:
        resolved = ClauseTree.resolve_span(virtual_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )


def test_synthetic_zero_node_body_start_uses_char_span_start() -> None:
    """The synthetic clause_path='0' node (pre-heading body text) has no
    heading line of its own — its char_span covers the body text directly.

    Before the fix, body_start was computed as char_span[1] + 1 for every
    node (including '0'), overshooting by len(first_paragraph) + 1.

    Per the structural discriminator (segmenter.py), the synthetic node's
    char_span covers exactly its own first line of text — mirroring what
    ``_ClauseBuilder.add_body`` actually records — not the full multi-line
    accumulated text, so the fixture's char_span is built from the first
    line only.
    """
    body_text = "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    first_line = body_text.split("\n", 1)[0]
    node = ClauseNode(
        clause_path="0",
        heading=None,
        text=body_text,
        char_span=(0, len(first_line)),
    )
    tree = _tree_from_nodes(node)
    result = segment(tree)

    zero = result.resolve_path("0")
    assert zero is not None
    n_a = result.resolve_path("0.a")
    n_b = result.resolve_path("0.b")
    assert n_a is not None
    assert n_b is not None

    # The virtual text for the "0" node IS its own body_text — there is no
    # preceding heading line.
    for child in (n_a, n_b):
        resolved = ClauseTree.resolve_span(body_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )


def test_docx_tab_indented_body_line_does_not_shift_child_span(tmp_path: Path) -> None:
    """A tab/space-indented DOCX body paragraph must not desync content_offset.

    Before the fix, add_body() received the raw (unstripped) paragraph text,
    so a leading-whitespace-indented "(a) ..." line broke the segmenter's
    'stripped == line' assumption and made content_offset short by the
    leading-whitespace length.
    """
    doc = Document()
    doc.add_heading("Obligations", level=1)
    doc.add_paragraph("   (a) Alpha Corp shall deliver.")
    doc.add_paragraph("(b) Beta Ltd shall pay.")
    path = tmp_path / "indented.docx"
    doc.save(str(path))

    ingested = ingest_docx(path, "d", "v1")
    segmented = segment(ingested.tree)

    obligations = next(n for n in segmented.nodes if n.heading == "Obligations")
    assert not obligations.text.startswith(" "), (
        "Ingester body text must be stripped before accumulation"
    )

    virtual_text = "Obligations" + "\n" + "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    promoted = [c for c in obligations.children if c.clause_path.split(".")[-1] in ("a", "b")]
    assert len(promoted) == 2
    for child in promoted:
        resolved = ClauseTree.resolve_span(virtual_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )


def test_docx_genuine_zero_numbered_heading_does_not_shift_child_spans(
    tmp_path: Path,
) -> None:
    """A REAL '0.'-numbered heading (e.g. '0. Preamble') must be segmented
    using the heading-relative body_start (char_span[1] + 1), not the
    synthetic-pre-heading-node's body_start (char_span[0]).

    Both the genuine "0." clause and the synthetic pre-heading node share
    clause_path == "0" and (when the heading text is empty) heading is None
    too, so clause_path and heading cannot discriminate the synthetic node
    even together — see test_docx_bare_zero_numbered_heading_does_not_shift_child_spans
    for the empty-heading-text case that clause_path/heading alone would
    misclassify. The real discriminator is structural (see segmenter.py).
    Before the fix, discriminating on clause_path == "0" alone caused every
    promoted child under a genuine "0." heading to have its span shifted
    left by the whole heading line's length + 1.
    """
    doc = Document()
    doc.add_paragraph("0. Preamble")
    doc.add_paragraph("(a) Alpha Corp shall deliver.")
    doc.add_paragraph("(b) Beta Ltd shall pay.")
    path = tmp_path / "zero_heading.docx"
    doc.save(str(path))

    ingested = ingest_docx(path, "d", "v1")
    segmented = segment(ingested.tree)

    root = segmented.nodes[0]
    assert root.clause_path == "0"
    assert root.heading == "Preamble"

    virtual_text = "0. Preamble" + "\n" + "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    promoted = [c for c in root.children if c.clause_path.split(".")[-1] in ("a", "b")]
    assert len(promoted) == 2, "Expected both (a) and (b) to be promoted"
    for child in promoted:
        resolved = ClauseTree.resolve_span(virtual_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )


@pytest.mark.parametrize("heading_line", ["0.", "0)"])
def test_docx_bare_zero_numbered_heading_does_not_shift_child_spans(
    tmp_path: Path,
    heading_line: str,
) -> None:
    """A REAL '0.'-numbered heading whose heading text is EMPTY (the
    paragraph is exactly "0." or "0)", so heading is None) must still be
    segmented using the heading-relative body_start (char_span[1] + 1), not
    the synthetic-pre-heading-node's body_start (char_span[0]).

    This is the case that clause_path == "0" and heading is None cannot
    discriminate even together, since the genuine node and the synthetic
    pre-heading node both satisfy both conditions. Before the segmenter fix,
    this case misfired and shifted every promoted child's span left by the
    whole heading line's length + 1 (issue #39, finding 1, fix round 2).
    """
    doc = Document()
    doc.add_paragraph(heading_line)
    doc.add_paragraph("(a) Alpha Corp shall deliver.")
    doc.add_paragraph("(b) Beta Ltd shall pay.")
    path = tmp_path / "bare_zero_heading.docx"
    doc.save(str(path))

    ingested = ingest_docx(path, "d", "v1")
    segmented = segment(ingested.tree)

    root = segmented.nodes[0]
    assert root.clause_path == "0"
    assert root.heading is None

    virtual_text = heading_line + "\n" + "(a) Alpha Corp shall deliver.\n(b) Beta Ltd shall pay."
    promoted = [c for c in root.children if c.clause_path.split(".")[-1] in ("a", "b")]
    assert len(promoted) == 2, "Expected both (a) and (b) to be promoted"
    for child in promoted:
        resolved = ClauseTree.resolve_span(virtual_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )


def test_docx_non_zero_heading_length_coincidence_does_not_shift_child_spans(
    tmp_path: Path,
) -> None:
    """A genuine NON-"0" bare-numbered heading whose heading line happens to
    be the same length as its first body line must still be segmented using
    the heading-relative body_start (char_span[1] + 1), not the
    synthetic-pre-heading-node's body_start (char_span[0]).

    This is the false-positive class the length-only discriminator (with the
    ``clause_path == "0"`` conjunct dropped) opens up: heading "10.1." is 5
    characters and first body line "(a) A" is also 5 characters, so
    ``char_span[1] - char_span[0] == len(text.split("\\n", 1)[0])`` holds even
    though this is an ordinary "10.1" clause, not the synthetic pre-heading
    node. ``clause_path == "0"`` rules this node out (its clause_path is
    "10.1"), which is exactly why that conjunct must be kept rather than
    replaced by the length check alone. Without the conjunct, body_start
    would be miscomputed as char_span[0]=0 instead of char_span[1]+1=6,
    shifting every promoted child's span left by 6 characters (reproduced
    end-to-end: resolve_span('10.1.a') would return '.' instead of 'A').
    """
    doc = Document()
    doc.add_paragraph("10.1.")
    doc.add_paragraph("(a) A")
    doc.add_paragraph("(b) B")
    path = tmp_path / "length_coincidence.docx"
    doc.save(str(path))

    ingested = ingest_docx(path, "d", "v1")

    pre_segment_root = ingested.tree.nodes[0]
    assert pre_segment_root.clause_path == "10.1"
    assert pre_segment_root.heading is None
    # Sanity-check the length coincidence this test depends on, against the
    # pre-segmentation node (post-segmentation text is just the preamble,
    # which is empty here since every body line is consumed as a lettered
    # item).
    assert pre_segment_root.char_span[1] - pre_segment_root.char_span[0] == len(
        pre_segment_root.text.split("\n", 1)[0]
    )

    segmented = segment(ingested.tree)

    root = segmented.nodes[0]
    assert root.clause_path == "10.1"
    assert root.heading is None

    virtual_text = "10.1." + "\n" + "(a) A\n(b) B"
    promoted = [c for c in root.children if c.clause_path.split(".")[-1] in ("a", "b")]
    assert len(promoted) == 2, "Expected both (a) and (b) to be promoted"
    for child in promoted:
        resolved = ClauseTree.resolve_span(virtual_text, child.char_span)
        assert resolved == child.text, (
            f"{child.clause_path}: resolve_span returned {resolved!r} but "
            f"child.text is {child.text!r} (span {child.char_span})"
        )
