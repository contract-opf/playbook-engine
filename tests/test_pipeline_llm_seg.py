"""Pipeline integration tests for LLM segmentation wiring (issues #74, #76).

Verifies the three acceptance criteria for the synchronous LLM-segmentation
path (#74):

1. End-to-end on a synthetic corpus with an injected fake ``segment_fn``:
   ``mine_corpus(..., use_llm_segmentation=True)`` produces
   ``observations.jsonl`` whose clauses reflect the LLM tree and carry the
   LLM ``taxonomy_id`` (classified, not all-None).
2. The produced ``ClauseTree`` satisfies the existing downstream contract:
   diff/deviation run without error, and ``playbook validate`` on the
   resulting playbook passes (``ValidationResult.ok`` is True — the
   equivalent of ``playbook validate`` exiting 0).
3. With a fake ``segment_fn`` returning gate-failing output, the run raises
   ``SegmentationQAError`` (fail loud) rather than silently degrading.

Plus the batch-segmentation pre-pass wiring (#76, near the bottom of this
file): ``mine_corpus(..., use_batch_segmentation=True)`` extracts every
version up front and segments the whole corpus via one mocked Message
Batches client, feeding the batched ``SegNode`` output into the same
grounding + QA + observation flow — same classified-observations and
fail-loud-QA contracts as the synchronous path, exercised through the real
``segment_documents_batch`` wiring (not a fake segment function).

SECURITY NOTE: All fixtures use programmatically constructed RTF text with
synthetic, fictional content.  No real agreement files are referenced, and
the fake ``segment_fn``/batches client never calls a live LLM — no network,
no API key. Fictional party names only (e.g. "Alpha Corp", "Beta University").
"""

from __future__ import annotations

import json
import subprocess
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
import yaml

from playbook_engine import extraction
from playbook_engine import pipeline as pipeline_module
from playbook_engine import publisher as publisher_module
from playbook_engine.aar import (
    build_after_action_data,
    build_after_action_report,
    write_after_action_report,
)
from playbook_engine.clause_classifier import AMBIGUITY_THRESHOLD
from playbook_engine.clause_differ import ClauseDiff, TextHunk
from playbook_engine.clause_tree import ClauseNode, ClauseTree
from playbook_engine.config import load_config
from playbook_engine.docx_ingester import TextUnit, TrackedChange, TrackedChanges
from playbook_engine.entity_registry import EntityRegistry
from playbook_engine.extraction import ExtractionCache, ExtractorLabel, extract_blocks
from playbook_engine.llm_segmenter_batch import (
    NormalizeTrailError,
    NormalizeTrailResult,
    SegmentationVerdictCache,
    segment_documents_batch,
)
from playbook_engine.observation_builder import read_observations_jsonl
from playbook_engine.pipeline import (
    PipelineError,
    _attribution_for_diff,
    _bridge_tracked_changes_if_needed,
    _classified_from_taxonomy_by_path,
    _extract_blocks_for_bridge,
    _llm_tracked_changes,
    mine_corpus,
    project_playbook,
)
from playbook_engine.segmentation_grounding import Block, SegNode
from playbook_engine.taxonomy import load_taxonomy
from playbook_engine.validator import validate_document

# ---------------------------------------------------------------------------
# RTF fixture helpers (same convention as test_pipeline_project.py /
# test_pipeline_provenance.py — extract_blocks supports .rtf via pandoc, so
# these fixtures exercise the real extractor, not a synthetic Block list).
# ---------------------------------------------------------------------------

_RTF_PROLOGUE = (
    r"{\rtf1\ansi\deff0"
    r"{\fonttbl{\f0\froman\fcharset0 Times New Roman;}}"
    r"\f0\fs24 "
)
_RTF_EPILOGUE = r"}"


def _rtf(body: str) -> str:
    return _RTF_PROLOGUE + body + _RTF_EPILOGUE


def _write_rtf(path: Path, body: str) -> None:
    path.write_text(_rtf(body), encoding="utf-8")


_TAXONOMY_PATH = Path(__file__).parent.parent / "spec" / "taxonomy" / "affiliation-agreement.yaml"

# Two clauses, each "N. Heading\par body text\par" — extract_blocks (RTF via
# pandoc) turns each \par-delimited paragraph into its own Block, in order:
# b0=heading1, b1=body1, b2=heading2, b3=body2. See _fake_segment_fn below,
# which relies on exactly this block-per-paragraph shape.
_V1_BODY = (
    r"1. Indemnification\par "
    r"Alpha Corp shall indemnify Beta University against third-party claims "
    r"arising from the placement programme.\par "
    r"2. Governing Law\par "
    r"This agreement is governed by the laws of the State of California.\par "
)

# v2: signed copy — same clauses, slightly different Indemnification body
# (so the diff/deviation stage has a genuine changed clause to assess), plus
# a signatures block so detect_signed() anchors it as signed.
_V2_BODY = (
    r"1. Indemnification\par "
    r"Alpha Corp shall indemnify and hold harmless Beta University against "
    r"third-party claims arising from the placement programme.\par "
    r"2. Governing Law\par "
    r"This agreement is governed by the laws of the State of California.\par "
    r"3. Signatures\par "
    r"By: Jane Doe, Alpha Corp\par "
)

# Heading text -> taxonomy_id, used by the fake segment_fn to classify each
# clause in the same pass as segmentation (mirrors what the real LLM
# segmenter does per llm_segmenter.SEGMENTER_SYSTEM_PROMPT).
_HEADING_TAXONOMY = {
    "1. Indemnification": "indemnification",
    "2. Governing Law": "governing_law",
}

# A second, unrelated document used by the mixed-corpus tests below (issue
# #83) — deliberately NOT in _HEADING_TAXONOMY (its heading is "1.
# Confidentiality", never "1. Indemnification"/"2. Governing Law") so
# _dispatching_segment_fn routes it to the always-gate-failing fake by
# content alone (segment_fn's contract has no doc_id parameter — see
# segmentation_qa.segment_verify_repair's docstring). A distinct fictional
# counterparty name (Gamma Industries) from _V1_BODY/_V2_BODY's Alpha
# Corp/Beta University keeps the two documents' provenance unambiguous.
_BAD_DOC_BODY = (
    r"1. Confidentiality\par "
    r"Gamma Industries shall keep all disclosed information confidential.\par "
)

# A QA-quarantined document whose body names a configured
# provenance.known_entities entry as the LITERAL FIRST TOKEN of the
# paragraph that _gate_failing_segment_fn leaves uncovered (issue #83
# regression test). _gate_failing_segment_fn covers only the FIRST block
# (the heading paragraph), so the coverage gate's "trailing" uncovered text
# begins with a real newline immediately followed by this paragraph's first
# word — extract_blocks (RTF via pandoc) separates \par paragraphs with a
# single "\n", not a space, so there is no whitespace of any kind between
# that newline and the entity name. This is deliberately NOT prefixed with
# any lead-in text (e.g. "Held by ..."): an earlier attempt at this fix used
# such a prefix specifically to keep the entity name from being the first
# token, which sidesteps rather than tests the dangerous case — a raw
# newline followed immediately by a known-entity name is exactly the input
# that broke a whole-word alias scrub built on Python's repr() (a leading
# "\n" renders as the two literal characters backslash+"n", and that
# trailing "n" reads as a word character immediately before the name,
# defeating a (?<!\w) boundary check). This fix instead never embeds the raw
# gap text in the QA error message at all (segmentation_qa._check_coverage),
# so this input must produce no leak regardless of that matching edge case.
_ENTITY_LEAK_NAME = "Fictional University"
_ENTITY_LEAK_BODY = (
    r"1. Confidentiality\par "
    rf"{_ENTITY_LEAK_NAME} shall retain all right, title, and interest in the "
    r"disclosed materials, and no license or other right is granted except as "
    r"expressly set forth in this section.\par "
)


# ---------------------------------------------------------------------------
# Fake segment_fn — deterministic, block-pair based (no LLM, no network)
# ---------------------------------------------------------------------------


def _fake_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Pair consecutive (heading, body) blocks into clause nodes.

    Matches the RTF fixtures' block-per-paragraph shape from extract_blocks:
    each clause is exactly two blocks (a numbered heading, then its body).
    A block whose text isn't in ``_HEADING_TAXONOMY`` (e.g. "3. Signatures")
    is still covered by a node (every block must be accounted for), but with
    ``taxonomy_id=None`` — the LLM's own convention for non-clause noise.
    """
    del canonical_text
    nodes: list[SegNode] = []
    order = 1
    i = 0
    while i < len(blocks):
        heading_block = blocks[i]
        tid = _HEADING_TAXONOMY.get(heading_block.text)
        if tid is None:
            # Non-clause block (e.g. a signatures heading + its one body
            # block) — still emit a node covering both so coverage holds.
            end_block = blocks[i + 1] if i + 1 < len(blocks) else heading_block
            nodes.append(
                SegNode(
                    node_id=f"n{order}",
                    parent_id=None,
                    order=order,
                    heading=heading_block.text,
                    taxonomy_id=None,
                    start_block_id=heading_block.block_id,
                    end_block_id=end_block.block_id,
                )
            )
            i += 2 if end_block is not heading_block else 1
        else:
            body_block = blocks[i + 1]
            nodes.append(
                SegNode(
                    node_id=f"n{order}",
                    parent_id=None,
                    order=order,
                    heading=heading_block.text,
                    taxonomy_id=tid,
                    start_block_id=heading_block.block_id,
                    end_block_id=body_block.block_id,
                    start_quote=heading_block.text[:10],
                    end_quote=body_block.text[-10:],
                )
            )
            i += 2
        order += 1
    return nodes


def _gate_failing_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Return a segmentation that fails the coverage gate on every attempt.

    Only covers the first block, leaving the rest of ``canonical_text``
    uncovered — a deterministic, always-failing candidate so
    ``segment_verify_repair`` exhausts every repair attempt and
    ``SegmentationQAError`` propagates (acceptance criterion 3).
    """
    del canonical_text
    first = blocks[0]
    return [
        SegNode(
            node_id="n1",
            parent_id=None,
            order=1,
            heading=first.text,
            taxonomy_id=None,
            start_block_id=first.block_id,
            end_block_id=first.block_id,
        )
    ]


#: A (synthetic) known-entity name glued directly to a trailing "_2023" with
#: NO space/separator before the underscore (issue #97 regression test) —
#: the same word-boundary-defeating shape issue #96 proved defeats
#: entity_registry._fuzzy_name_pattern's (?<!\w)NAME(?!\w) whole-word
#: boundary check (underscore is a \w character). Used as an out-of-enum
#: taxonomy_id below so the taxonomy gate rejects it — proving the fix
#: (segmentation_qa._safe_taxonomy_id_repr) closes the leak at the source
#: rather than relying on that scrub, which cannot catch a glued token.
_TAXONOMY_LEAK_ENTITY_NAME = "Fictional University"
_TAXONOMY_LEAK_REJECTED_ID = f"{_TAXONOMY_LEAK_ENTITY_NAME}_2023"


def _taxonomy_leak_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Segment as ONE node spanning every block — passes grounding,
    coverage, reconstruction, and tree trivially (a single leaf whose span
    is the entire canonical text has nothing to conflict with) — but
    assigns an out-of-enum ``taxonomy_id`` (``_TAXONOMY_LEAK_REJECTED_ID``),
    so the run fails ONLY the taxonomy gate (issue #97 regression test),
    deterministically on every repair attempt.
    """
    del canonical_text
    return [
        SegNode(
            node_id="n1",
            parent_id=None,
            order=1,
            heading=blocks[0].text,
            taxonomy_id=_TAXONOMY_LEAK_REJECTED_ID,
            start_block_id=blocks[0].block_id,
            end_block_id=blocks[-1].block_id,
        )
    ]


#: A non-``str`` (JSON-number-shaped) rejected ``taxonomy_id`` (issue #97
#: fix round 1 regression test) — exactly the shape a model response that
#: violates its own structured-output schema for this field can produce at
#: runtime (``llm_segmenter._parse_seg_nodes`` assigns
#: ``raw["taxonomy_id"]`` with no type check into a plain, unvalidated
#: ``SegNode`` dataclass). Used as an out-of-enum taxonomy_id so the
#: taxonomy gate rejects it, proving
#: ``segmentation_qa._safe_taxonomy_id_repr``'s ``isinstance`` guard keeps
#: this a fail-loud ``SegmentationQAError`` (document quarantined) rather
#: than an uncaught ``TypeError`` that ``pipeline._compute_doc_result``'s
#: broad ``except Exception`` swallows into a per-file warning.
_TAXONOMY_NON_STRING_REJECTED_ID = 1234


def _taxonomy_non_string_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Segment as ONE node spanning every block — passes grounding,
    coverage, reconstruction, and tree trivially (a single leaf whose span
    is the entire canonical text has nothing to conflict with) — but
    assigns a non-``str`` ``taxonomy_id`` (``_TAXONOMY_NON_STRING_REJECTED_ID``),
    so the run fails ONLY the taxonomy gate (issue #97 fix round 1
    regression test), deterministically on every repair attempt.
    """
    del canonical_text
    return [
        SegNode(
            node_id="n1",
            parent_id=None,
            order=1,
            heading=blocks[0].text,
            taxonomy_id=_TAXONOMY_NON_STRING_REJECTED_ID,  # type: ignore[arg-type]
            start_block_id=blocks[0].block_id,
            end_block_id=blocks[-1].block_id,
        )
    ]


#: An UNHASHABLE (JSON-array-shaped) rejected ``taxonomy_id`` (issue #97 fix
#: round 2 regression test) — a highly plausible schema violation (a model
#: assigning two taxonomy categories to one node), and exactly the shape
#: that defeated the round-1 fix: ``_check_taxonomy`` evaluated
#: ``taxonomy_id not in allowed`` against a ``set`` BEFORE ever calling
#: ``segmentation_qa._safe_taxonomy_id_repr``, so a ``list`` (unhashable)
#: raised an uncaught ``TypeError: unhashable type`` at that membership
#: test — never even reaching the round-1 fix's ``isinstance`` guard. Used
#: as an out-of-enum taxonomy_id so the taxonomy gate rejects it, proving
#: the round-2 fix (checking ``isinstance`` in ``_check_taxonomy`` itself,
#: before the ``set`` lookup) keeps this a fail-loud ``SegmentationQAError``
#: (document quarantined) rather than an uncaught ``TypeError`` that
#: ``pipeline._compute_doc_result``'s broad ``except Exception`` swallows
#: into a per-file warning.
_TAXONOMY_LIST_REJECTED_ID = ["indemnification", "governing_law"]


def _taxonomy_list_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Segment as ONE node spanning every block — passes grounding,
    coverage, reconstruction, and tree trivially (a single leaf whose span
    is the entire canonical text has nothing to conflict with) — but
    assigns an unhashable ``taxonomy_id`` (``_TAXONOMY_LIST_REJECTED_ID``),
    so the run fails ONLY the taxonomy gate (issue #97 fix round 2
    regression test), deterministically on every repair attempt.
    """
    del canonical_text
    return [
        SegNode(
            node_id="n1",
            parent_id=None,
            order=1,
            heading=blocks[0].text,
            taxonomy_id=_TAXONOMY_LIST_REJECTED_ID,  # type: ignore[arg-type]
            start_block_id=blocks[0].block_id,
            end_block_id=blocks[-1].block_id,
        )
    ]


def _fail_only_second_version_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Succeed via ``_fake_segment_fn`` for _V1_BODY's content, fail every QA
    gate for _V2_BODY's content (distinguished by its "Signatures" clause,
    absent from _V1_BODY) — lets a two-version single-document corpus
    quarantine only its SECOND version, so by the time the
    SegmentationQAError fires, ``_compute_doc_result``'s per-version loop has
    already recorded a genuine "ok" row for the first version (issue #83:
    exercises more than the degenerate zero-versions-mined case
    ``test_gate_failing_segment_fn_quarantines_document`` covers).
    """
    if "Signatures" in canonical_text:
        return _gate_failing_segment_fn(canonical_text, blocks)
    return _fake_segment_fn(canonical_text, blocks)


def _dispatching_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Route to ``_fake_segment_fn`` for _V1_BODY-shaped content (identified
    by its "Indemnification" heading), to ``_gate_failing_segment_fn`` for
    anything else (issue #83's mixed-corpus tests below) — one ``mine_corpus``
    call with a SINGLE ``llm_segment_fn`` exercising a corpus where one
    document mines cleanly and another is QA-quarantined. Dispatches on
    content because ``segment_fn``'s contract has no doc_id/version
    parameter to key on directly.
    """
    if "Indemnification" in canonical_text:
        return _fake_segment_fn(canonical_text, blocks)
    return _gate_failing_segment_fn(canonical_text, blocks)


# ---------------------------------------------------------------------------
# Corpus + config factory
# ---------------------------------------------------------------------------


def _make_corpus(tmp_path: Path, *, two_versions: bool) -> tuple[Path, Path, Path]:
    """Build a synthetic corpus + config; return (corpus_dir, config_path, out_dir)."""
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-001"
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "v1.rtf", _V1_BODY)
    if two_versions:
        _write_rtf(deal_dir / "v2.rtf", _V2_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    out_dir = tmp_path / "out"
    return corpus_dir, config_path, out_dir


def _make_single_doc_corpus_with_known_entities(
    tmp_path: Path, doc_id: str, body: str, known_entities: list[str]
) -> tuple[Path, Path, Path]:
    """Single-document/single-version corpus, parameterized on the document
    folder name, RTF *body*, and ``provenance.known_entities`` (issue #83) —
    for tests that need control over exactly what a QA-gate failure's error
    message embeds, and/or whether the document folder itself is named after
    a known entity.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / doc_id
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "v1.rtf", body)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"], "known_entities": known_entities},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    out_dir = tmp_path / "out"
    return corpus_dir, config_path, out_dir


def _make_mixed_corpus_with_known_entities(
    tmp_path: Path, quarantined_doc_id: str, known_entities: list[str]
) -> tuple[Path, Path, Path]:
    """Two documents: "deal-001" mines cleanly (_V1_BODY), and a second,
    named *quarantined_doc_id*, whose only version fails every QA gate
    (issue #83) — for exercising ``mine_corpus``/``build_after_action_data``
    on a corpus that quarantines only PART of itself, with
    ``provenance.known_entities`` configured. Pair with
    ``_dispatching_segment_fn``.
    """
    corpus_dir = tmp_path / "corpus"
    good_dir = corpus_dir / "deal-001"
    good_dir.mkdir(parents=True)
    _write_rtf(good_dir / "v1.rtf", _V1_BODY)

    bad_dir = corpus_dir / quarantined_doc_id
    bad_dir.mkdir(parents=True)
    _write_rtf(bad_dir / "v1.rtf", _BAD_DOC_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"], "known_entities": known_entities},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    out_dir = tmp_path / "out"
    return corpus_dir, config_path, out_dir


#: Default malformed hints.yaml (unbalanced flow-sequence bracket) — content
#: has nothing to do with any entity name, so this isolates the
#: path-embedding leak (issue #96) from any unrelated leak surface in the
#: YAML parser's own message (see callers below that override
#: hints_yaml_text to exercise those other surfaces deliberately).
_GENERIC_MALFORMED_HINTS_YAML = "order: [v1, v2\nsigned_version: v1\n"


def _make_mixed_corpus_with_malformed_hints(
    tmp_path: Path,
    quarantined_doc_id: str,
    known_entities: list[str],
    hints_yaml_text: str = _GENERIC_MALFORMED_HINTS_YAML,
) -> tuple[Path, Path, Path]:
    """Two documents: "deal-001" mines cleanly (_V1_BODY, deterministic
    path — no LLM segmentation needed since ``Hints.load`` runs on both
    paths), and a second, named *quarantined_doc_id*, whose ``hints.yaml``
    is malformed (issue #96) — ``HintsError`` propagates and quarantines
    only this document. Shaped like ``_make_mixed_corpus_with_known_entities``
    (a real second document with real content is required: unlike
    ``SegmentationQAError``, ``HintsError`` carries no ``partial_corpus_doc``
    — see pipeline.py's quarantine handler — so a single-document corpus
    would leave BOTH observations.jsonl and corpus_manifest.json empty and
    ``project_playbook`` would refuse to compile at all
    (``PipelineError: Observation store is empty``), before this fix is even
    exercised).
    """
    corpus_dir = tmp_path / "corpus"
    good_dir = corpus_dir / "deal-001"
    good_dir.mkdir(parents=True)
    _write_rtf(good_dir / "v1.rtf", _V1_BODY)

    bad_dir = corpus_dir / quarantined_doc_id
    bad_dir.mkdir(parents=True)
    _write_rtf(bad_dir / "v1.rtf", _V1_BODY)
    (bad_dir / "hints.yaml").write_text(hints_yaml_text, encoding="utf-8")

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"], "known_entities": known_entities},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    out_dir = tmp_path / "out"
    return corpus_dir, config_path, out_dir


# ---------------------------------------------------------------------------
# docling mocking helpers (issue #81) — mirror test_extraction.py's
# _mock_docling_subprocess, but split into an always-failing variant (forces
# extract_blocks's live per-file backend-error fallback) and a
# markdown-controlled succeeding variant, since the fallback tests below
# need BOTH within the same file (and sometimes within the same test).
# ---------------------------------------------------------------------------


#: The real shutil.which/subprocess.run, captured once at import time — the
#: fakes below must only special-case docling invocations and delegate
#: everything else (notably pandoc's OWN `shutil.which("pandoc")` check AND
#: its own `subprocess.run(["pandoc", ...])` call — both go through these
#: same two extraction.py module attributes) to the real implementation, or
#: a docling-fails-and-recovers scenario would ALSO break the RTF legacy
#: fallback and mask the recovery behind an unrelated ExtractionError.
_real_which = extraction.shutil.which
_real_subprocess_run = extraction.subprocess.run


def _mock_docling_present_and_failing(monkeypatch: pytest.MonkeyPatch) -> None:
    """docling is on PATH but raises on every conversion — forces
    extract_blocks's live per-file backend-error fallback."""

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else _real_which(cmd)

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        if cmd[0] == "docling":
            raise subprocess.CalledProcessError(1, cmd, stderr="docling: conversion failed")
        return _real_subprocess_run(cmd, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)


def _mock_docling_present_and_succeeding(
    monkeypatch: pytest.MonkeyPatch, markdown: str, *, stem: str
) -> None:
    """docling is on PATH and succeeds, writing *markdown* as its output for
    the file whose stem is *stem*."""

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else _real_which(cmd)

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        if cmd[0] != "docling":
            return _real_subprocess_run(cmd, **kwargs)  # type: ignore[arg-type]
        outdir = Path(cmd[cmd.index("--output") + 1])
        (outdir / f"{stem}.md").write_text(markdown, encoding="utf-8")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)


# Markdown docling would produce for _V1_BODY's content — matches
# _fake_segment_fn's expected (heading, body) block-pair shape exactly (see
# _HEADING_TAXONOMY above): "1. Indemnification"/"2. Governing Law" headings,
# each followed by its body paragraph.
_V1_DOCLING_MARKDOWN = (
    "# 1. Indemnification\n"
    "\n"
    "Alpha Corp shall indemnify Beta University against third-party claims "
    "arising from the placement programme.\n"
    "\n"
    "# 2. Governing Law\n"
    "\n"
    "This agreement is governed by the laws of the State of California.\n"
)


# ---------------------------------------------------------------------------
# AC-1: end-to-end — observations carry the LLM taxonomy_id (classified)
# ---------------------------------------------------------------------------


def test_llm_segmentation_produces_classified_observations(tmp_path: Path) -> None:
    """mine_corpus(use_llm_segmentation=True) with a fake segment_fn produces
    observations whose taxonomy_id reflects the LLM tree — not all None.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    obs_path = out_dir / "observations.jsonl"
    assert obs_path.exists()
    raw_obs = read_observations_jsonl(obs_path)
    assert raw_obs, "mine_corpus must write at least one observation"

    taxonomy_ids = [o["taxonomy_id"] for o in raw_obs]
    assert not all(tid is None for tid in taxonomy_ids), (
        "LLM-segmented observations must carry the LLM's taxonomy_id (classified), not be all-None"
    )
    assert set(taxonomy_ids) == {"indemnification", "governing_law"}

    # basis="judge" — carried straight from the LLM pass, no classify_tree call.
    assert all(o["basis"] in ("judge", "deterministic") for o in raw_obs)


def test_llm_segmented_normalized_tree_is_written(tmp_path: Path) -> None:
    """The per-version ClauseTree written to normalized/ reflects the LLM tree
    (two top-level clauses) and carries the real document_id/version identity
    — not run_gates' "doc"/"v1" placeholder defaults.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    tree_path = out_dir / "normalized" / "deal-001" / "v1.clauses.json"
    assert tree_path.exists()
    tree = ClauseTree.load(tree_path)
    assert tree.document_id == "deal-001"
    assert tree.version == "v1"
    assert tree.source_file == "v1.rtf"
    assert [n.heading for n in tree.nodes] == ["1. Indemnification", "2. Governing Law"]


# ---------------------------------------------------------------------------
# Issue #86: LLM-segmenter taxonomy assignments must not be asserted as
# confidence=1.0/basis="judge" — a single unverified LLM pass over untrusted
# counterparty text is not a real judge verdict, and downstream
# confidence-based review gating must actually be able to flag it.
# ---------------------------------------------------------------------------


def test_classified_from_taxonomy_by_path_is_not_asserted_confidence_1_judge() -> None:
    """An LLM-assigned taxonomy_id must not come back as confidence=1.0/basis="judge".

    That combination masquerades a single unverified LLM pass as a real,
    separately-verified judge verdict, and is indistinguishable downstream
    from an actual ClassificationJudge call. Assert the replacement: a
    dedicated basis distinct from "judge", and a confidence low enough that
    ``ClauseClassification.is_ambiguous`` (below ``AMBIGUITY_THRESHOLD``) is
    True — i.e. this assignment is flagged as needing review, not treated as
    certain.
    """
    tree = ClauseTree(
        document_id="deal-001",
        version="v1",
        source_file="v1.rtf",
        nodes=[
            ClauseNode(
                clause_path="1",
                heading="Indemnification",
                text="Alpha Corp shall indemnify Beta University.",
                char_span=(0, 10),
            )
        ],
    )
    classified = _classified_from_taxonomy_by_path(tree, {"1": "indemnification"})

    assert len(classified) == 1
    cc = classified[0].classification
    assert cc.taxonomy_id == "indemnification"
    assert not (cc.confidence == 1.0 and cc.basis == "judge"), (
        "LLM-segmented taxonomy assignment must not be asserted as a real "
        "judge verdict (confidence=1.0, basis='judge')"
    )
    assert cc.basis != "judge"
    assert cc.confidence < AMBIGUITY_THRESHOLD
    assert cc.is_ambiguous, "a below-threshold LLM-segmenter confidence must read as ambiguous"


def test_llm_segmentation_low_confidence_surfaces_in_after_action_review(tmp_path: Path) -> None:
    """A below-threshold LLM-segmenter classification confidence must actually
    reach the after-action report's needs-attention section — the concrete
    downstream confidence-based review gate in this repo (aar._build_needs_attention,
    fed by classification_confidences -> build_observations -> Observation.confidence).

    Before the fix, confidence=1.0 meant these clauses could never trip this
    gate; the LLM misclassifying an indemnification clause would present as a
    certainty in the observation store.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )
    project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    classified_obs = [o for o in raw_obs if o["taxonomy_id"] is not None]
    assert classified_obs, "must have at least one LLM-classified observation to check"
    assert all(o["confidence"] < 0.5 for o in classified_obs), (
        "LLM-segmented classifications must carry a confidence below the "
        "after-action report's low-confidence threshold"
    )

    data = build_after_action_data(out_dir)
    needs_attention = data["needs_attention"]
    assert any("low confidence" in "".join(item.get("reasons", [])) for item in needs_attention), (
        f"expected a low-confidence needs_attention item; got {needs_attention}"
    )


# ---------------------------------------------------------------------------
# AC-2: downstream contract — diff/deviation + playbook validate
# ---------------------------------------------------------------------------


def test_llm_segmentation_two_versions_diff_and_playbook_validate(tmp_path: Path) -> None:
    """Two LLM-segmented versions run through diff/deviation without error,
    and the resulting playbook passes playbook-schema validation.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=True)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs, "two-version diff must still produce observations"
    # The signed version (v2) has 2 real clauses + the "Signatures" noise
    # node (taxonomy_id=None) — deviation assessment must not choke on the
    # unclassified node, and classified clauses must still show up.
    assert {"indemnification", "governing_law"} <= {o["taxonomy_id"] for o in raw_obs}

    playbook = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    result = validate_document(playbook)
    assert result.ok, f"playbook validate must pass: {[str(e) for e in result.errors]}"


# ---------------------------------------------------------------------------
# AC-3: QA-gate failure fails loud but ISOLATED — the failing document is
# quarantined (recorded in quarantine.json), not silently degraded to the
# deterministic segmenter and not aborting the whole corpus run.
# ---------------------------------------------------------------------------


def test_gate_failing_segment_fn_quarantines_document(tmp_path: Path) -> None:
    """A segment_fn whose output never passes the QA gates must quarantine that
    document — recorded in quarantine.json with a SegmentationQAError reason —
    rather than silently degrading the tree, silently dropping the version, or
    aborting the whole run. Here the only document fails, so the run completes
    with zero observations but a populated quarantine.json.

    Also (issue #83): the quarantined version must still get a durable
    corpus_manifest.json row — status="failed", the QA error text, and a
    non-None extractor — instead of the extractor label being erased and the
    version falling through to the "not attempted" default fill (the version
    WAS attempted: extraction succeeded and segmentation ran, only the QA
    gate failed). Checked against the WRITTEN manifest file on disk, not an
    in-memory dict — asserting on the in-memory structure is precisely the
    gap that made an earlier attempt at this fix a silent no-op. The
    document must remain visibly quarantined, never masquerade as a normal
    scope-gated document: in_scope=False + scope_rationale (OPF §3.6) and an
    explicit x_quarantined marker.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    # Does NOT raise — the run completes despite the QA failure.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_gate_failing_segment_fn,
    )

    assert not read_observations_jsonl(out_dir / "observations.jsonl"), (
        "the quarantined document must contribute no observations"
    )
    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]
    assert "SegmentationQAError" in quarantine[0]["reason"]

    # issue #83 — read the WRITTEN corpus_manifest.json, not an in-memory dict.
    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    assert [d["document_id"] for d in manifest] == ["deal-001"]
    doc = manifest[0]
    assert doc["in_scope"] is False
    assert doc["scope_rationale"]
    assert doc["x_quarantined"] is True
    assert doc["versions_mined"] == 0, "no version ingested successfully before the QA failure"
    assert [v["version"] for v in doc["version_ingest"]] == ["v1"]
    row = doc["version_ingest"][0]
    assert row["status"] == "failed"
    assert row["extractor"] == "rtf", (
        "extractor must survive even though segmentation never succeeded"
    )
    assert "reason" in row  # issue #81 shape preserved (None: no ExtractorLabel resolved)
    assert row["reason"] is None
    # Same QA error text quarantine.json recorded (minus the exception type
    # name prefix quarantine.json's f"{type(exc).__name__}: {exc}" adds).
    assert quarantine[0]["reason"] == f"SegmentationQAError: {row['error']}"
    # The error text is built only from gate names/offsets/lengths, never a
    # raw slice of source content (issue #83) — assert the shape rather than
    # just "some non-empty string", so a future regression that re-embeds
    # canonical_text is caught here too, not just by the dedicated
    # entity-leak test below.
    assert "coverage gate" in row["error"]
    assert "char(s) uncovered" in row["error"]


def test_gate_failing_segment_fn_quarantine_preserves_earlier_successful_version(
    tmp_path: Path,
) -> None:
    """issue #83: when only a LATER version fails QA, the quarantined
    document's partial corpus_manifest.json record must still carry the
    EARLIER version's genuine "ok" row alongside the failing version's row —
    proving the partial record is a snapshot of everything this loop got
    through, not just the one version that failed.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=True)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fail_only_second_version_segment_fn,
    )

    assert not read_observations_jsonl(out_dir / "observations.jsonl"), (
        "the quarantined document must contribute no observations, even "
        "though its first version individually ingested fine"
    )
    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    doc = manifest[0]
    assert doc["in_scope"] is False
    assert doc["x_quarantined"] is True
    assert doc["versions_mined"] == 1, "v1 ingested fine before v2's QA failure aborted L1"
    rows_by_version = {v["version"]: v for v in doc["version_ingest"]}
    assert set(rows_by_version) == {"v1", "v2"}
    assert rows_by_version["v1"]["status"] == "ok"
    assert rows_by_version["v1"]["extractor"] is not None
    assert rows_by_version["v2"]["status"] == "failed"
    assert rows_by_version["v2"]["extractor"] == "rtf"
    assert quarantine[0]["reason"] == f"SegmentationQAError: {rows_by_version['v2']['error']}"


def test_quarantined_document_partial_record_survives_compile(tmp_path: Path) -> None:
    """issue #83 amendment guardrail: a quarantined document's partial
    corpus_doc must not corrupt or masquerade through ``project_playbook``'s
    assemble+validate — the compile/publish path the amendment explicitly
    calls out. A mixed corpus (one document mines cleanly, the other is
    QA-quarantined) must still produce a valid ``playbook.opf.json``: the
    quarantined document contributes zero evidence, is excluded from the
    in-scope stat, and its ``version_ingest`` survives
    ``playbook_assembler._sanitize_corpus_documents_for_schema`` (which
    strips engine-internal keys like "reason" before publication) without
    tripping schema/normative validation.
    """
    corpus_dir, config_path, out_dir = _make_mixed_corpus_with_known_entities(
        tmp_path, "deal-002", known_entities=[]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_dispatching_segment_fn,
    )

    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-002"]

    # Observations only ever cite the document that actually mined.
    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs, "the successfully-mined document must still contribute observations"
    assert {o["citation"]["document_id"] for o in raw_obs} == {"deal-001"}

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    by_id = {d["document_id"]: d for d in manifest}
    assert set(by_id) == {"deal-001", "deal-002"}
    assert by_id["deal-001"]["in_scope"] is True
    assert "x_quarantined" not in by_id["deal-001"]
    assert by_id["deal-002"]["in_scope"] is False
    assert by_id["deal-002"]["x_quarantined"] is True

    # project_playbook self-validates (assemble_playbook -> validate_document)
    # and raises AssemblyError on any blocking failure — must not raise.
    playbook = project_playbook(out_dir, cfg, taxonomy)
    assert validate_document(playbook).ok

    docs_by_id = {d["document_id"]: d for d in playbook["corpus"]["documents"]}
    assert docs_by_id["deal-002"]["in_scope"] is False
    assert docs_by_id["deal-002"]["x_quarantined"] is True
    assert docs_by_id["deal-002"]["scope_rationale"]
    # Sanitized for publication (issue #81's stripping contract): only the
    # closed schema key set survives on each version_ingest entry — "reason"
    # (engine-internal) must NOT leak into the published playbook.
    for vi in docs_by_id["deal-002"]["version_ingest"]:
        assert set(vi) <= {"version", "status", "error", "extractor"}

    # The quarantined document must not inflate playbook-level stats as if
    # it were a normal scope-gated document.
    assert playbook["corpus"]["stats"]["documents_total"] == 2
    assert playbook["corpus"]["stats"]["documents_in_scope"] == 1

    # No clause position in the published evidence can trace back to the
    # quarantined document — it produced zero observations, so none of its
    # content can appear as evidence.
    cited_doc_ids = {
        obs["example_ref"]["document_id"]
        for clause in playbook["evidence"]["clauses"]
        for obs in clause.get("observed_positions", [])
    }
    assert "deal-002" not in cited_doc_ids


def test_quarantine_error_never_leaks_raw_entity_name(tmp_path: Path) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #83). A QA-quarantined document's
    version_ingest[].error carries the SegmentationQAError message verbatim.
    Before this fix, the coverage gate's message embedded a raw slice of
    source text (``canonical_text[cursor:]``, rendered with ``!r``) — so with
    ``provenance.known_entities`` configured, a QA-quarantined document whose
    body names that entity as the FIRST TOKEN of the uncovered text (see
    ``_ENTITY_LEAK_BODY``'s docstring for exactly why that specific shape
    matters) leaked the RAW counterparty name into corpus_manifest.json,
    playbook.opf.json, AND review.json — and would trip publisher's hard,
    unsuppressible step-4 entity backstop on compile. This fix instead never
    embeds a raw text slice in the QA error message at all (see
    ``segmentation_qa._check_coverage``), so no alias-matching edge case can
    reintroduce this leak. corpus_manifest.json is never length-truncated
    (there is no length cap anywhere in this fix — the message is already
    short, gate-name/offset/length-only), so proving the raw name is absent
    there specifically proves the fix is "never embed it", not "truncate it
    small enough not to notice".
    """
    corpus_dir, config_path, out_dir = _make_single_doc_corpus_with_known_entities(
        tmp_path, "deal-001", _ENTITY_LEAK_BODY, known_entities=[_ENTITY_LEAK_NAME]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_gate_failing_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]

    # issue #83: read the WRITTEN corpus_manifest.json, not an in-memory dict.
    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    manifest = json.loads(manifest_text)
    row = manifest[0]["version_ingest"][0]
    assert row["status"] == "failed"
    assert _ENTITY_LEAK_NAME not in manifest_text, (
        "raw entity name must never reach corpus_manifest.json via a "
        "quarantined document's version_ingest[].error"
    )
    # Sanity: the dangerous input really did trigger the coverage gate's
    # trailing-gap path (not some other gate) — otherwise this test would
    # vacuously pass without exercising the code this fix touched.
    assert "coverage gate" in row["error"]
    assert "trailing gap" in row["error"]

    playbook = project_playbook(out_dir, cfg, taxonomy)
    published_text = (out_dir / "playbook.opf.json").read_text(encoding="utf-8")
    assert _ENTITY_LEAK_NAME not in published_text
    assert _ENTITY_LEAK_NAME not in json.dumps(playbook), "double-check the in-memory dict too"

    # publisher's hard, unsuppressible step-4 backstop (_entity_backstop_scan)
    # must find zero hits — the compiled playbook must already be born-safe
    # by the time it reaches publish.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [_ENTITY_LEAK_NAME], "sanity: the entity really was registered"
    assert publisher_module._entity_backstop_scan(playbook, known_entity_names) == []


def test_taxonomy_gate_error_never_leaks_raw_entity_name(tmp_path: Path) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #97). The taxonomy gate
    (``segmentation_qa._check_taxonomy``) used to echo a REJECTED
    ``taxonomy_id`` verbatim via ``!r``. That value is LLM-assigned, and the
    gate fires precisely when it is OUTSIDE the allowed vocabulary — so the
    one case that reaches this message is the case where the model returned
    something unconstrained. Since #83, a ``SegmentationQAError`` message is
    persisted verbatim into ``corpus_manifest.json``'s
    ``version_ingest[].error`` for a quarantined document, and that file is
    compiled straight into ``playbook.opf.json`` — so arbitrary model output
    had a path into the compiled artifact.

    ``_taxonomy_leak_segment_fn`` assigns ``_TAXONOMY_LEAK_REJECTED_ID`` —
    ``_TAXONOMY_LEAK_ENTITY_NAME`` glued directly to a trailing "_2023" with
    NO separating space, the same shape issue #96 proved defeats
    ``entity_registry._fuzzy_name_pattern``'s whole-word boundary scrub —
    proving this fix closes the leak at the source
    (``segmentation_qa._safe_taxonomy_id_repr``) rather than relying on that
    scrub, which cannot catch a glued token.
    """
    corpus_dir, config_path, out_dir = _make_single_doc_corpus_with_known_entities(
        tmp_path, "deal-001", _BAD_DOC_BODY, known_entities=[_TAXONOMY_LEAK_ENTITY_NAME]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_taxonomy_leak_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]

    # issue #83 precedent: read the WRITTEN corpus_manifest.json, not an
    # in-memory dict.
    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    manifest = json.loads(manifest_text)
    row = manifest[0]["version_ingest"][0]
    assert row["status"] == "failed"
    assert _TAXONOMY_LEAK_ENTITY_NAME not in manifest_text, (
        "raw entity name must never reach corpus_manifest.json via the "
        "taxonomy gate's rejected-id message"
    )
    assert _TAXONOMY_LEAK_REJECTED_ID not in manifest_text
    # Sanity: the dangerous input really did trigger the taxonomy gate (not
    # some other gate) — otherwise this test would vacuously pass without
    # exercising the code this fix touched.
    assert "taxonomy gate" in row["error"]
    assert "not in the allowed taxonomy_ids" in row["error"]
    # Acceptance criterion: the message still names the offending clause path.
    assert "'1'" in row["error"]
    # Acceptance criterion: still diagnostically useful — the placeholder
    # names the withheld value's length even though it withholds the value.
    assert f"{len(_TAXONOMY_LEAK_REJECTED_ID)} chars" in row["error"]

    playbook = project_playbook(out_dir, cfg, taxonomy)
    published_text = (out_dir / "playbook.opf.json").read_text(encoding="utf-8")
    assert _TAXONOMY_LEAK_ENTITY_NAME not in published_text
    assert _TAXONOMY_LEAK_ENTITY_NAME not in json.dumps(playbook), (
        "double-check the in-memory dict too"
    )

    # publisher's hard, unsuppressible step-4 backstop (_entity_backstop_scan)
    # must find zero hits — the compiled playbook must already be born-safe
    # by the time it reaches publish.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [_TAXONOMY_LEAK_ENTITY_NAME], (
        "sanity: the entity really was registered"
    )
    assert publisher_module._entity_backstop_scan(playbook, known_entity_names) == []


def test_taxonomy_gate_non_string_id_quarantines_with_taxonomy_reason(tmp_path: Path) -> None:
    """REGRESSION (issue #97 fix round 1). Before this fix,
    ``segmentation_qa._safe_taxonomy_id_repr`` called
    ``_TAXONOMY_ID_SHAPE_RE.fullmatch(taxonomy_id)``/``len(taxonomy_id)``
    unguarded, both of which raise ``TypeError`` on a non-``str``
    ``taxonomy_id`` — exactly the value a model response that violates its
    own structured-output schema for this field can produce
    (``llm_segmenter._parse_seg_nodes`` assigns ``raw["taxonomy_id"]`` with
    no type check into a plain, unvalidated ``SegNode`` dataclass).

    That ``TypeError`` escapes ``run_gates``/``segment_verify_repair`` (they
    catch only ``SegmentationQAError``) and used to be caught by
    ``pipeline._compute_doc_result``'s broad ``except Exception`` as a
    per-version warning instead — on this single-version document, that
    means ``_compute_doc_result`` returns ``None`` (no version survived), so
    ``mine_corpus`` quarantines it under the WRONG, generic reason "all
    versions failed extraction/ingest" and ``corpus_manifest.json`` is
    written as ``[]`` — the document's extractor label and QA status vanish
    entirely, exactly what issue #83 existed to prevent (see
    ``pipeline.py``'s ``except SegmentationQAError`` comment: "a QA-gate
    failure on the LLM path must never be swallowed into a per-file warning
    + skipped version"). This test proves the fix restores that fail-loud
    contract: a clean ``SegmentationQAError`` naming the taxonomy gate, with
    the #83 partial record intact in ``corpus_manifest.json``.
    """
    corpus_dir, config_path, out_dir = _make_single_doc_corpus_with_known_entities(
        tmp_path, "deal-001", _BAD_DOC_BODY, known_entities=[]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_taxonomy_non_string_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    # Acceptance criterion (core regression): quarantined with a
    # TAXONOMY-GATE reason, not the generic "all versions failed
    # extraction/ingest" a swallowed TypeError used to produce.
    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]
    reason = quarantine[0]["reason"]
    assert "SegmentationQAError" in reason
    assert "taxonomy gate" in reason
    assert "all versions failed extraction/ingest" not in reason

    # Acceptance criterion (issue #83 partial record): corpus_manifest.json
    # must NOT be "[]" — the document's extractor label and QA status must
    # survive, exactly what a swallowed TypeError used to erase.
    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    manifest = json.loads(manifest_text)
    assert manifest != [], (
        "corpus_manifest.json must not be empty — a swallowed TypeError "
        "used to make the document vanish from it entirely"
    )
    assert [d["document_id"] for d in manifest] == ["deal-001"]
    row = manifest[0]["version_ingest"][0]
    assert row["status"] == "failed"
    assert "taxonomy gate" in row["error"]
    assert "not in the allowed taxonomy_ids" in row["error"]
    # The placeholder still names the withheld value's type, proving the
    # gate went through _safe_taxonomy_id_repr's guarded path rather than
    # crashing (or, worse, blindly echoing the raw model value).
    assert "non-string value, type int" in row["error"]


def test_taxonomy_gate_list_id_quarantines_with_taxonomy_reason(tmp_path: Path) -> None:
    """REGRESSION (issue #97 fix round 2). The round-1 fix above guards
    ``segmentation_qa._safe_taxonomy_id_repr`` against a non-``str``
    ``taxonomy_id``, but that guard was unreachable for an UNHASHABLE value:
    ``_check_taxonomy`` evaluated ``taxonomy_id not in allowed`` against a
    ``set`` BEFORE ever calling ``_safe_taxonomy_id_repr``, so a ``list``
    (e.g. a model assigning two taxonomy categories to one node — a highly
    plausible schema violation, unlike an arbitrary scalar) raised an
    uncaught ``TypeError: unhashable type`` at that membership test.

    That ``TypeError`` escapes ``run_gates``/``segment_verify_repair`` (they
    catch only ``SegmentationQAError``) and used to be caught by
    ``pipeline._compute_doc_result``'s broad ``except Exception`` as a
    per-version warning instead — on this single-version document, that
    means ``_compute_doc_result`` returns ``None`` (no version survived), so
    ``mine_corpus`` quarantines it under the WRONG, generic reason "all
    versions failed extraction/ingest" and ``corpus_manifest.json`` is
    written as ``[]`` — the document's extractor label and QA status vanish
    entirely, exactly what issue #83 existed to prevent (see
    ``pipeline.py``'s ``except SegmentationQAError`` comment: "a QA-gate
    failure on the LLM path must never be swallowed into a per-file warning
    + skipped version"). This test proves the round-2 fix (checking
    ``isinstance`` in ``_check_taxonomy`` itself, before the ``set`` lookup)
    restores that fail-loud contract for an unhashable id too: a clean
    ``SegmentationQAError`` naming the taxonomy gate, with the #83 partial
    record intact in ``corpus_manifest.json``.
    """
    corpus_dir, config_path, out_dir = _make_single_doc_corpus_with_known_entities(
        tmp_path, "deal-001", _BAD_DOC_BODY, known_entities=[]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_taxonomy_list_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    # Acceptance criterion (core regression): quarantined with a
    # TAXONOMY-GATE reason, not the generic "all versions failed
    # extraction/ingest" a swallowed TypeError used to produce.
    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]
    reason = quarantine[0]["reason"]
    assert "SegmentationQAError" in reason
    assert "taxonomy gate" in reason
    assert "all versions failed extraction/ingest" not in reason

    # Acceptance criterion (issue #83 partial record): corpus_manifest.json
    # must NOT be "[]" — the document's extractor label and QA status must
    # survive, exactly what a swallowed TypeError used to erase.
    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    manifest = json.loads(manifest_text)
    assert manifest != [], (
        "corpus_manifest.json must not be empty — a swallowed TypeError "
        "used to make the document vanish from it entirely"
    )
    assert [d["document_id"] for d in manifest] == ["deal-001"]
    row = manifest[0]["version_ingest"][0]
    assert row["status"] == "failed"
    assert "taxonomy gate" in row["error"]
    assert "not in the allowed taxonomy_ids" in row["error"]
    # The placeholder still names the withheld value's type, proving the
    # gate went through _safe_taxonomy_id_repr's guarded path rather than
    # crashing (or, worse, blindly echoing the raw model value).
    assert "non-string value, type list" in row["error"]


# A document folder named after a known entity, glued directly to a trailing
# "_2023" with NO space/separator before the underscore (issue #96 regression
# test). Underscore counts as a \w character, so this is deliberately chosen
# to defeat entity_registry._fuzzy_name_pattern's (?<!\w)NAME(?!\w) whole-word
# boundary check on the trailing side — a realistic corpus-folder naming
# convention (underscore-joined year suffix), not a contrived one. This
# proves the regression test below is load-bearing on version_orderer's
# "never embed the path" fix itself, not incidentally saved by pipeline.py's
# separate defense-in-depth pseudonymization pass over quarantine.json's
# reason text (see that pass's own comment in pipeline.py: never make the
# scrub load-bearing — mirrors the #83/#95 precedent this ticket (#96) cites).
_HINTS_LEAK_ENTITY_NAME = "Fictional University"
_HINTS_LEAK_DOC_ID = f"{_HINTS_LEAK_ENTITY_NAME}_2023"


def test_hints_error_never_leaks_raw_entity_name_via_path(tmp_path: Path) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #96). version_orderer.Hints.load
    used to raise HintsError with the full hints.yaml path embedded in the
    message (e.g. f"{path}: not valid YAML: {exc}"). That path is
    <corpus>/<document-folder>/hints.yaml, and the document folder is
    typically named after the counterparty (doc_id == doc_dir.name, per
    pipeline._compute_doc_result) — so a document quarantined by a malformed
    hints.yaml, in a corpus where the document folder is named after a
    provenance.known_entities entry, used to write the raw entity name
    straight into quarantine.json's reason field — even though #83 already
    aliases quarantine.json's document_id field separately — and from there
    into the after-action report too (aar._load_quarantine ->
    _build_needs_attention embeds quarantine's reason verbatim; see
    playbook_engine/aar.py:576). (HintsError carries no partial_corpus_doc —
    unlike SegmentationQAError — so it never reaches corpus_manifest.json at
    all; see _make_mixed_corpus_with_malformed_hints.) This fix instead
    names the file by its bare filename ("hints.yaml"), never embedding the
    document-folder path, so no alias-matching edge case is needed to keep
    the name out (see version_orderer.HintsError's own docstring for why
    this specific folder name defeats the defense-in-depth scrub, proving
    this isn't a scrub-only fix).
    """
    corpus_dir, config_path, out_dir = _make_mixed_corpus_with_malformed_hints(
        tmp_path, _HINTS_LEAK_DOC_ID, known_entities=[_HINTS_LEAK_ENTITY_NAME]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        entity_registry_path=tmp_path / "registry.json",
    )

    # issue #83 precedent: read the WRITTEN quarantine.json, not an in-memory dict.
    quarantine_text = (out_dir / "quarantine.json").read_text(encoding="utf-8")
    quarantine = json.loads(quarantine_text)
    assert len(quarantine) == 1
    assert _HINTS_LEAK_ENTITY_NAME not in quarantine_text, (
        "raw entity name must never reach quarantine.json via a malformed "
        "hints.yaml's HintsError, whose message used to embed the full "
        "hints.yaml path (whose parent directory is the document folder, "
        "typically named after the counterparty)"
    )
    # Sanity: the dangerous input really did trigger HintsError on the
    # malformed-YAML path (not some other quarantine reason) — otherwise this
    # test would vacuously pass without exercising the code this fix touched.
    reason = quarantine[0]["reason"]
    assert reason.startswith("HintsError:")
    # Acceptance criterion: diagnostically useful — a human can still tell
    # HOW the file was malformed even with the path gone. WHICH document is
    # still identifiable via the record's own document_id field (#83).
    assert "not valid YAML" in reason
    assert "hints.yaml" in reason

    # corpus_manifest.json never gets an entry for the quarantined document at
    # all (HintsError carries no partial_corpus_doc — see the docstring
    # above); this just confirms the surviving "deal-001" entry doesn't
    # incidentally carry the other document's entity name either.
    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    assert _HINTS_LEAK_ENTITY_NAME not in manifest_text

    # The after-action report IS the verified second consumer of
    # quarantine.json's reason field (aar._load_quarantine ->
    # _build_needs_attention, playbook_engine/aar.py:576, embeds q["reason"]
    # verbatim into "needs_attention" reasons) — read the WRITTEN
    # report.md/report.json, not an in-memory dict.
    write_after_action_report(out_dir, out_dir / "report.md")
    report_md = (out_dir / "report.md").read_text(encoding="utf-8")
    report_json = (out_dir / "report.json").read_text(encoding="utf-8")
    assert _HINTS_LEAK_ENTITY_NAME not in report_md
    assert _HINTS_LEAK_ENTITY_NAME not in report_json

    # Belt-and-braces (issue #96 review correction: an earlier attempt's
    # comments overstated this as THE documented propagation path — it
    # isn't; the after-action report above is). playbook.opf.json has no
    # quarantine section, and quarantine.json is not among corpus_manifest's
    # sources either — so neither artifact was ever reachable by this leak.
    # These assertions stay true either way; they just aren't proof of
    # anything this fix changed.
    playbook = project_playbook(out_dir, cfg, taxonomy)
    published_text = (out_dir / "playbook.opf.json").read_text(encoding="utf-8")
    assert _HINTS_LEAK_ENTITY_NAME not in published_text
    assert _HINTS_LEAK_ENTITY_NAME not in json.dumps(playbook), (
        "double-check the in-memory dict too"
    )

    # publisher's hard, unsuppressible step-4 backstop (_entity_backstop_scan)
    # must find zero hits — the compiled playbook must already be born-safe
    # by the time it reaches publish.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [_HINTS_LEAK_ENTITY_NAME], (
        "sanity: the entity really was registered"
    )
    assert publisher_module._entity_backstop_scan(playbook, known_entity_names) == []


def test_hints_error_never_leaks_raw_entity_name_via_yaml_content_snippet(
    tmp_path: Path,
) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #96, fix round 1). HintsError's
    "not valid YAML" branch used to interpolate ``{exc}`` (the underlying
    ``yaml.YAMLError``) in full, and PyYAML's own ``__str__`` embeds a raw
    snippet of hints.yaml's SOURCE CONTENT around the parse error via
    ``Mark.get_snippet()`` (e.g. '...line 1, column 8:\\n order: [Fictional
    University_2023_v1.do ... \\n           ^...'). This is a leak vector the
    {path}->{path.name} fix alone does NOT address (it only touches the
    path, never {exc}) — so this test uses a document folder that carries
    NO entity name at all (isolating it from the {path} leak covered by the
    test above) and instead puts the entity name INSIDE hints.yaml's own
    malformed ``order:`` content, GLUED directly to a trailing
    "_2023_v1.docx" suffix with no separating space — the extension-
    inclusive ``order:`` convention docs/CORPUS-LAYOUT.md documents, and the
    same glued shape ``_HINTS_LEAK_DOC_ID`` uses above to defeat
    ``entity_registry._fuzzy_name_pattern``'s whole-word boundary check.

    Because that glued form has no true word boundary on its trailing side,
    pipeline.py's defense-in-depth pseudonymization pass over
    quarantine.json's ``reason`` field CANNOT catch it (``\\s+``-joined
    whole-word matching cannot cross "y_2023"). So this vector can only be
    closed by never embedding the parser's raw snippet in the first place
    (``version_orderer._yaml_error_detail``, built from the exception's
    TYPE plus line/column only, never ``str(exc)``) — not by the scrub.
    """
    entity_name = "Fictional University"
    hints_yaml_text = f"order: [{entity_name}_2023_v1.docx, unterminated\n"
    corpus_dir, config_path, out_dir = _make_mixed_corpus_with_malformed_hints(
        tmp_path, "deal-002", known_entities=[entity_name], hints_yaml_text=hints_yaml_text
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        entity_registry_path=tmp_path / "registry.json",
    )

    # issue #83 precedent: read the WRITTEN quarantine.json, not an in-memory dict.
    quarantine_text = (out_dir / "quarantine.json").read_text(encoding="utf-8")
    quarantine = json.loads(quarantine_text)
    assert len(quarantine) == 1
    reason = quarantine[0]["reason"]
    # Sanity: the dangerous input really did trigger HintsError on the
    # malformed-YAML path (not some other quarantine reason) — otherwise
    # this test would vacuously pass without exercising the branch this fix
    # touched.
    assert reason.startswith("HintsError:")
    assert "not valid YAML" in reason
    assert "hints.yaml" in reason
    assert entity_name not in quarantine_text, (
        "raw entity name must never reach quarantine.json via HintsError's "
        "'not valid YAML' branch, even when glued to a trailing suffix that "
        "defeats the whole-word defense-in-depth scrub — this vector must "
        "be closed at the source, not rely on the scrub"
    )
    # Acceptance criterion: diagnostically useful — a human can still tell
    # HOW the file was malformed (down to line/column) even with the raw
    # snippet gone.
    assert "line" in reason and "column" in reason

    write_after_action_report(out_dir, out_dir / "report.md")
    report_md = (out_dir / "report.md").read_text(encoding="utf-8")
    report_json = (out_dir / "report.json").read_text(encoding="utf-8")
    assert entity_name not in report_md
    assert entity_name not in report_json

    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    assert entity_name not in manifest_text

    playbook = project_playbook(out_dir, cfg, taxonomy)
    published_text = (out_dir / "playbook.opf.json").read_text(encoding="utf-8")
    assert entity_name not in published_text
    assert entity_name not in json.dumps(playbook), "double-check the in-memory dict too"

    # publisher's hard, unsuppressible step-4 backstop (_entity_backstop_scan)
    # must find zero hits — the compiled playbook must already be born-safe
    # by the time it reaches publish.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [entity_name], "sanity: the entity really was registered"
    assert publisher_module._entity_backstop_scan(playbook, known_entity_names) == []


def test_hints_error_never_leaks_raw_entity_name_via_yaml_problem_token(
    tmp_path: Path,
) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #96, fix round 2 — the leak the
    round-1 fix MISSED). Even with ``str(exc)``'s raw source snippet gone
    (see the sibling ``..._via_yaml_content_snippet`` test above),
    ``yaml.YAMLError.problem`` ALONE is not safe to interpolate either: it
    reads as a fixed grammar-violation phrase in the common case, but
    PyYAML's ``composer.py`` builds it with ``%r`` around a token taken
    straight from the document being parsed for an undefined-alias
    reference: ``"found undefined alias %r" % anchor``. hints.yaml's
    documented grammar (docs/CORPUS-LAYOUT.md) never uses YAML
    anchors/aliases, but ``yaml.safe_load`` still SCANS the full YAML
    grammar regardless of what's documented (``SafeLoader`` only restricts
    which Python types get *constructed*, never which syntax gets
    *scanned/parsed*), so a hand-typed slip like this (a plausible
    copy/paste error, not the documented grammar) reaches this branch for
    real.

    Uses the REQUIRED glued form ``order: [*Northwind_2023]`` (NOT
    ``order: [*Northwind]``) — a single-word known entity name
    ("Northwind") glued directly to a trailing "_2023" year suffix with NO
    separating space, mirroring the same document-folder/order-list naming
    convention the sibling tests above use. This is deliberate: the
    un-glued form is the ONE token shape
    ``entity_registry._fuzzy_name_pattern``'s whole-word
    ``(?<!\\w)NAME(?!\\w)`` boundary CAN match, so a fixture using it would
    only prove the defense-in-depth scrub works (see
    ``test_quarantine_reason_defense_in_depth_pseudonymization_fires``
    below, which proves exactly that with a DIFFERENT exception type), not
    that THIS vector is closed at the source. The glued form defeats that
    boundary (``_`` is itself a ``\\w`` character) and can only be closed
    by ``version_orderer._yaml_error_detail`` never embedding ``problem``
    in the first place — confirmed by mutation in this ticket's own
    verification (reverting ONLY ``_yaml_error_detail``'s body to include
    ``exc.problem`` reproduces the raw name in quarantine.json).
    """
    entity_name = "Northwind"
    corpus_dir, config_path, out_dir = _make_mixed_corpus_with_malformed_hints(
        tmp_path,
        "deal-002",
        known_entities=[entity_name],
        hints_yaml_text=f"order: [*{entity_name}_2023]\n",
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        entity_registry_path=tmp_path / "registry.json",
    )

    # issue #83 precedent: read the WRITTEN quarantine.json, not an in-memory dict.
    quarantine_text = (out_dir / "quarantine.json").read_text(encoding="utf-8")
    quarantine = json.loads(quarantine_text)
    assert len(quarantine) == 1
    reason = quarantine[0]["reason"]
    # Sanity: the dangerous input really did trigger HintsError on the
    # undefined-alias path (not some other quarantine reason) — otherwise
    # this test would vacuously pass without exercising the branch this fix
    # touched.
    assert reason.startswith("HintsError:")
    assert "not valid YAML" in reason
    assert entity_name not in quarantine_text, (
        "raw entity name must never reach quarantine.json via HintsError's "
        "'not valid YAML' branch, even via an accidental YAML alias "
        "reference glued to a trailing suffix that defeats the whole-word "
        "defense-in-depth scrub — this vector must be closed at the "
        "source, not rely on the scrub"
    )
    # Acceptance criterion: diagnostically useful — a human can still tell
    # HOW the file was malformed (down to line/column) even with the raw
    # token gone.
    assert "line" in reason and "column" in reason

    write_after_action_report(out_dir, out_dir / "report.md")
    report_md = (out_dir / "report.md").read_text(encoding="utf-8")
    report_json = (out_dir / "report.json").read_text(encoding="utf-8")
    assert entity_name not in report_md
    assert entity_name not in report_json

    playbook = project_playbook(out_dir, cfg, taxonomy)
    published_text = (out_dir / "playbook.opf.json").read_text(encoding="utf-8")
    assert entity_name not in published_text
    assert entity_name not in json.dumps(playbook), "double-check the in-memory dict too"

    # publisher's hard, unsuppressible step-4 backstop (_entity_backstop_scan)
    # must find zero hits — the compiled playbook must already be born-safe
    # by the time it reaches publish.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [entity_name], "sanity: the entity really was registered"
    assert publisher_module._entity_backstop_scan(playbook, known_entity_names) == []


def test_quarantine_reason_defense_in_depth_pseudonymization_fires(tmp_path: Path) -> None:
    """issue #96 (defense in depth): ``pipeline.mine_corpus``'s
    pseudonymization pass over ``quarantine.json``'s ``reason`` field had
    ZERO regression coverage before this fix — every quarantine-reason
    privacy test above (both ``test_hints_error_never_leaks_raw_entity_
    name_via_*`` tests, and #83's ``test_quarantine_error_never_leaks_raw_
    entity_name``) deliberately uses a name GLUED to trailing characters
    specifically so ``entity_registry._fuzzy_name_pattern``'s whole-word
    boundary check cannot fire — proving the source-level "never embed it"
    fixes work even when the scrub can't help, but never exercising the
    scrub itself. So that line could be deleted (or silently broken) and
    the rest of the suite would still pass, even though issue #96
    explicitly requires it ("do this as well, not instead") as the
    catch-all for a FUTURE exception type whose message embeds a name in a
    form the scrub CAN match.

    ``NormalizeTrailError`` (llm_segmenter_batch.py) is exactly such a
    type: unlike ``HintsError`` (fixed at the source by this issue),
    nothing about its message is specific to hints.yaml, so a
    caller-injected ``normalize_trail_fn`` can raise it with an arbitrary
    message — here, one that embeds a known entity name as a normal,
    space-separated, word-bounded token (the injection pattern mirrors
    ``test_mine_corpus_forwards_normalize_trail_opt_in`` above). That is
    precisely the shape the scrub CAN catch, so this test proves it
    actually does, in the written artifact — not merely that the raw name
    is absent (which the alias-presence assertion below distinguishes from
    a scrub that merely deleted or corrupted the name).

    NOTE: this issue's own investigation (see its "Out of scope" item)
    already verified that ``NormalizeTrailError``'s REAL call sites in
    llm_segmenter_batch.py build only structural messages (indices/keys/
    block types; ``json.JSONDecodeError.__str__`` omits the document body)
    — none of them actually embed source content today. This test's
    injected message is a deliberately adversarial STAND-IN that proves the
    catch-all mechanism itself works, not a claim that a real call site is
    unsafe.

    Mutation-proof (per issue #96's Notes): with the rest of this fix
    applied, reverting ONLY pipeline.py's
    ``"reason": pseudonymize_text(q["reason"], known_entities, entity_registry)``
    line back to ``"reason": q["reason"]`` makes this test fail with the raw
    entity name present in quarantine.json.
    """
    entity_name = "Fictional University"
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=True)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.provenance.known_entities = [entity_name]

    def _raising_normalize_fn(
        version_trees: dict[str, ClauseTree],
        taxonomy_by_version: dict[str, dict[str, str | None]],
    ) -> NormalizeTrailResult:
        raise NormalizeTrailError(f"model response is not valid JSON: {entity_name} sent bad JSON")

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
        normalize_trail_across_versions=True,
        normalize_trail_fn=_raising_normalize_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    # issue #83 precedent: read the WRITTEN quarantine.json, not an in-memory dict.
    quarantine_text = (out_dir / "quarantine.json").read_text(encoding="utf-8")
    quarantine = json.loads(quarantine_text)
    assert [q["document_id"] for q in quarantine] == ["deal-001"]
    reason = quarantine[0]["reason"]
    # Sanity: the dangerous input really did trigger NormalizeTrailError (not
    # some other quarantine reason) — otherwise this test would vacuously
    # pass without exercising the code this fix touched.
    assert reason.startswith("NormalizeTrailError:")
    assert "sent bad JSON" in reason, "diagnostically useful: failure mode text must survive"
    assert entity_name not in quarantine_text, (
        "raw entity name must never reach quarantine.json's reason field, "
        "even from an exception type this issue does not otherwise fix at "
        "the source — this is exactly what the defense-in-depth "
        "pseudonymization pass exists to catch"
    )
    # Prove the scrub actually FIRED (substituted the alias in), not merely
    # that the raw name is absent — e.g. from an unrelated corruption.
    registry = EntityRegistry.load(tmp_path / "registry.json")
    known_entity_names = list(registry.alias_map().values())
    assert known_entity_names == [entity_name], "sanity: the entity really was registered"
    alias = next(a for a, name in registry.alias_map().items() if name == entity_name)
    assert alias in reason, "the scrub must substitute the alias in, not merely delete the name"


def test_quarantine_and_manifest_document_ids_reconcile_across_pseudonymization(
    tmp_path: Path,
) -> None:
    """issue #83 (item 2 — regression, reproduced): quarantine.json used to
    be written BEFORE the born-safe pseudonymization pass, carrying the RAW
    document_id, while corpus_manifest.json is written after and carries the
    ALIASED id. A quarantined document now ALSO gets a partial
    corpus_manifest.json entry (see pipeline._build_quarantine_corpus_doc),
    and a consumer that cross-references the two files by document_id (e.g.
    aar._build_needs_attention, to avoid double-counting a quarantined
    document as if the compiled playbook covered it too) could never
    recognize the overlap while the ids were spelled differently — silently
    OVERSTATING coverage. Exercised end-to-end through the real pipeline
    (not a hand-written manifest/quarantine fixture) with the quarantined
    document's folder actually NAMED after the known entity, so this proves
    the pipeline reconciles the ids — not just that aar.py's arithmetic is
    correct given already-matching inputs.
    """
    entity_name = "Fictional University"
    corpus_dir, config_path, out_dir = _make_mixed_corpus_with_known_entities(
        tmp_path, entity_name, known_entities=[entity_name]
    )
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_dispatching_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert len(quarantine) == 1
    quarantined_id = quarantine[0]["document_id"]
    assert quarantined_id != entity_name
    assert "fictional" not in quarantined_id.lower(), (
        "quarantine.json's document_id must be aliased, not the raw entity-derived slug"
    )
    assert entity_name not in (out_dir / "quarantine.json").read_text(encoding="utf-8")

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    manifest_ids = [d["document_id"] for d in manifest]
    assert len(manifest_ids) == 2
    assert quarantined_id in manifest_ids, (
        "quarantine.json's (aliased) document_id must match corpus_manifest.json's "
        "partial-record entry EXACTLY — otherwise a consumer that cross-references "
        "the two files by document_id can never recognize the overlap"
    )

    # aar.build_after_action_data's coverage line must not double-count the
    # quarantined document's own partial manifest entry as "covered".
    data = build_after_action_data(out_dir)
    summary = [
        i for i in data["needs_attention"] if any("quarantined vs" in r for r in i["reasons"])
    ]
    assert len(summary) == 1
    # Exactly ONE document (deal-001) genuinely contributed playbook content;
    # the quarantined document's own partial record must not inflate this to 2.
    assert any("covers 1 of 2" in r for r in summary[0]["reasons"]), summary[0]["reasons"]
    report = build_after_action_report(out_dir)
    assert "covers 1 of 2" in report


# ---------------------------------------------------------------------------
# Segmentation verdict cache on the synchronous LLM path (issue #91)
# ---------------------------------------------------------------------------


def test_segmentation_cache_hits_on_sync_path(tmp_path: Path) -> None:
    """A SegmentationVerdictCache passed as segmentation_cache must be honored
    on the synchronous per-document LLM path too (not just the batch
    pre-pass): a second mine_corpus run (with the L1-L4 stage cache disabled
    so _compute_doc_result actually re-executes) hits the segmentation cache
    and never re-invokes the injected segment_fn.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    seg_cache = SegmentationVerdictCache(tmp_path / "seg_cache.jsonl")
    call_count = 0

    def _counting_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
        nonlocal call_count
        call_count += 1
        return _fake_segment_fn(canonical_text, blocks)

    common_kwargs: dict[str, Any] = {
        "corpus_dir": corpus_dir,
        "config": cfg,
        "taxonomy": taxonomy,
        "use_llm_segmentation": True,
        "llm_segment_fn": _counting_segment_fn,
        "segmentation_cache": seg_cache,
        "no_cache": True,  # disable the unrelated L1-L4 stage cache
    }

    mine_corpus(out_dir=out_dir, **common_kwargs)
    assert call_count == 1

    out_dir_2 = tmp_path / "out2"
    mine_corpus(out_dir=out_dir_2, **common_kwargs)
    # No second LLM call for the second run — the content-hash cache
    # satisfied the version without touching segment_fn at all.
    assert call_count == 1

    raw_obs = read_observations_jsonl(out_dir_2 / "observations.jsonl")
    assert {o["taxonomy_id"] for o in raw_obs} == {"indemnification", "governing_law"}


# ---------------------------------------------------------------------------
# refresh_extraction (issue #78) — operator-invoked --no-cache must force
# real re-extraction despite a warm extraction_cache.jsonl, while the judge
# path's internally-forced no_cache=True must NOT (regression guard for
# issue #132's judge-warm-cache intent). Mirrors
# test_segmentation_cache_hits_on_sync_path's two-out_dir/shared-cache shape,
# but counts the underlying RTF extractor call instead of segment_fn calls.
# ---------------------------------------------------------------------------


def test_refresh_extraction_forces_reextraction_despite_warm_cache(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """mine_corpus(no_cache=True, refresh_extraction=True) — the shape
    cli.py's ``mine --no-cache`` produces — must re-invoke the extractor
    for a version whose content is already in a warm ``extraction_cache``.

    Before the fix, ``no_cache`` only gated the L1-L4 ``ArtifactStore``/
    ``JudgmentCache``; ``ExtractionCache`` was threaded through completely
    independently and always hit, so ``--no-cache`` silently replayed the
    same stale blocks.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    extraction_cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    call_count = 0
    real_extract_rtf_lines = extraction._extract_rtf_lines

    def _counting_extract_rtf_lines(path: Path) -> list[tuple[str, int]]:
        nonlocal call_count
        call_count += 1
        return real_extract_rtf_lines(path)

    monkeypatch.setattr(extraction, "_extract_rtf_lines", _counting_extract_rtf_lines)

    common_kwargs: dict[str, Any] = {
        "corpus_dir": corpus_dir,
        "config": cfg,
        "taxonomy": taxonomy,
        "use_llm_segmentation": True,
        "llm_segment_fn": _fake_segment_fn,
        "extraction_cache": extraction_cache,
        "no_cache": True,  # disable the L1-L4 stage cache so _compute_doc_result re-executes
    }

    mine_corpus(out_dir=out_dir, **common_kwargs)
    assert call_count == 1, "first run must extract (cache miss)"

    out_dir_2 = tmp_path / "out2"
    mine_corpus(out_dir=out_dir_2, refresh_extraction=True, **common_kwargs)
    assert call_count == 2, (
        "refresh_extraction=True must force re-extraction despite the warm extraction_cache"
    )

    # The refreshed run's output is unaffected — same source, same result.
    raw_obs = read_observations_jsonl(out_dir_2 / "observations.jsonl")
    assert {o["taxonomy_id"] for o in raw_obs} == {"indemnification", "governing_law"}


def test_judge_forced_no_cache_does_not_force_reextraction(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """mine_corpus(no_cache=True) ALONE — refresh_extraction left at its
    default False, exactly how ``playbook judge``'s cli.py wiring calls it —
    must NOT force re-extraction: the judge path forces no_cache=True to
    bypass the L1-L4 stage cache's stale needs_review sentinels, and that
    must not also force every judge round to re-extract/re-OCR every version
    from scratch (issue #132's judge-warm-cache intent, which this issue
    #78 fix must not regress).
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    extraction_cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")
    call_count = 0
    real_extract_rtf_lines = extraction._extract_rtf_lines

    def _counting_extract_rtf_lines(path: Path) -> list[tuple[str, int]]:
        nonlocal call_count
        call_count += 1
        return real_extract_rtf_lines(path)

    monkeypatch.setattr(extraction, "_extract_rtf_lines", _counting_extract_rtf_lines)

    common_kwargs: dict[str, Any] = {
        "corpus_dir": corpus_dir,
        "config": cfg,
        "taxonomy": taxonomy,
        "use_llm_segmentation": True,
        "llm_segment_fn": _fake_segment_fn,
        "extraction_cache": extraction_cache,
        "no_cache": True,  # the judge path's forced stage-cache bypass
    }

    mine_corpus(out_dir=out_dir, **common_kwargs)
    assert call_count == 1

    out_dir_2 = tmp_path / "out2"
    mine_corpus(out_dir=out_dir_2, **common_kwargs)  # refresh_extraction defaults to False
    assert call_count == 1, (
        "no_cache=True alone must not force re-extraction — extraction_cache must stay warm"
    )


# ---------------------------------------------------------------------------
# Deterministic path unaffected — regression guard
# ---------------------------------------------------------------------------


def test_default_use_llm_segmentation_false_unaffected(tmp_path: Path) -> None:
    """use_llm_segmentation defaults to False: the deterministic segmenter
    path (segment(ingest(...).tree)) must remain the default, unchanged.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    # No use_llm_segmentation kwarg at all — must not require llm_segment_fn
    # and must not touch the LLM path (no fake injected, so any accidental
    # LLM-path call would try to construct a real anthropic client and fail).
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs
    # Deterministic path: the fast-path heading matcher in clause_classifier
    # (exact/Jaccard match against the taxonomy label) still classifies
    # obvious headings like "Indemnification" without any judge — the stub
    # _NullClassificationJudge only affects nodes that reach the judge. The
    # regression guard that matters here is basis: LLM-segmented
    # observations get basis="judge" straight from the LLM pass (see the
    # AC-1 test above); the deterministic path must never produce that
    # basis, since no judge (real or stub) here ever returns basis="judge"
    # (_NullClassificationJudge always returns basis="unclassified").
    bases = {o["basis"] for o in raw_obs}
    assert "judge" not in bases, (
        f"deterministic path must not produce basis='judge' observations; got {bases}"
    )


# ---------------------------------------------------------------------------
# Cross-version normalization wiring (issue #75) — mine_corpus must forward the
# opt-in down to _compute_doc_result, and the normalized labels must be used.
# ---------------------------------------------------------------------------


def test_mine_corpus_forwards_normalize_trail_opt_in(tmp_path: Path) -> None:
    """mine_corpus(normalize_trail_across_versions=True, normalize_trail_fn=...)
    must actually invoke the injected normalize_fn over the version trail and
    feed its normalized taxonomy back into classification.

    Regression guard: the opt-in params were once accepted by mine_corpus but
    never forwarded to _compute_doc_result, so the pass silently never ran.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=True)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    received: list[dict[str, dict[str, str | None]]] = []

    def _fake_normalize_fn(
        version_trees: dict[str, ClauseTree],
        taxonomy_by_version: dict[str, dict[str, str | None]],
    ) -> NormalizeTrailResult:
        # Record the call, then normalize by dropping every "indemnification"
        # label to None — an observable change that proves the returned result
        # is actually used, not merely that the fn was called.
        received.append(taxonomy_by_version)
        remapped = {
            vid: {path: (None if tid == "indemnification" else tid) for path, tid in labels.items()}
            for vid, labels in taxonomy_by_version.items()
        }
        return NormalizeTrailResult(taxonomy_by_version=remapped, boundary_flags=[])

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
        normalize_trail_across_versions=True,
        normalize_trail_fn=_fake_normalize_fn,
    )

    # Called exactly once, over a trail of both versions.
    assert len(received) == 1, "normalize_trail_fn must be invoked once per agreement"
    assert len(received[0]) == 2, "the normalize pass must see every version in the trail"

    # The normalized result was used: the dropped "indemnification" label no
    # longer appears in any observation, while the untouched label survives.
    non_null = {o["taxonomy_id"] for o in read_observations_jsonl(out_dir / "observations.jsonl")}
    non_null.discard(None)
    assert "indemnification" not in non_null, "normalized labels must replace the per-version ones"
    assert "governing_law" in non_null, "untouched labels must survive normalization"


def test_normalize_trail_not_run_for_single_version(tmp_path: Path) -> None:
    """The cross-version pass is a no-op for a single-version agreement
    (len(version_trees) == 1) even when opted in — nothing to normalize across.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    calls: list[dict[str, dict[str, str | None]]] = []

    def _spy_normalize_fn(
        version_trees: dict[str, ClauseTree],
        taxonomy_by_version: dict[str, dict[str, str | None]],
    ) -> NormalizeTrailResult:
        calls.append(taxonomy_by_version)
        return NormalizeTrailResult(taxonomy_by_version=taxonomy_by_version, boundary_flags=[])

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
        normalize_trail_across_versions=True,
        normalize_trail_fn=_spy_normalize_fn,
    )

    assert calls == [], "single-version agreement must not trigger the cross-version pass"


# ---------------------------------------------------------------------------
# Batch segmentation pre-pass wiring (issue #76) — mocked Message Batches
# client exercised through the real segment_documents_batch/mine_corpus
# wiring, not a fake segment_fn.
# ---------------------------------------------------------------------------


def _seg_nodes_to_response_text(seg_nodes: list[SegNode]) -> str:
    """Serialize SegNodes into the structured-output JSON text a batch result carries.

    Mirrors ``llm_segmenter._parse_seg_nodes``'s expected ``{"nodes": [...]}``
    shape exactly (see that function) — the fake batches client below returns
    this text as a "succeeded" result's message content.
    """
    return json.dumps(
        {
            "nodes": [
                {
                    "node_id": n.node_id,
                    "parent_id": n.parent_id,
                    "order": n.order,
                    "heading": n.heading,
                    "taxonomy_id": n.taxonomy_id,
                    "start_block_id": n.start_block_id,
                    "end_block_id": n.end_block_id,
                    "start_quote": n.start_quote,
                    "end_quote": n.end_quote,
                }
                for n in seg_nodes
            ]
        }
    )


def _gate_failing_response_for(blocks: list[Block]) -> list[SegNode]:
    """Same always-failing shape as ``_gate_failing_segment_fn`` above, reused
    by the fake batches client's canned "gate_failing" mode."""
    first = blocks[0]
    return [
        SegNode(
            node_id="n1",
            parent_id=None,
            order=1,
            heading=first.text,
            taxonomy_id=None,
            start_block_id=first.block_id,
            end_block_id=first.block_id,
        )
    ]


class _FakeBatchesResource:
    """Minimal fake of the Anthropic Message Batches surface.

    Computes each custom_id's canned response text lazily from
    *blocks_by_custom_id* at ``.results()`` time via *segment_fn* — the same
    live-recompute-from-blocks approach ``_fake_segment_fn`` itself uses,
    rather than hardcoding block ids into a canned JSON string that could
    silently drift from the fixture text.
    """

    def __init__(
        self,
        blocks_by_custom_id: dict[str, list[Block]],
        segment_fn: Any,
        *,
        gate_failing: bool = False,
    ) -> None:
        self._blocks_by_custom_id = blocks_by_custom_id
        self._segment_fn = segment_fn
        self._gate_failing = gate_failing
        self.create_calls: list[dict[str, Any]] = []
        self.retrieve_calls: list[str] = []
        self.results_calls: list[str] = []
        self._requests_by_batch_id: dict[str, list[dict[str, Any]]] = {}

    def create(self, **kwargs: Any) -> Any:
        self.create_calls.append(kwargs)
        batch_id = f"batch_{len(self.create_calls)}"
        self._requests_by_batch_id[batch_id] = kwargs["requests"]
        return SimpleNamespace(id=batch_id, processing_status="ended")

    def retrieve(self, batch_id: str) -> Any:
        self.retrieve_calls.append(batch_id)
        return SimpleNamespace(id=batch_id, processing_status="ended")

    def results(self, batch_id: str) -> list[Any]:
        self.results_calls.append(batch_id)
        out = []
        for req in self._requests_by_batch_id[batch_id]:
            custom_id = req["custom_id"]
            blocks = self._blocks_by_custom_id[custom_id]
            seg_nodes = (
                _gate_failing_response_for(blocks)
                if self._gate_failing
                else self._segment_fn("", blocks)
            )
            text = _seg_nodes_to_response_text(seg_nodes)
            result = SimpleNamespace(
                type="succeeded",
                message=SimpleNamespace(content=[SimpleNamespace(type="text", text=text)]),
            )
            out.append(SimpleNamespace(custom_id=custom_id, result=result))
        return out


class _FakeBatchClient:
    """Fake Anthropic client exposing only the batches surface segment_documents_batch uses."""

    def __init__(
        self, blocks_by_custom_id: dict[str, list[Block]], *, gate_failing: bool = False
    ) -> None:
        self.messages = SimpleNamespace(
            batches=_FakeBatchesResource(
                blocks_by_custom_id, _fake_segment_fn, gate_failing=gate_failing
            )
        )


def _make_batch_client(corpus_dir: Path, *, gate_failing: bool = False) -> _FakeBatchClient:
    """Build a fake batches client pre-registered with every version's blocks.

    Extracts every ``.rtf`` file under *corpus_dir* the same way
    ``_collect_batch_items`` does and keys each version's block stream by its
    ``{doc_id}/{version}`` custom_id, so the fake's canned response for a
    given custom_id always matches the fixture it corresponds to.
    """
    blocks_by_custom_id: dict[str, list[Block]] = {}
    for doc_dir in sorted(d for d in corpus_dir.iterdir() if d.is_dir()):
        for vf in sorted(doc_dir.glob("*.rtf")):
            _canonical_text, blocks, _extractor = extract_blocks(vf)
            blocks_by_custom_id[f"{doc_dir.name}/{vf.stem}"] = blocks
    return _FakeBatchClient(blocks_by_custom_id, gate_failing=gate_failing)


def test_batch_segmentation_produces_classified_observations(tmp_path: Path) -> None:
    """mine_corpus(use_batch_segmentation=True) with a mocked batches client
    produces observations whose taxonomy_id reflects the batched SegNode
    output — same classified-observations contract as the synchronous path,
    but through the real segment_documents_batch wiring.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    client = _make_batch_client(corpus_dir)

    def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
        return segment_documents_batch(
            items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
        )

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        use_batch_segmentation=True,
        segment_documents_batch_fn=_batch_fn,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs, "mine_corpus must write at least one observation"

    taxonomy_ids = [o["taxonomy_id"] for o in raw_obs]
    assert not all(tid is None for tid in taxonomy_ids), (
        "batch-segmented observations must carry the batched taxonomy_id, not be all-None"
    )
    assert set(taxonomy_ids) == {"indemnification", "governing_law"}

    # Exactly one corpus-wide batch call was made (not one per document).
    assert len(client.messages.batches.create_calls) == 1
    submitted_ids = {r["custom_id"] for r in client.messages.batches.create_calls[0]["requests"]}
    assert submitted_ids == {"deal-001/v1"}


def test_batch_segmentation_two_versions_diff_and_playbook_validate(tmp_path: Path) -> None:
    """Two batch-segmented versions run through diff/deviation without error,
    and the resulting playbook passes playbook-schema validation — the
    downstream contract must hold identically to the synchronous LLM path.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=True)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    client = _make_batch_client(corpus_dir)

    def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
        return segment_documents_batch(
            items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
        )

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        use_batch_segmentation=True,
        segment_documents_batch_fn=_batch_fn,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs, "two-version diff must still produce observations"
    assert {"indemnification", "governing_law"} <= {o["taxonomy_id"] for o in raw_obs}

    # Both versions' custom_ids went into the single corpus-wide batch call.
    submitted_ids = {r["custom_id"] for r in client.messages.batches.create_calls[0]["requests"]}
    assert submitted_ids == {"deal-001/v1", "deal-001/v2"}

    playbook = project_playbook(out_dir=out_dir, config=cfg, taxonomy=taxonomy)

    result = validate_document(playbook)
    assert result.ok, f"playbook validate must pass: {[str(e) for e in result.errors]}"


def test_batch_segmentation_gate_failure_quarantines_document(tmp_path: Path) -> None:
    """A batched SegNode result that fails the QA gates at grounding time must
    quarantine that document (fail loud, recorded in quarantine.json) — no
    repair loop, no silent fallback to a worse tree, no aborting the whole run.
    Same isolated-quarantine contract as the synchronous path.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    client = _make_batch_client(corpus_dir, gate_failing=True)

    def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
        return segment_documents_batch(
            items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
        )

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        use_batch_segmentation=True,
        segment_documents_batch_fn=_batch_fn,
    )

    assert not read_observations_jsonl(out_dir / "observations.jsonl"), (
        "the quarantined document must contribute no observations"
    )
    quarantine = json.loads((out_dir / "quarantine.json").read_text(encoding="utf-8"))
    assert [q["document_id"] for q in quarantine] == ["deal-001"]
    assert "SegmentationQAError" in quarantine[0]["reason"]


def test_batch_segmentation_cache_avoids_second_batch_call(tmp_path: Path) -> None:
    """A SegmentationVerdictCache passed as segmentation_cache must be honored:
    a second mine_corpus run (with the L1-L4 stage cache disabled so
    _compute_doc_result actually re-executes) hits the segmentation cache and
    makes no second batch .create() call.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    client = _make_batch_client(corpus_dir)
    seg_cache = SegmentationVerdictCache(tmp_path / "seg_cache.jsonl")

    def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
        return segment_documents_batch(
            items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
        )

    common_kwargs: dict[str, Any] = {
        "corpus_dir": corpus_dir,
        "config": cfg,
        "taxonomy": taxonomy,
        "use_llm_segmentation": True,
        "use_batch_segmentation": True,
        "segment_documents_batch_fn": _batch_fn,
        "segmentation_cache": seg_cache,
        "no_cache": True,  # disable the unrelated L1-L4 stage cache
    }

    mine_corpus(out_dir=out_dir, **common_kwargs)
    assert len(client.messages.batches.create_calls) == 1

    out_dir_2 = tmp_path / "out2"
    mine_corpus(out_dir=out_dir_2, **common_kwargs)
    # No new batch submitted for the second run — the content-hash cache
    # satisfied every item without touching the client at all.
    assert len(client.messages.batches.create_calls) == 1

    raw_obs = read_observations_jsonl(out_dir_2 / "observations.jsonl")
    assert {o["taxonomy_id"] for o in raw_obs} == {"indemnification", "governing_law"}


def test_batch_segmentation_default_false_unaffected(tmp_path: Path) -> None:
    """use_batch_segmentation defaults to False: the synchronous per-document
    LLM path (use_llm_segmentation alone) must remain unaffected — no
    unrequested batch call, no requirement to pass segment_documents_batch_fn.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    # No use_batch_segmentation kwarg, no segment_documents_batch_fn — any
    # accidental batch-path call would try to construct a real anthropic
    # client and fail.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert raw_obs
    assert {o["taxonomy_id"] for o in raw_obs} == {"indemnification", "governing_law"}


# ---------------------------------------------------------------------------
# Issue #92: batch pre-pass must exclude L1-L4 stage-cache hits
# ---------------------------------------------------------------------------


def test_batch_prepass_skips_stage_cache_hits(tmp_path: Path) -> None:
    """A document whose L1-L4 stage cache is already warm must be excluded
    from the batch-segmentation pre-pass entirely on a re-run.

    Without the fix, every document's version files are extracted and
    submitted to the (paid) Message Batch on every run — including documents
    the per-document loop below is about to replay verbatim from
    ``ArtifactStore``, whose batched results are then simply thrown away.

    Runs ``mine_corpus`` three times over the same ``out_dir`` (so the
    on-disk stage cache carries over between calls) with a fake
    ``segment_documents_batch_fn`` that records the custom_ids it was asked
    to submit on each call:

    1. Both documents are new -> both submitted.
    2. Nothing changed -> both are cache hits -> zero items submitted.
    3. Only deal-002 changes -> deal-001 (still cached) must be excluded;
       only deal-002's version is submitted.
    """
    corpus_dir = tmp_path / "corpus"
    deal1_dir = corpus_dir / "deal-001"
    deal2_dir = corpus_dir / "deal-002"
    deal1_dir.mkdir(parents=True)
    deal2_dir.mkdir(parents=True)
    _write_rtf(deal1_dir / "v1.rtf", _V1_BODY)
    _write_rtf(deal2_dir / "v1.rtf", _V1_BODY)

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    submitted_by_call: list[set[str]] = []

    def _make_tracking_batch_fn() -> Any:
        # A fresh fake client per call (mirrors _make_batch_client's registry
        # picking up whatever the corpus currently contains on disk).
        client = _make_batch_client(corpus_dir)

        def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
            submitted_by_call.append({item.custom_id for item in items})
            return segment_documents_batch(
                items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
            )

        return _batch_fn

    run_kwargs: dict[str, Any] = {
        "corpus_dir": corpus_dir,
        "config": config,
        "taxonomy": taxonomy,
        "out_dir": out_dir,
        "use_llm_segmentation": True,
        "use_batch_segmentation": True,
    }

    # 1) First run: both documents are new — both go into the pre-pass.
    mine_corpus(segment_documents_batch_fn=_make_tracking_batch_fn(), **run_kwargs)
    assert submitted_by_call[-1] == {"deal-001/v1", "deal-002/v1"}

    # 2) Second run, same out_dir (warm stage cache): nothing changed on
    # disk — both documents are cache hits, so the pre-pass must submit
    # zero items.
    mine_corpus(segment_documents_batch_fn=_make_tracking_batch_fn(), **run_kwargs)
    assert submitted_by_call[-1] == set()

    # 3) Third run: only deal-002 changes. deal-001 is still a cache hit and
    # must stay excluded from the pre-pass — only deal-002's version goes in.
    _write_rtf(deal2_dir / "v1.rtf", _V2_BODY)
    mine_corpus(segment_documents_batch_fn=_make_tracking_batch_fn(), **run_kwargs)
    assert submitted_by_call[-1] == {"deal-002/v1"}

    # Sanity: the corpus as a whole still produced observations for both
    # documents across the store (final manifest reflects both).
    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    assert {o["citation"]["document_id"] for o in raw_obs} >= {"deal-001", "deal-002"}


# ---------------------------------------------------------------------------
# Issue #81: structured fallback reason, fixed manifest mislabel, and
# config.extraction.max_fallback enforcement.
# ---------------------------------------------------------------------------


def test_live_fallback_records_legacy_and_backend_error_reason(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A live per-file docling failure on the SYNCHRONOUS LLM path must be
    recorded as extractor="legacy" + reason="backend-error" in
    corpus_manifest.json's version_ingest — not mislabeled "docling" (the
    old up-front detect_extractor(vf) PATH-check guess, corrected only on a
    cache-hit replay) and not silently invisible.

    Must fail against pre-fix code: pre-fix, version_ingest["extractor"]
    would be the up-front "docling" guess (shutil.which reports it present)
    with no "reason" key at all, since no extraction_cache is passed here to
    trigger the old cache-hit-relabel correction.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    _mock_docling_present_and_failing(monkeypatch)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["extractor"] == "legacy"
    assert ingest["reason"] == "backend-error"
    assert "detail" not in ingest, "the raw exception text must never reach version_ingest"


def test_batch_path_live_fallback_records_legacy_and_backend_error_reason(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The batch pre-pass (_collect_batch_items -> _BatchExtraction) must
    also surface the real extractor label for a live per-file docling
    fallback — not just the synchronous path.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    _mock_docling_present_and_failing(monkeypatch)

    client = _make_batch_client(corpus_dir)

    def _batch_fn(items: Any, *, taxonomy_ids: Any, cache: Any = None, **_kwargs: Any) -> Any:
        return segment_documents_batch(
            items, taxonomy_ids=taxonomy_ids, client=client, cache=cache, poll_interval_s=0
        )

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        use_batch_segmentation=True,
        segment_documents_batch_fn=_batch_fn,
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["extractor"] == "legacy"
    assert ingest["reason"] == "backend-error"


def test_max_fallback_exceeded_raises_pipeline_error_naming_tuples(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Exceeding config.extraction.max_fallback must fail the run, naming
    the offending (document_id, version, reason) tuples — but the manifest
    must still be written correctly first (an operator who exceeds the
    budget still gets a complete, inspectable corpus_manifest.json, not a
    half-finished run)."""
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.extraction.max_fallback = 0  # tolerate zero degradations

    _mock_docling_present_and_failing(monkeypatch)

    with pytest.raises(PipelineError) as exc_info:
        mine_corpus(
            corpus_dir=corpus_dir,
            config=cfg,
            taxonomy=taxonomy,
            out_dir=out_dir,
            use_llm_segmentation=True,
            llm_segment_fn=_fake_segment_fn,
        )
    message = str(exc_info.value)
    assert "deal-001" in message
    assert "v1" in message
    assert "backend-error" in message

    # The manifest was still written correctly before the raise.
    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["extractor"] == "legacy"
    assert ingest["reason"] == "backend-error"


def test_max_fallback_unbounded_by_default_does_not_raise(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """max_fallback is None by default (unbounded) — a live fallback must
    not fail the run unless a budget was actually configured."""
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    assert cfg.extraction.max_fallback is None

    _mock_docling_present_and_failing(monkeypatch)

    # Does not raise.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )


def test_max_fallback_not_tripped_by_declared_legacy(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """extraction.extractor: legacy (a deliberate config choice) must NEVER
    count against max_fallback, even at max_fallback=0 — only real
    degradations (env-missing/backend-error) count (verifier correction on
    issue #81)."""
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.extraction.max_fallback = 0
    cfg.extraction.extractor = "legacy"

    # docling IS on PATH but must never even be attempted (declared legacy) —
    # pandoc's own subprocess.run call (the real RTF legacy path) must still
    # go through untouched.
    def _boom_if_docling_invoked(
        cmd: list[str], **kwargs: object
    ) -> subprocess.CompletedProcess[str]:
        if cmd[0] == "docling":
            raise AssertionError("docling must never be invoked under a declared 'legacy' run")
        return _real_subprocess_run(cmd, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(
        extraction.shutil,
        "which",
        lambda cmd: "/usr/bin/docling" if cmd == "docling" else _real_which(cmd),
    )
    monkeypatch.setattr(extraction.subprocess, "run", _boom_if_docling_invoked)

    # Does NOT raise.
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["extractor"] == "legacy"
    assert ingest["reason"] == "declared"


def test_reason_survives_pseudonymization_rewrite(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """version_ingest[].reason must survive _pseudonymize_corpus_documents's
    dict-spread rewrite (issue #81's explicitly required test) — the
    pseudonymization pass copies whole version_ingest dicts, so extra keys
    should pass through, but this proves it rather than assuming it."""
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.provenance.known_entities = ["Beta University"]  # triggers the pseudonymization pass

    _mock_docling_present_and_failing(monkeypatch)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["reason"] == "backend-error", (
        "reason must survive the pseudonymization pass unchanged"
    )
    assert ingest["extractor"] == "legacy"


def test_fallback_detail_never_leaks_raw_entity_name(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """CRITICAL PRIVACY REQUIREMENT (issue #81): ExtractorLabel.detail
    (str(exc) from a docling failure) embeds the absolute source path, which
    embeds the counterparty/entity name baked into the corpus folder
    structure. It must never reach corpus_manifest.json — defeating the
    born-safe pseudonymization contract that artifact otherwise upholds.
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "Fictional University"
    deal_dir.mkdir(parents=True)
    _write_rtf(deal_dir / "v1.rtf", _V1_BODY)

    cfg_dict = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {
            "our_party_aliases": ["Alpha Corp"],
            "known_entities": ["Fictional University"],
        },
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg_dict), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)

    _mock_docling_present_and_failing(monkeypatch)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
        entity_registry_path=tmp_path / "registry.json",
    )

    manifest_text = (out_dir / "corpus_manifest.json").read_text(encoding="utf-8")
    manifest = json.loads(manifest_text)
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["reason"] == "backend-error", "sanity: the fallback must actually have happened"
    assert "detail" not in ingest, "version_ingest must never carry a 'detail' key at all"
    assert "Fictional University" not in manifest_text


def test_stage_cache_reason_version_bump_forces_recompute_on_upgrade(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A warm L1-L4 stage cache (out/.cache) written under a DIFFERENT
    version_ingest_reason_version must not be replayed — otherwise
    config.extraction.max_fallback (and corpus_manifest.json's reason field)
    would go silently blind on any out_dir mined before this fix (issue
    #81's stage-cache invalidation requirement).

    Simulates "before this fix" by pinning
    pipeline._VERSION_INGEST_REASON_VERSION to a different sentinel for run
    1 (so its L1-L4 stage-cache entry is filed under a fingerprint the real
    code will never compute) while docling succeeds cleanly. Restoring the
    real constant for run 2 AND making docling now fail on the exact same
    file makes any stale replay directly observable: if the cache
    incorrectly hit the run-1 entry, the manifest would still show
    "docling"/reason=None; only a genuine recompute observes the new
    failure and records "legacy"/"backend-error".
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.extraction.max_fallback = 0

    real_reason_version = pipeline_module._VERSION_INGEST_REASON_VERSION

    # Run 1: docling succeeds cleanly; cached under a SENTINEL fingerprint
    # standing in for "the fingerprint shape before this fix existed".
    monkeypatch.setattr(pipeline_module, "_VERSION_INGEST_REASON_VERSION", "sentinel-pre-81")
    _mock_docling_present_and_succeeding(monkeypatch, _V1_DOCLING_MARKDOWN, stem="v1")
    mine_corpus(
        corpus_dir=corpus_dir,
        config=cfg,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_fake_segment_fn,
    )
    manifest_1 = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest_1 = manifest_1[0]["version_ingest"][0]
    assert ingest_1["extractor"] == "docling"
    assert ingest_1.get("reason") is None

    # Restore the real constant AND make docling now fail on this file.
    monkeypatch.setattr(pipeline_module, "_VERSION_INGEST_REASON_VERSION", real_reason_version)
    _mock_docling_present_and_failing(monkeypatch)

    with pytest.raises(PipelineError):
        mine_corpus(
            corpus_dir=corpus_dir,
            config=cfg,
            taxonomy=taxonomy,
            out_dir=out_dir,
            use_llm_segmentation=True,
            llm_segment_fn=_fake_segment_fn,
        )

    manifest_2 = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest_2 = manifest_2[0]["version_ingest"][0]
    assert ingest_2["extractor"] == "legacy", "the stale run-1 cache entry must not be replayed"
    assert ingest_2["reason"] == "backend-error"


def test_pre_81_extraction_cache_entry_is_migrated_and_counted_as_a_fallback(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The SEPARATE, more persistent extraction_cache.jsonl (kept warm across
    every judge/mine/compile round, independent of the L1-L4 stage cache) must
    never replay a pre-#81 entry as ``reason=None`` — the blindness that made
    ``config.extraction.max_fallback``, corpus_manifest.json and the review
    flags report zero fallbacks on a corpus that DID fall back (issue #81,
    project-note item 6 — the exact scenario that deadlocked the first
    implementation attempt).

    Originally this was enforced by bumping the cache's ``format_version``, so
    the entry simply missed and the whole corpus was re-extracted. It is now
    enforced by the #81 ladder rung, which reconstructs the entry's reason
    exactly — a "legacy" value under a "docling" KEY environment can only be a
    live per-file fallback — and re-files it under the current format. The
    outcome this test cares about is unchanged and asserted below (the
    manifest records ``reason="backend-error"``, and max_fallback=0 raises),
    but it is now reached WITHOUT re-extracting the file: the seeded blocks
    survive into the manifest, and the cache reports one migration and zero
    invalidations.

    Seeds a pre-#81-shaped entry (``extractor="legacy"``, no
    reason/fallback_from/detail keys — as a real pre-#81 build would have
    written for a file that fell back under a docling-available environment)
    and re-mines into a fresh out_dir whose ONLY warm state is
    extraction_cache.jsonl.
    """
    corpus_dir, config_path, out_dir = _make_corpus(tmp_path, two_versions=False)
    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    cfg = load_config(config_path)
    cfg.extraction.max_fallback = 0

    vf = corpus_dir / "deal-001" / "v1.rtf"
    extraction_cache_path = out_dir / "extraction_cache.jsonl"
    extraction_cache = ExtractionCache(extraction_cache_path)

    # Seed the pre-#81-shaped entry under the PRE-#81 KEY — format_version
    # hardcoded to "1" (NOT via extraction._extraction_cache_payload, which
    # would pick up today's ladder head) so this genuinely simulates on-disk
    # state from before this fix, under the environment a docling-available
    # host would have resolved ("docling" — the per-file fallback pins the
    # KEY to the environment, not the post-fallback adapter label; see
    # extraction._extraction_cache_payload).
    #
    # The block stream mirrors what the legacy RTF adapter really produced for
    # this fixture (one block per paragraph — see _V1_BODY and
    # _fake_segment_fn): the entry is now REPLAYED rather than discarded, so
    # it has to be a faithful pre-#81 entry, not a placeholder.
    lines = [
        "1. Indemnification",
        "Alpha Corp shall indemnify Beta University against third-party claims "
        "arising from the placement programme.",
        "2. Governing Law",
        "This agreement is governed by the laws of the State of California.",
    ]
    canonical_text = "\n".join(lines)
    blocks: list[dict[str, object]] = []
    cursor = 0
    for index, line in enumerate(lines):
        blocks.append(
            {"block_id": f"b{index}", "page": 0, "char_span": [cursor, cursor + len(line)]}
        )
        cursor += len(line) + 1  # +1 for the joining newline
    pre_81_key_payload = {
        "file_sha256": extraction._sha256_file(vf),
        "format_version": "1",
        "extractor_env": "docling",
    }
    pre_81_value = {
        "canonical_text": canonical_text,
        "blocks": blocks,
        "extractor": "legacy",
        # No "reason"/"fallback_from" keys — exactly the pre-#81 shape.
    }
    extraction_cache._store.put(pre_81_key_payload, pre_81_value)

    _mock_docling_present_and_failing(monkeypatch)

    with pytest.raises(PipelineError) as exc_info:
        mine_corpus(
            corpus_dir=corpus_dir,
            config=cfg,
            taxonomy=taxonomy,
            out_dir=out_dir,
            use_llm_segmentation=True,
            llm_segment_fn=_fake_segment_fn,
            extraction_cache=extraction_cache,
        )
    assert "backend-error" in str(exc_info.value)

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest = manifest[0]["version_ingest"][0]
    assert ingest["extractor"] == "legacy"
    assert ingest["reason"] == "backend-error", (
        "a pre-#81 extraction_cache.jsonl entry must never be replayed as reason=None — "
        "the #81 ladder rung reconstructs the fallback it actually recorded"
    )
    assert extraction_cache.migrated_count == 1, (
        "the entry's blocks were still correct, so the reason must have been recovered by "
        "migration rather than by re-extracting the file"
    )
    assert extraction_cache.invalidated_count == 0


# ---------------------------------------------------------------------------
# Issue #85: the DOCX tracked-changes side-channel must reach tracked_by_vid
# (and therefore tracked_changes_overlay) on the LLM-segmentation path too —
# previously only the deterministic branch of _compute_doc_result ever built
# it (issue #88's own regression guard, test_pipeline_project.py's
# test_docx_redline_observation_carries_tracked_changes_attribution, only
# exercises that deterministic branch).
# ---------------------------------------------------------------------------


# Heading text -> taxonomy_id for the (unnumbered) DOCX fixtures in
# test_pipeline_project.py (_docx_no_tracked_changes/_docx_with_tracked_insertion):
# two Heading-1 paragraphs, "Obligations" then "Governing Law". Neither has an
# explicit numeric prefix, so BOTH docx_ingester.ingest_docx (the tracked-changes
# side-channel's own parse) and segmentation_grounding.build() (the LLM tree's
# clause_path) independently assign sequential paths by sibling order alone —
# "1" then "2" — so they land on the same clause_path without this fixture
# needing numbered headings.
_DOCX_HEADING_TAXONOMY = {
    "Obligations": "relationship_of_parties",
    "Governing Law": "governing_law",
}


def _docx_fake_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Pair consecutive (heading, body) blocks into clause nodes.

    Same block-pairing shape as _fake_segment_fn above, adapted to the
    unnumbered DOCX heading fixtures (see _DOCX_HEADING_TAXONOMY).
    """
    del canonical_text
    nodes: list[SegNode] = []
    for order, i in enumerate(range(0, len(blocks), 2), start=1):
        heading_block, body_block = blocks[i], blocks[i + 1]
        nodes.append(
            SegNode(
                node_id=f"n{order}",
                parent_id=None,
                order=order,
                heading=heading_block.text,
                taxonomy_id=_DOCX_HEADING_TAXONOMY.get(heading_block.text),
                start_block_id=heading_block.block_id,
                end_block_id=body_block.block_id,
                start_quote=heading_block.text[:10],
                end_quote=body_block.text[-10:],
            )
        )
    return nodes


def test_llm_tracked_changes_returns_side_channel_for_tracked_docx(tmp_path: Path) -> None:
    """Direct unit coverage for pipeline._llm_tracked_changes: a
    tracked-changes DOCX yields a non-None TrackedChanges (plus its
    docx_ingester unit stream — issue #118), with no warning logged on a
    clean parse.
    """
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)
    messages: list[str] = []

    tracked, units = _llm_tracked_changes(path, messages.append)

    assert tracked is not None
    assert any(c.author == "Alice" for c in tracked.changes)
    assert units, "docx_ingester's unit stream must be non-empty for a real DOCX"
    assert messages == []


def test_llm_tracked_changes_degrades_to_none_on_unreadable_docx(tmp_path: Path) -> None:
    """A DOCX python-docx cannot open must degrade to (None, []) + a
    progress warning, never raise — an LLM-segmented version's tree has
    already been produced successfully by the time this best-effort call
    runs, so a failure here must not retroactively fail the whole version.
    """
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx")
    messages: list[str] = []

    tracked, units = _llm_tracked_changes(path, messages.append)

    assert tracked is None
    assert units == []
    assert len(messages) == 1
    assert "tracked-changes side-channel unavailable" in messages[0]


def test_llm_segmented_docx_carries_tracked_changes_attribution(tmp_path: Path) -> None:
    """A tracked-change insertion in a DOCX segmented via the LLM path must
    still reach tracked_by_vid and, through it, tracked_changes_overlay.

    Regression guard for issue #85: previously the LLM-segmentation branches
    of _compute_doc_result never populated tracked_by_vid at all, so this
    attribution was structurally absent on every LLM-segmented corpus — the
    entire real-world case, since the deterministic path is the exception.
    Mirrors test_pipeline_project.py's
    test_docx_redline_observation_carries_tracked_changes_attribution (issue
    #88), routed through use_llm_segmentation=True with a fake segment_fn
    instead of the deterministic path.
    """
    from tests.test_pipeline_project import _docx_no_tracked_changes, _docx_with_tracked_insertion

    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-002"
    deal_dir.mkdir(parents=True)
    _docx_no_tracked_changes(deal_dir, "v1.docx")
    _docx_with_tracked_insertion(deal_dir, "v2.docx")

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")
    out_dir = tmp_path / "out"

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_docx_fake_segment_fn,
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    deal_obs = [o for o in raw_obs if o["citation"]["document_id"] == "deal-002"]
    assert deal_obs, "deal-002 must have observations"

    attributed = [o for o in deal_obs if o.get("attribution") is not None]
    assert attributed, (
        "LLM-segmented DOCX must carry tracked-changes attribution — "
        f"got attributions={[o.get('attribution') for o in deal_obs]}"
    )
    assert attributed[0]["attribution"] == {
        "author": "Alice",
        "date": "2024-03-15T10:00:00Z",
        "tracked_type": "insertion",
    }


# ---------------------------------------------------------------------------
# _bridge_tracked_changes_if_needed / _extract_blocks_for_bridge (issue #118)
# ---------------------------------------------------------------------------


def _tc(
    text: str,
    char_span: tuple[int, int] | None,
    author: str = "Alice",
    change_type: str = "insertion",
    clause_path: str = "1",
) -> TrackedChange:
    return TrackedChange(
        change_type=change_type,  # type: ignore[arg-type]
        author=author,
        date="2024-01-01",
        text=text,
        clause_path=clause_path,
        char_span=char_span,
    )


def _units(texts: list[str]) -> list[TextUnit]:
    units: list[TextUnit] = []
    offset = 0
    for t in texts:
        units.append(TextUnit(text=t, char_span=(offset, offset + len(t))))
        offset += len(t) + 1
    return units


def _blocks(texts: list[str]) -> list[Block]:
    blocks: list[Block] = []
    offset = 0
    for i, t in enumerate(texts):
        blocks.append(Block(block_id=f"b{i}", page=0, char_span=(offset, offset + len(t)), text=t))
        offset += len(t) + 1
    return blocks


def test_bridge_not_needed_on_legacy_extractor_leaves_spans_untouched() -> None:
    """Legacy adapter reuses docx_ingester's own paragraph-join text
    verbatim (#112's originally-landed case) — no bridge, spans pass
    through exactly as given."""
    units = _units(["1. Clause", "Body text."])
    blocks = _blocks(["Clause", "Body text."])
    tc = _tc("Body text.", units[1].char_span)
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])

    result = _bridge_tracked_changes_if_needed(tracked, units, blocks, ExtractorLabel("legacy"))

    assert result is not None
    assert result.changes[0].char_span == units[1].char_span


def test_bridge_not_needed_when_extractor_label_none_leaves_spans_untouched() -> None:
    """Deterministic path never calls this at all, but a None label
    (unknown extractor) must degrade the same safe way as "legacy" — never
    strip a span it has no reason to distrust."""
    units = _units(["1. Clause"])
    blocks = _blocks(["Clause"])
    tc = _tc("Clause", units[0].char_span)
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])

    result = _bridge_tracked_changes_if_needed(tracked, units, blocks, None)

    assert result is not None
    assert result.changes[0].char_span == units[0].char_span


def test_bridge_docling_extractor_with_blocks_translates_span() -> None:
    units = _units(["1. Clause", "Payment due in thirty days."])
    blocks = _blocks(["Clause", "Payment due in thirty days."])
    tc = _tc("Payment due in thirty days.", units[1].char_span)
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])

    result = _bridge_tracked_changes_if_needed(tracked, units, blocks, ExtractorLabel("docling"))

    assert result is not None
    span = result.changes[0].char_span
    assert span is not None
    assert span == blocks[1].char_span


def test_bridge_docling_extractor_without_blocks_strips_spans_to_none() -> None:
    """Safety gate (issue #118): a non-legacy extractor with no confirmed
    bridge target must never leave a raw, potentially wrong-coordinate-space
    span in place — clears it instead of trusting numeric coincidence."""
    units = _units(["1. Clause", "Payment due in thirty days."])
    tc = _tc("Payment due in thirty days.", units[1].char_span)
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])

    result = _bridge_tracked_changes_if_needed(tracked, units, None, ExtractorLabel("docling"))

    assert result is not None
    assert result.changes[0].char_span is None
    # Everything else about the change must survive untouched.
    assert result.changes[0].author == "Alice"
    assert result.changes[0].clause_path == "1"


def test_bridge_no_tracked_changes_returns_input_unchanged() -> None:
    assert _bridge_tracked_changes_if_needed(None, [], None, ExtractorLabel("docling")) is None

    empty = TrackedChanges(document_id="doc", version="v2", changes=[])
    result = _bridge_tracked_changes_if_needed(empty, [], None, ExtractorLabel("docling"))
    assert result is empty


def test_extract_blocks_for_bridge_returns_none_on_unreadable_file(tmp_path: Path) -> None:
    """Best-effort, never raises — tracked-changes attribution is a bonus
    signal and must not retroactively fail an already-successful version."""
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx")

    result = _extract_blocks_for_bridge(path, None, False, "auto")

    assert result is None


# ---------------------------------------------------------------------------
# issue #118 fix round 1, finding 1: _extract_blocks_for_bridge must be a
# cache hit — never a second real extraction — once the segmentation call
# has already populated extraction_cache, regardless of the run's own
# refresh_extraction (--no-cache) setting.
# ---------------------------------------------------------------------------


def test_extract_blocks_for_bridge_hits_cache_populated_by_segmentation_call(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Regression guard for issue #118 fix round 1, finding 1.

    ``_llm_segment_file`` (via ``segment_to_tree``) calls ``extract_blocks``
    first, which populates ``extraction_cache`` for this exact ``(path,
    cache, extractor)`` key. ``_extract_blocks_for_bridge``'s later re-fetch
    must land on that SAME cache entry: pipeline.py's only production call
    site (``_compute_doc_result``, ~pipeline.py:1818) now always passes
    ``refresh=False`` here, regardless of the run's own
    ``refresh_extraction``/``--no-cache`` setting — because the bridge must
    translate spans into the coordinate space the tree was just built from,
    and a second real extraction is not guaranteed to reproduce it even if
    byte-identical.

    Proven directly against the real per-format extractor entry point
    (``extraction._extract_legacy_lines``), not a mock of ``extract_blocks``
    itself, which would hide a regression at the call site. Mirrors the
    issue's own empirical proof: with a warm cache,
    ``_extract_blocks_for_bridge(p, cache, False, ...)`` performs zero real
    extractions; ``(p, cache, True, ...)`` performs one — the contrast that
    shows ``refresh`` genuinely matters here, not just a no-op parameter.
    """
    from tests.test_docx_ingester import _tracked_docx

    path = _tracked_docx(tmp_path)
    cache = ExtractionCache(tmp_path / "extraction_cache.jsonl")

    real_extract_legacy_lines = extraction._extract_legacy_lines
    calls: list[Path] = []

    def counting_extract_legacy_lines(p: Path, suffix: str) -> list[tuple[str, int]]:
        calls.append(p)
        return real_extract_legacy_lines(p, suffix)

    monkeypatch.setattr(extraction, "_extract_legacy_lines", counting_extract_legacy_lines)

    # Warm the cache exactly as the segmentation call's own extract_blocks
    # invocation does (segment_to_tree -> extract_blocks(path, cache=...,
    # refresh=refresh_extraction, extractor=...)). extractor="legacy" pins
    # this test to the real per-format extractor regardless of whether
    # docling happens to be on the host running it.
    extract_blocks(path, cache=cache, refresh=False, extractor="legacy")
    assert len(calls) == 1, "sanity: the warming call itself must be a real extraction"

    hit = _extract_blocks_for_bridge(path, cache, False, "legacy")
    assert hit is not None
    assert len(calls) == 1, (
        "the bridge re-fetch with refresh=False must be a cache hit — no second "
        "real extraction — once the segmentation call has already populated the "
        "cache for this file"
    )

    refreshed = _extract_blocks_for_bridge(path, cache, True, "legacy")
    assert refreshed is not None
    assert len(calls) == 2, (
        "contrast case: refresh_extraction=True passed directly to "
        "_extract_blocks_for_bridge DOES force a second real extraction — "
        "proving the production fix is the CALL SITE now always passing False, "
        "not some cache behavior that would have masked the bug either way"
    )


# ---------------------------------------------------------------------------
# issue #118 fix round 1, finding 2: end-to-end coverage for the two
# production bridge call sites in _compute_doc_result (pipeline.py:1787 and
# :1820), unreachable by every other test in this file — either the legacy
# extractor is used (no bridge needed) or the bridge functions are called
# directly rather than through mine_corpus.
#
# Headings carry an EXPLICIT number ("1.1.", "1.2.") so docx_ingester's own
# clause_path ("1.1"/"1.2", via _parse_clause_number's explicit-number
# branch) diverges from the LLM/agent segmenter's flat root-level clause_path
# ("1"/"2", via segmentation_grounding's own sequential numbering) — the same
# hierarchical-vs-flat mismatch #112 first measured on the real corpus. With
# that mismatch, enrich_clause_diff's clause-path filter can never admit the
# tracked change on its own — only a bridge that leaves a real, overlapping
# char_span in place can. That makes this fixture an honest test of the two
# call sites specifically: bypassing them (as if deleted) hits the safety
# gate's None-stripping (_bridge_tracked_changes_if_needed, "blocks is
# None") and the clause-path mismatch then leaves nothing to attribute to.
# ---------------------------------------------------------------------------

_NUMBERED_DOCX_HEADING_TAXONOMY = {
    "1.1. Confidentiality": "relationship_of_parties",
    "1.2. Term": "governing_law",
}


def _numbered_docx_fake_segment_fn(canonical_text: str, blocks: list[Block]) -> list[SegNode]:
    """Same (heading, body) block-pairing shape as ``_docx_fake_segment_fn``,
    but assigns FLAT root-level clause paths ("1", "2") that do not mirror
    the DOCX's own "1.1."/"1.2." numbering — exactly how a real LLM/agent
    segmenter's own numbering is independent of source-document numbering."""
    del canonical_text
    nodes: list[SegNode] = []
    for order, i in enumerate(range(0, len(blocks), 2), start=1):
        heading_block, body_block = blocks[i], blocks[i + 1]
        nodes.append(
            SegNode(
                node_id=f"n{order}",
                parent_id=None,
                order=order,
                heading=heading_block.text,
                taxonomy_id=_NUMBERED_DOCX_HEADING_TAXONOMY.get(heading_block.text),
                start_block_id=heading_block.block_id,
                end_block_id=body_block.block_id,
                start_quote=heading_block.text[:10],
                end_quote=body_block.text[-10:],
            )
        )
    return nodes


def _numbered_docx_no_tracked_changes(tmp_path: Path, filename: str) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("1.1. Confidentiality", level=1)
    doc.add_paragraph("Party A shall provide services to client.")
    doc.add_heading("1.2. Term", level=1)
    doc.add_paragraph("This agreement is governed by the laws of the State of California.")
    path = tmp_path / filename
    doc.save(str(path))
    return path


def _numbered_docx_with_tracked_insertion(tmp_path: Path, filename: str) -> Path:
    """Redlined version: 'promptly' inserted by Alice via w:ins in the
    Confidentiality clause, same convention as
    test_docx_ingester._tracked_docx / test_pipeline_project's
    _docx_with_tracked_insertion, but under numbered headings (see module
    section comment above). ALSO carries a second insertion by Bob, in the
    Term clause — two distinct authors in this version's side channel, so
    tracked_changes_overlay.round_level_fallback_attribution (single-author
    only) cannot fire and silently supply the Confidentiality clause's
    attribution regardless of whether the bridge ran; only enrich_clause_diff
    (clause-path OR bridged char_span) can attribute Alice's change."""
    from docx import Document
    from lxml import etree

    from tests.test_pipeline_project import _w

    doc = Document()
    doc.add_heading("1.1. Confidentiality", level=1)
    p = doc.add_paragraph()
    p.add_run("Party A shall ")

    ins_elem = etree.SubElement(p._p, _w("ins"))
    ins_elem.set(_w("id"), "1")
    ins_elem.set(_w("author"), "Alice")
    ins_elem.set(_w("date"), "2024-03-15T10:00:00Z")
    r_ins = etree.SubElement(ins_elem, _w("r"))
    t_ins = etree.SubElement(r_ins, _w("t"))
    t_ins.text = "promptly "

    p.add_run("provide services to client.")
    doc.add_heading("1.2. Term", level=1)
    p2 = doc.add_paragraph()
    p2.add_run("This agreement is governed ")

    ins_elem2 = etree.SubElement(p2._p, _w("ins"))
    ins_elem2.set(_w("id"), "2")
    ins_elem2.set(_w("author"), "Bob")
    ins_elem2.set(_w("date"), "2024-03-16T09:00:00Z")
    r_ins2 = etree.SubElement(ins_elem2, _w("r"))
    t_ins2 = etree.SubElement(r_ins2, _w("t"))
    t_ins2.text = "without delay "

    p2.add_run("by the laws of the State of California.")

    path = tmp_path / filename
    doc.save(str(path))
    return path


# Markdown docling would produce for the fixtures above. Neither heading
# matches docling's single-level _LIST_ITEM_RE ("1.1." is a multi-level
# dotted number — see test_span_bridge.py's own
# test_normalize_strips_multi_level_dotted_numbering), so the numbering is
# NOT stripped here; only the EXTRACTOR differs from docx_ingester's own
# parse, not the text.
_NUMBERED_V1_DOCLING_MARKDOWN = (
    "1.1. Confidentiality\n"
    "\n"
    "Party A shall provide services to client.\n"
    "\n"
    "1.2. Term\n"
    "\n"
    "This agreement is governed by the laws of the State of California.\n"
)

_NUMBERED_V2_DOCLING_MARKDOWN = (
    "1.1. Confidentiality\n"
    "\n"
    "Party A shall promptly provide services to client.\n"
    "\n"
    "1.2. Term\n"
    "\n"
    "This agreement is governed without delay by the laws of the State of California.\n"
)


def _mock_docling_present_multi_stem(
    monkeypatch: pytest.MonkeyPatch, markdown_by_stem: dict[str, str]
) -> None:
    """Like ``_mock_docling_present_and_succeeding``, but serves DIFFERENT
    markdown per input-file stem — both corpus versions below must actually
    go through docling for ``extractor_label`` to resolve to "docling" and
    reach the non-legacy bridge branch."""

    def fake_which(cmd: str) -> str | None:
        return "/usr/bin/docling" if cmd == "docling" else _real_which(cmd)

    def fake_run(cmd: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        if cmd[0] != "docling":
            return _real_subprocess_run(cmd, **kwargs)  # type: ignore[arg-type]
        stem = Path(cmd[2]).stem  # cmd == ["docling", "convert", str(path), ...]
        outdir = Path(cmd[cmd.index("--output") + 1])
        (outdir / f"{stem}.md").write_text(markdown_by_stem[stem], encoding="utf-8")
        return subprocess.CompletedProcess(cmd, 0, stdout="", stderr="")

    monkeypatch.setattr(extraction.shutil, "which", fake_which)
    monkeypatch.setattr(extraction.subprocess, "run", fake_run)


def test_llm_segmented_docx_via_docling_bridges_tracked_changes_attribution(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """End-to-end regression guard for issue #118 fix round 1, finding 2.

    Mines a tracked-changes DOCX through the SYNCHRONOUS LLM-segmentation
    branch of ``_compute_doc_result`` (pipeline.py:~1790-1822) with docling
    mocked present-and-succeeding, so ``extractor_label == "docling"`` and
    the non-legacy bridge branch (pipeline.py:1817-1822 —
    ``_extract_blocks_for_bridge`` + ``_bridge_tracked_changes_if_needed``)
    is actually exercised through ``mine_corpus`` — proven by asserting the
    resulting observation's ``attribution`` resolves to the redline author.

    Then re-mines the SAME corpus with the bridge call site bypassed
    (``_extract_blocks_for_bridge`` monkeypatched to always return ``None``,
    as if the call site were deleted) and asserts attribution is honestly
    ABSENT rather than the same correct result — proving this test would
    actually catch that regression, unlike every other test in this file
    (issue #118 fix round 1 measured: neutering the production bridge call
    leaves ``make all`` green except two direct unit tests of the bridge
    functions themselves, never through ``mine_corpus``).
    """
    corpus_dir = tmp_path / "corpus"
    deal_dir = corpus_dir / "deal-118"
    deal_dir.mkdir(parents=True)
    _numbered_docx_no_tracked_changes(deal_dir, "v1.docx")
    _numbered_docx_with_tracked_insertion(deal_dir, "v2.docx")

    cfg = {
        "agreement_type": {
            "id": "educational-affiliation",
            "name": "Educational Affiliation Agreement",
        },
        "baseline": {"template": None},
        "taxonomy": str(_TAXONOMY_PATH),
        "provenance": {"our_party_aliases": ["Alpha Corp"]},
    }
    config_path = tmp_path / "playbook.config.yaml"
    config_path.write_text(yaml.dump(cfg), encoding="utf-8")

    taxonomy = load_taxonomy(_TAXONOMY_PATH)
    config = load_config(config_path)

    _mock_docling_present_multi_stem(
        monkeypatch,
        {"v1": _NUMBERED_V1_DOCLING_MARKDOWN, "v2": _NUMBERED_V2_DOCLING_MARKDOWN},
    )

    out_dir = tmp_path / "out-bridged"
    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir,
        use_llm_segmentation=True,
        llm_segment_fn=_numbered_docx_fake_segment_fn,
        extraction_cache=ExtractionCache(out_dir / "extraction_cache.jsonl"),
    )

    manifest = json.loads((out_dir / "corpus_manifest.json").read_text(encoding="utf-8"))
    ingest_by_version = {vi["version"]: vi for vi in manifest[0]["version_ingest"]}
    assert ingest_by_version["v2"]["extractor"] == "docling", (
        "this test only proves what it claims to if the tracked-changes version "
        "actually went through docling, not a legacy fallback"
    )

    raw_obs = read_observations_jsonl(out_dir / "observations.jsonl")
    deal_obs = [o for o in raw_obs if o["citation"]["document_id"] == "deal-118"]
    assert deal_obs, "deal-118 must have observations"

    attributed = [o for o in deal_obs if o.get("attribution") is not None]
    assert attributed, (
        "a correctly bridged tracked-changes span must resolve to Alice's "
        f"insertion — got attributions={[o.get('attribution') for o in deal_obs]}"
    )
    assert attributed[0]["attribution"]["author"] == "Alice"
    assert attributed[0]["attribution"]["tracked_type"] == "insertion"

    # Bypass the bridge (as if pipeline.py:1817-1822 were deleted) and
    # re-mine the SAME corpus into a fresh out_dir.
    monkeypatch.setattr(pipeline_module, "_extract_blocks_for_bridge", lambda *a, **kw: None)

    out_dir_bypassed = tmp_path / "out-bypassed"
    mine_corpus(
        corpus_dir=corpus_dir,
        config=config,
        taxonomy=taxonomy,
        out_dir=out_dir_bypassed,
        use_llm_segmentation=True,
        llm_segment_fn=_numbered_docx_fake_segment_fn,
        extraction_cache=ExtractionCache(out_dir_bypassed / "extraction_cache.jsonl"),
    )

    raw_obs_bypassed = read_observations_jsonl(out_dir_bypassed / "observations.jsonl")
    deal_obs_bypassed = [o for o in raw_obs_bypassed if o["citation"]["document_id"] == "deal-118"]
    assert deal_obs_bypassed, "deal-118 must still have observations"
    assert all(o.get("attribution") is None for o in deal_obs_bypassed), (
        "with the bridge call site bypassed, the '1.1' vs '1' clause-path "
        "mismatch must leave every observation's attribution honestly absent "
        "(never a wrong guess) — proving the wiring this ticket exists to "
        "deliver is actually exercised end-to-end, not dead code"
    )


# ---------------------------------------------------------------------------
# _attribution_for_diff round-level fallback (issue #118)
# ---------------------------------------------------------------------------


def _clause_diff(hunks: tuple[TextHunk, ...]) -> ClauseDiff:
    return ClauseDiff(
        taxonomy_id="ind",
        clause_path_before="1",
        clause_path_after="1",
        kind="modified",
        hunks=hunks,
        text_before="original text",
        text_after="revised text",
    )


def test_attribution_for_diff_falls_back_to_round_level_when_no_hunk_matches() -> None:
    """No per-hunk match (dissimilar text) but exactly one distinct author
    in the version's side channel, AND single_round=True (the net diff IS
    that one round) — the fallback tier fires where enrich_clause_diff
    alone would leave this unattributed."""
    hunk = TextHunk(kind="insert", old_text="", new_text="wholly unrelated new language")
    tc = _tc("completely different redline text", None, author="Alice")
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])
    cd = _clause_diff((hunk,))

    result = _attribution_for_diff(cd, tracked, single_round=True)

    assert result is not None
    assert result.author == "Alice"
    assert result.date is None


def test_attribution_for_diff_round_level_fallback_refuses_with_two_authors() -> None:
    hunk = TextHunk(kind="insert", old_text="", new_text="wholly unrelated new language")
    tc1 = _tc("completely different redline text", None, author="Alice")
    tc2 = _tc("yet more different text", None, author="Bob", clause_path="9")
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc1, tc2])
    cd = _clause_diff((hunk,))

    result = _attribution_for_diff(cd, tracked, single_round=True)

    assert result is None


def test_attribution_for_diff_round_level_fallback_refuses_when_not_single_round() -> None:
    """issue #118 fix round 2, finding 1 regression guard: exactly one
    distinct author in the side channel — the fallback's OWN firing
    condition is satisfied — but single_round=False (a net diff spanning
    more than one negotiation round) must suppress it outright. Consulting
    only the signed version's side channel for a net diff that can bundle
    an earlier, differently-authored round's change is not "that round";
    firing here would convert an honest unknown into a confidently wrong
    attribution to the signed round's sole author."""
    hunk = TextHunk(kind="insert", old_text="", new_text="wholly unrelated new language")
    tc = _tc("completely different redline text", None, author="Alice")
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])
    cd = _clause_diff((hunk,))

    result = _attribution_for_diff(cd, tracked, single_round=False)

    assert result is None


def test_attribution_for_diff_prefers_direct_match_over_round_level_fallback() -> None:
    """When enrich_clause_diff DOES find a real per-hunk match, that match
    is returned — the fallback tier only ever fills a gap, never overrides
    a real signal. Also proves the direct match is unaffected by
    single_round=False: a real per-hunk match must survive multi-round net
    diffs exactly as it does for single-round ones."""
    hunk = TextHunk(kind="insert", old_text="", new_text="consequential damages excluded")
    tc = _tc("consequential damages excluded", None, author="Alice")
    tracked = TrackedChanges(document_id="doc", version="v2", changes=[tc])
    cd = _clause_diff((hunk,))

    result = _attribution_for_diff(cd, tracked, single_round=False)

    assert result is not None
    assert result.author == "Alice"
    assert result.date == "2024-01-01"  # real match carries the real date, not None
