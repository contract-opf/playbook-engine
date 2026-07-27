"""Tests for intake_plan.py — universal corpus intake (issue #186).

Builds a scrambled, flat-loose-files copy of the persistent
``examples/judge-fixture/corpus`` fixture (deal-alpha: v1/v2 rtf; deal-beta:
v1 rtf) in-test — renamed to generic ``doc_NNN`` names with no per-agreement
subfolders (an "unknown" layout: staging.detect_layout would refuse to guess
this one). No second fixture corpus is committed.

Original chronological order within a cluster is recovered via filesystem
mtime (the weakest evidence tier — see ``intake_plan._fallback_timestamp``),
which this test controls explicitly since the RTF fixture carries no
embedded-metadata dates for the stronger tiers to key off.
"""

from __future__ import annotations

import hashlib
import json
import os
import shutil
import time
from pathlib import Path

import pytest
from click.testing import CliRunner, Result
from docx import Document

from playbook_engine.cli import cli
from playbook_engine.intake_plan import build_staging_plan, execute_staging_plan
from playbook_engine.staging import _STAGING_MARKER, UnknownLayoutError, stage

FIXTURE_CORPUS = Path(__file__).parent.parent / "examples" / "judge-fixture" / "corpus"
_DAY = 86400.0


def _scrambled_corpus(dest: Path) -> None:
    """Write a flat, loosely-named, scrambled copy of FIXTURE_CORPUS into *dest*.

    Mapping (scrambled name -> original) is deliberately out of order and
    stripped of any agreement-identifying folder/name signal:
      doc_001.rtf <- deal-beta/v1.rtf   (mtime: -3 days)
      doc_002.rtf <- deal-alpha/v2.rtf  (mtime: -1 day)
      doc_003.rtf <- deal-alpha/v1.rtf  (mtime: -2 days)

    mtimes are set explicitly (not just copy order) so version_orderer's
    timestamp tie-break can recover alpha's original v1-before-v2 order from
    content-tied permutations (two non-signed versions with no signed
    anchor have exactly one pairwise distance, so cost alone can't order
    them — see version_orderer.order_versions).
    """
    dest.mkdir(parents=True, exist_ok=True)
    base = time.time()
    mapping = [
        (FIXTURE_CORPUS / "deal-beta" / "v1.rtf", "doc_001.rtf", -3 * _DAY),
        (FIXTURE_CORPUS / "deal-alpha" / "v2.rtf", "doc_002.rtf", -1 * _DAY),
        (FIXTURE_CORPUS / "deal-alpha" / "v1.rtf", "doc_003.rtf", -2 * _DAY),
    ]
    for original, name, offset in mapping:
        target = dest / name
        shutil.copy2(original, target)
        t = base + offset
        os.utime(target, (t, t))


def _add_lorem_ipsum_docx(dest: Path, name: str = "doc_004.docx") -> None:
    """Write an unrelated lorem-ipsum DOCX into *dest* — no vocabulary overlap
    with the legal-boilerplate fixture corpus, so its nearest-neighbour
    content distance to every real file is ~1.0 (see intake_plan.UNRELATED_CEILING)."""
    doc = Document()
    doc.add_paragraph("Lorem Ipsum")
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet consectetur adipiscing elit sed do eiusmod "
        "tempor incididunt ut labore et dolore magna aliqua."
    )
    doc.save(str(dest / name))


def _hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _resolved_files_by_deal(out_dir: Path) -> dict[str, list[str]]:
    """Map each staged agreement folder to its ordered list of content hashes
    (following symlinks) — for comparing staged output byte-for-byte,
    independent of deal_id naming AND of which source tree (scrambled temp
    copy vs. the persistent fixture) the files were staged from."""
    result: dict[str, list[str]] = {}
    for agreement_dir in sorted(p for p in out_dir.iterdir() if p.is_dir()):
        files = sorted(
            p
            for p in agreement_dir.iterdir()
            if (p.is_file() or p.is_symlink()) and p.name != "hints.yaml"
        )
        result[agreement_dir.name] = [_hash(f) for f in files]
    return result


# ---------------------------------------------------------------------------
# 1. Reconstruction
# ---------------------------------------------------------------------------


class TestBuildStagingPlan:
    def test_plan_reconstructs_scrambled_fixture(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        plan = build_staging_plan(scrambled)

        assert plan["layout"] == "unknown"
        assert plan["unassigned"] == []
        assert len(plan["deals"]) == 2

        by_size = sorted(plan["deals"], key=lambda d: len(d["files"]))
        singleton, pair = by_size[0], by_size[1]

        # deal-beta: exactly one version, unsigned.
        assert len(singleton["files"]) == 1
        assert singleton["files"][0]["path"] == "doc_001.rtf"
        assert singleton["files"][0]["signed"] is False

        # deal-alpha: two versions, original chronological order recovered
        # (v1 -> doc_003.rtf, v2 -> doc_002.rtf) despite scrambled naming,
        # and neither is signed (matches the un-scrambled fixture — same
        # signed anchors: both None).
        assert [f["path"] for f in pair["files"]] == ["doc_003.rtf", "doc_002.rtf"]
        assert [f["proposed_version"] for f in pair["files"]] == [1, 2]
        assert all(f["signed"] is False for f in pair["files"])

    def test_unrelated_file_unassigned(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        _add_lorem_ipsum_docx(scrambled)

        plan = build_staging_plan(scrambled)

        assert plan["unassigned"], "lorem-ipsum file should land in unassigned"
        unassigned_paths = [u["path"] for u in plan["unassigned"]]
        assert "doc_004.docx" in unassigned_paths
        for u in plan["unassigned"]:
            assert u["reason"]

        # The two real deals are still correctly assembled (the junk file
        # didn't merge into or otherwise disturb either cluster).
        assert len(plan["deals"]) == 2

        # CLI plan-only run over the same scrambled+junk corpus exits 0.
        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(tmp_path / "out"), "--plan-only"]
        )
        assert result.exit_code == 0, result.output

    def test_evidence_recorded_per_file(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        _add_lorem_ipsum_docx(scrambled)

        plan = build_staging_plan(scrambled)

        for deal in plan["deals"]:
            for f in deal["files"]:
                assert f["evidence"], f"{f['path']} has no evidence recorded"

    def test_progress_reports_one_line_per_file(self, tmp_path: Path) -> None:
        """Issue #44: build_staging_plan/_gather_evidence ingest every file
        with no progress feedback, so a large loose corpus sits silent for
        minutes. A supplied progress callback must receive one line per
        supported file ingested, plus a notice before the O(n^2) pairwise
        distance computation."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        lines: list[str] = []
        build_staging_plan(scrambled, progress=lines.append)

        # One line per ingested file (3 in the scrambled fixture), each
        # naming the file it just processed.
        per_file_lines = [ln for ln in lines if "doc_00" in ln]
        assert len(per_file_lines) == 3
        assert any("doc_001.rtf" in ln for ln in per_file_lines)
        assert any("doc_002.rtf" in ln for ln in per_file_lines)
        assert any("doc_003.rtf" in ln for ln in per_file_lines)

        # A notice line precedes the pairwise-distance computation.
        assert any("pairwise distance" in ln for ln in lines)

    def test_plan_only_cli_streams_per_file_progress(self, tmp_path: Path) -> None:
        """Manual verification from issue #44: `playbook stage --plan-only`
        must stream per-file lines instead of sitting silent between
        'src :'/'out :' and the finished plan."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(tmp_path / "out"), "--plan-only"]
        )
        assert result.exit_code == 0, result.output
        assert "doc_001.rtf" in result.output
        assert "pairwise distance" in result.output


# ---------------------------------------------------------------------------
# 2. stage() refuses to guess on an unknown layout
# ---------------------------------------------------------------------------


class TestPlanNeverStagesDirectly:
    def test_stage_raises_on_unknown_layout(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        with pytest.raises(UnknownLayoutError, match="staging_plan"):
            stage(scrambled, tmp_path / "out")

    def test_plan_never_stages_directly(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        runner = CliRunner()
        result = runner.invoke(cli, ["stage", str(scrambled), "--out", str(tmp_path / "out")])

        assert result.exit_code != 0
        assert "staging_plan" in result.output


# ---------------------------------------------------------------------------
# 3. execute_staging_plan matches direct staging of the un-scrambled fixture
# ---------------------------------------------------------------------------


class TestExecuteStagingPlan:
    def test_execute_plan_equals_canonical(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        plan = build_staging_plan(scrambled)
        out_from_plan = tmp_path / "out_plan"
        plan_result = execute_staging_plan(plan, scrambled, out_from_plan)

        out_canonical = tmp_path / "out_canonical"
        canonical_result = stage(FIXTURE_CORPUS, out_canonical)

        assert plan_result.agreement_count == canonical_result.agreement_count == 2
        assert plan_result.staged_count == canonical_result.staged_count == 3

        # Compare by resolved source file content, not by deal_id/agreement
        # folder naming (the plan can't recover "deal-alpha"/"deal-beta" as
        # names — only their content/order/signed-ness).
        plan_by_deal = _resolved_files_by_deal(out_from_plan)
        canonical_by_deal = _resolved_files_by_deal(out_canonical)

        plan_groups = {frozenset(files) for files in plan_by_deal.values()}
        canonical_groups = {frozenset(files) for files in canonical_by_deal.values()}
        assert plan_groups == canonical_groups

        # Version order within each group also matches (list, not just set).
        plan_order_by_group = {frozenset(files): files for files in plan_by_deal.values()}
        canonical_order_by_group = {frozenset(files): files for files in canonical_by_deal.values()}
        for group in plan_groups:
            assert plan_order_by_group[group] == canonical_order_by_group[group]

    def test_execute_writes_hints_with_ordering(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        plan = build_staging_plan(scrambled)
        out_dir = tmp_path / "out"
        execute_staging_plan(plan, scrambled, out_dir)

        for deal in plan["deals"]:
            hints_path = out_dir / deal["deal_id"] / "hints.yaml"
            assert hints_path.exists()

    def test_plan_json_roundtrip(self, tmp_path: Path) -> None:
        """A plan serialized to JSON and read back executes identically —
        the real CLI workflow (--plan-only writes it, --plan reads it back)."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        plan = build_staging_plan(scrambled)
        plan_path = tmp_path / "staging_plan.json"
        plan_path.write_text(json.dumps(plan), encoding="utf-8")

        reloaded = json.loads(plan_path.read_text(encoding="utf-8"))
        out_dir = tmp_path / "out"
        result = execute_staging_plan(reloaded, scrambled, out_dir)
        assert result.agreement_count == 2
        assert result.staged_count == 3


# ---------------------------------------------------------------------------
# 4. execute_staging_plan out_dir guard (issue #248) — same defect/fix as
#    staging.stage, since it shares staging._recreate_out_dir.
# ---------------------------------------------------------------------------


class TestExecuteStagingPlanOutDirGuard:
    def test_out_dir_equal_to_src_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)

        with pytest.raises(ValueError, match="overlaps the source corpus"):
            execute_staging_plan(plan, scrambled, scrambled)

        # src must be left untouched
        assert list(scrambled.iterdir())

    def test_out_dir_is_parent_of_src_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)

        with pytest.raises(ValueError, match="overlaps the source corpus"):
            execute_staging_plan(plan, scrambled, scrambled.parent)

        assert list(scrambled.iterdir())

    def test_out_dir_unrelated_nonstaging_dir_raises_without_deleting(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)

        out = tmp_path / "unrelated"
        out.mkdir()
        (out / "important.txt").write_text("do not delete me", encoding="utf-8")

        with pytest.raises(ValueError, match="refusing to delete non-staging directory"):
            execute_staging_plan(plan, scrambled, out)

        assert (out / "important.txt").read_text(encoding="utf-8") == "do not delete me"

    def test_out_dir_previous_staging_output_succeeds(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)

        out = tmp_path / "out"
        execute_staging_plan(plan, scrambled, out)  # first run — writes the marker
        result = execute_staging_plan(plan, scrambled, out)  # second run — carries marker
        assert result.agreement_count == 2


# ---------------------------------------------------------------------------
# 5. execute_staging_plan / stage --plan shape validation (issue #45) —
#    a hand/skill-edited plan is expected input, so malformed shape must be
#    rejected with a located ValueError/CLI error BEFORE out_dir is wiped,
#    not a raw KeyError/TypeError traceback fired mid-execution.
# ---------------------------------------------------------------------------


class TestExecuteStagingPlanShapeValidation:
    def test_missing_deal_id_raises_before_wipe(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        del plan["deals"][0]["deal_id"]

        out = tmp_path / "out"
        out.mkdir()
        marker = out / "pre-existing.txt"
        marker.write_text("keep me", encoding="utf-8")

        with pytest.raises(ValueError, match=r"deal 0: 'deal_id' must be"):
            execute_staging_plan(plan, scrambled, out)

        # out_dir must be untouched — the shape error fires before
        # _recreate_out_dir ever runs.
        assert marker.read_text(encoding="utf-8") == "keep me"

    def test_duplicate_deal_id_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        assert len(plan["deals"]) >= 2
        plan["deals"][1]["deal_id"] = plan["deals"][0]["deal_id"]

        with pytest.raises(ValueError, match="duplicate deal_id"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

    def test_missing_files_key_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        del plan["deals"][0]["files"]

        with pytest.raises(ValueError, match=r"'files' must be a non-empty list"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

    def test_missing_path_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        del plan["deals"][0]["files"][0]["path"]

        with pytest.raises(ValueError, match=r"'path' must be a non-empty string"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

    def test_string_proposed_version_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        plan["deals"][0]["files"][0]["proposed_version"] = "2"

        with pytest.raises(ValueError, match=r"'proposed_version' must be an int, got str"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

    def test_bool_proposed_version_rejected(self, tmp_path: Path) -> None:
        """bool is an int subclass in Python — must not slip through as version 1."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        plan["deals"][0]["files"][0]["proposed_version"] = True

        with pytest.raises(ValueError, match=r"'proposed_version' must be an int, got bool"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

    def test_non_list_deals_raises(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        with pytest.raises(ValueError, match="plan.deals must be a list"):
            execute_staging_plan({"deals": "not-a-list"}, scrambled, tmp_path / "out")

    def test_non_dict_plan_raises(self, tmp_path: Path) -> None:
        """A top-level JSON array (e.g. ``echo '[]' > plan.json``) reaches
        execute_staging_plan as a bare Python list, not a dict — the
        ``plan must be a JSON object`` branch, otherwise untested."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        with pytest.raises(ValueError, match=r"plan must be a JSON object, got list"):
            execute_staging_plan([], scrambled, tmp_path / "out")  # type: ignore[arg-type]

    def test_deal_id_path_traversal_rejected(self, tmp_path: Path) -> None:
        """deal_id is used directly as an out_dir subdirectory name
        (``out_dir / deal_id``) — a hand/skill-edited plan containing a
        traversal segment must be rejected, not followed outside out_dir."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        plan["deals"][0]["deal_id"] = "../escaped"

        out_dir = tmp_path / "out"
        with pytest.raises(ValueError, match=r"'deal_id' '\.\./escaped' must be a single path"):
            execute_staging_plan(plan, scrambled, out_dir)

        assert not (tmp_path / "escaped").exists()

    def test_deal_id_absolute_rejected(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        escape_target = tmp_path / "escaped-abs"
        plan["deals"][0]["deal_id"] = str(escape_target)

        with pytest.raises(ValueError, match=r"'deal_id' .* must be a single path"):
            execute_staging_plan(plan, scrambled, tmp_path / "out")

        assert not escape_target.exists()

    def test_file_path_traversal_rejected(self, tmp_path: Path) -> None:
        """files[].path is joined onto src_dir unchecked
        (``src_dir / f["path"]``) — a relative path containing '..' must be
        rejected rather than resolving outside src_dir."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        outside_secret = tmp_path / "outside-secret.rtf"
        outside_secret.write_text("not part of the corpus", encoding="utf-8")

        plan = build_staging_plan(scrambled)
        plan["deals"][0]["files"][0]["path"] = "../outside-secret.rtf"

        out_dir = tmp_path / "out"
        with pytest.raises(ValueError, match=r"'path' '\.\./outside-secret\.rtf' escapes src_dir"):
            execute_staging_plan(plan, scrambled, out_dir)

        assert not out_dir.exists()

    def test_file_path_absolute_rejected(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        outside_secret = tmp_path / "outside-secret.rtf"
        outside_secret.write_text("not part of the corpus", encoding="utf-8")

        plan = build_staging_plan(scrambled)
        plan["deals"][0]["files"][0]["path"] = str(outside_secret)

        out_dir = tmp_path / "out"
        with pytest.raises(ValueError, match=r"'path' .* must be relative to src_dir"):
            execute_staging_plan(plan, scrambled, out_dir)

        assert not out_dir.exists()

    def test_valid_plan_still_executes(self, tmp_path: Path) -> None:
        """Sanity check: validation doesn't reject well-formed plans."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)

        result = execute_staging_plan(plan, scrambled, tmp_path / "out")
        assert result.agreement_count == 2

    def test_symlinked_corpus_file_still_stages(self, tmp_path: Path) -> None:
        """A legitimate corpus file that is itself a symlink to a target
        outside src_dir (e.g. a content-addressed archive dir) must not be
        rejected — the plan's ``path`` is a plain, single-segment, in-corpus
        relative name that build_staging_plan produced itself; only the
        symlink's *resolved target* lives outside src_dir. The escape check
        must be lexical (rejecting '..' components), not resolution-based —
        a resolution-based check regresses this legitimate case (issue #45
        round 2)."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)

        archive = tmp_path / "archive"
        archive.mkdir()
        doc_003 = scrambled / "doc_003.rtf"
        real_target = archive / "doc_003_real.rtf"
        shutil.move(str(doc_003), str(real_target))
        doc_003.symlink_to(real_target)

        plan = build_staging_plan(scrambled)
        result = execute_staging_plan(plan, scrambled, tmp_path / "out")
        assert result.agreement_count == 2
        assert result.staged_count == 3

    def test_deal_id_collides_with_reserved_output_name_rejected(self, tmp_path: Path) -> None:
        """A deal_id equal to a filename ``stage --plan`` itself writes
        directly into out_dir (the executed-plan record, the scaffolded
        config, or the staging marker) would otherwise shadow that file with
        a same-named deal *directory*, surfacing as a raw IsADirectoryError
        deep in cli.py — after out_dir has already been wiped and files
        staged (issue #45 round 2)."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        plan = build_staging_plan(scrambled)
        plan["deals"][0]["deal_id"] = "staging_plan.json"

        out_dir = tmp_path / "out"
        with pytest.raises(ValueError, match=r"collides with a reserved output"):
            execute_staging_plan(plan, scrambled, out_dir)

        assert not out_dir.exists()


class TestStagePlanCliErrors:
    """Required verification (issue #45): `stage --plan` on a malformed plan
    exits 1 with a human-readable ERROR line on stderr, no unhandled
    exception escapes, and a pre-existing *staging* out_dir — including the
    user's hand-edited plan file and any content from a prior run — is left
    untouched.

    The out_dir built by ``_preexisting_staging_out`` deliberately carries
    the ``.playbook-staging`` marker (plus a ``staging_plan.json`` and prior
    staged content) so it is a directory ``staging._recreate_out_dir`` WOULD
    wipe absent the new pre-execution validation — unlike a plain
    non-staging directory, which issue #248's guard already refuses
    regardless of this fix. A test built on a non-staging out_dir would pass
    identically on pre-fix code (the #248 guard doing the work, not the
    validation this ticket adds) — see issue #45 review notes.
    """

    def _preexisting_staging_out(self, tmp_path: Path) -> Path:
        out = tmp_path / "out"
        out.mkdir()
        (out / _STAGING_MARKER).write_text("", encoding="utf-8")
        (out / "staging_plan.json").write_text('{"hand": "edited"}', encoding="utf-8")
        prior_deal = out / "deal-prior"
        prior_deal.mkdir()
        (prior_deal / "01__prior.rtf").write_text("prior staged content", encoding="utf-8")
        return out

    def _assert_untouched(self, out: Path) -> None:
        assert (out / _STAGING_MARKER).exists()
        assert (out / "staging_plan.json").read_text(encoding="utf-8") == '{"hand": "edited"}'
        assert (out / "deal-prior" / "01__prior.rtf").read_text(
            encoding="utf-8"
        ) == "prior staged content"

    def _assert_no_unhandled_exception(self, result: Result) -> None:
        # CliRunner.invoke stores whatever exception propagated out of the
        # command as result.exception — including a *clean* `raise
        # SystemExit(1)` (click.testing wraps it verbatim), so this
        # assertion actually discriminates a located, handled error from a
        # raw exception (e.g. KeyError) escaping the command: only the
        # former is `None` or a `SystemExit`. `"Traceback" not in
        # result.output` cannot make this distinction — CliRunner never
        # writes a traceback into the captured output stream for either
        # case.
        exc = result.exception
        assert exc is None or isinstance(exc, SystemExit), repr(exc)

    def test_invalid_json_exits_cleanly(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        out = self._preexisting_staging_out(tmp_path)

        bad_plan = tmp_path / "bad.json"
        bad_plan.write_text("{not valid json", encoding="utf-8")

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(out), "--plan", str(bad_plan)]
        )

        assert result.exit_code == 1, result.output
        self._assert_no_unhandled_exception(result)
        assert "ERROR" in result.stderr
        assert "is not valid JSON" in result.stderr
        self._assert_untouched(out)

    def test_missing_deal_id_exits_cleanly(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        out = self._preexisting_staging_out(tmp_path)

        plan = build_staging_plan(scrambled)
        del plan["deals"][0]["deal_id"]
        bad_plan = tmp_path / "bad.json"
        bad_plan.write_text(json.dumps(plan), encoding="utf-8")

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(out), "--plan", str(bad_plan)]
        )

        assert result.exit_code == 1, result.output
        self._assert_no_unhandled_exception(result)
        assert "ERROR: deal 0: 'deal_id' must be a non-empty string" in result.stderr
        self._assert_untouched(out)

    def test_string_proposed_version_exits_cleanly(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        out = self._preexisting_staging_out(tmp_path)

        plan = build_staging_plan(scrambled)
        plan["deals"][0]["files"][0]["proposed_version"] = "2"
        bad_plan = tmp_path / "bad.json"
        bad_plan.write_text(json.dumps(plan), encoding="utf-8")

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(out), "--plan", str(bad_plan)]
        )

        assert result.exit_code == 1, result.output
        self._assert_no_unhandled_exception(result)
        assert "'proposed_version' must be an int, got str" in result.stderr
        self._assert_untouched(out)

    def test_duplicate_deal_id_exits_cleanly(self, tmp_path: Path) -> None:
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        out = self._preexisting_staging_out(tmp_path)

        plan = build_staging_plan(scrambled)
        assert len(plan["deals"]) >= 2
        dup_id = plan["deals"][0]["deal_id"]
        plan["deals"][1]["deal_id"] = dup_id
        bad_plan = tmp_path / "bad.json"
        bad_plan.write_text(json.dumps(plan), encoding="utf-8")

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(out), "--plan", str(bad_plan)]
        )

        assert result.exit_code == 1, result.output
        self._assert_no_unhandled_exception(result)
        assert f"ERROR: deal 1: duplicate deal_id {dup_id!r}" in result.stderr
        self._assert_untouched(out)

    def test_plan_file_survives_malformed_plan_on_marker_bearing_out_dir(
        self, tmp_path: Path
    ) -> None:
        """Regression for the issue's own Notes: pre-execution validation
        protects the user's hand-edited plan file from being deleted by the
        rmtree on a doomed run. The documented workflow (cli.py, after every
        `stage --plan` run) re-persists the executed plan at
        ``out_dir/staging_plan.json`` — so running `stage --plan
        out/staging_plan.json` again against that same, now marker-bearing,
        out_dir is the normal case this protects, not an edge case. Pre-fix,
        `_recreate_out_dir` rmtrees out_dir — deleting the very plan file
        `--plan` points at, along with prior staged content — before the
        malformed `deal_id` is ever noticed."""
        scrambled = tmp_path / "scrambled"
        _scrambled_corpus(scrambled)
        out = self._preexisting_staging_out(tmp_path)

        plan = build_staging_plan(scrambled)
        del plan["deals"][0]["deal_id"]
        plan_path = out / "staging_plan.json"
        plan_path.write_text(json.dumps(plan), encoding="utf-8")

        runner = CliRunner()
        result = runner.invoke(
            cli, ["stage", str(scrambled), "--out", str(out), "--plan", str(plan_path)]
        )

        assert result.exit_code == 1, result.output
        self._assert_no_unhandled_exception(result)
        assert "ERROR: deal 0: 'deal_id' must be a non-empty string" in result.stderr
        assert json.loads(plan_path.read_text(encoding="utf-8")) == plan
        assert (out / "deal-prior" / "01__prior.rtf").read_text(
            encoding="utf-8"
        ) == "prior staged content"
