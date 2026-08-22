"""Block-stream extractor — deterministic document → (canonical text, blocks).

This is the extractor half of the LLM-segmentation seam (see
:mod:`playbook_engine.segmentation_grounding`): it turns a source document
into a flat, ordered stream of :class:`~playbook_engine.segmentation_grounding.Block`
plus the canonical text those blocks index into.  No structure detection, no
LLM — that happens downstream (the LLM segmenter groups blocks into clauses;
grounding reconstructs verbatim text/spans from this stream).

Supported formats:
  - **docling** (preferred, any of DOCX/PDF/RTF) — via a ``docling`` CLI
    subprocess (``docling convert <path> --to md --image-export-mode
    placeholder --output <tmpdir>``), like the ``pandoc`` RTF path below.
    docling converts the source into
    layout-aware Markdown; that Markdown is then parsed into one block per
    logical unit (heading, paragraph, list item, table row) in reading
    order.  This is used whenever ``docling`` is on ``PATH``
    (``shutil.which("docling")``) — see "docling vs. legacy adapters"
    below.  ``page`` is best-effort ``0`` (docling Markdown is not
    paginated — mirrors the RTF/DOCX convention).
  - **DOCX** (fallback) — walks the body XML the same way
    :mod:`playbook_engine.docx_ingester` does (reusing its
    ``_iter_body_blocks``/``_extract_para_text`` helpers), one block per
    non-empty paragraph or flattened table.  This matters because
    ``python-docx``'s ``paragraph.text`` silently drops text inside
    ``w:ins`` (tracked-change insertions) and ``doc.paragraphs`` skips
    table content entirely — using the same XML walk as the deterministic
    ingester keeps both paths in agreement on document content (issue #85).
    Tracked-change *deletions* are excluded (mirrors ``docx_ingester``):
    canonical text reflects current/accepted content, not withdrawn
    language. Headings are not treated specially here: this module only
    emits text blocks (heading detection is out of scope for this slice —
    the LLM infers boundaries from block text).
  - **PDF** (fallback) — via ``pdfplumber``.  One block per extracted text
    line, tagged with its 1-based source page number.
  - **RTF** (fallback) — via a ``pandoc`` subprocess (``pandoc <in.rtf> -t
    plain --wrap=none``).  One block per non-empty paragraph.  ``pandoc``
    is a system binary, not a Python dependency; when it is not on
    ``PATH`` this raises :class:`ExtractionError` rather than silently
    falling back.

docling vs. legacy adapters:
  ``extract_blocks`` prefers docling whenever ``shutil.which("docling")``
  finds the binary, regardless of file extension — docling handles
  DOCX/PDF/RTF uniformly and gives the LLM segmenter real heading/structure
  cues that the legacy per-format adapters above cannot. When docling is
  absent (e.g. host dev without the container), extraction falls back to
  the legacy adapters unchanged. Which path is used is logged at INFO *and*
  returned as ``extractor`` (``"docling"`` or ``"legacy"``) from
  :func:`extract_blocks` — see :func:`detect_extractor`. This was
  previously only visible via that ``logging.info`` line, which is
  suppressed by default Python logging config: a host install without
  docling would silently fall back to pdfplumber (no OCR) for scanned
  PDFs with no way for the operator to notice short of reading a
  scrolled-past log line (issue #129). docling itself is invoked as a
  subprocess only — it is never imported as a Python module, keeping the
  engine importable without ``torch`` on the host.

  This auto-preference is only the *default*. A corpus config can declare
  ``extraction: {extractor: docling}`` (fail loudly if docling is
  unavailable — no silent downgrade) or ``extraction: {extractor: legacy}``
  (force the legacy adapters even when docling IS installed, for
  deterministic container-free runs) via :func:`extract_blocks`'s
  ``extractor`` parameter — see :func:`_resolve_extractor_env` and issue
  #80. ``extractor="auto"`` (the default) is exactly the behavior described
  above, unchanged.

  Whenever the returned ``extractor`` label is ``"legacy"``, it is not
  self-explanatory WHY: a whole-run PATH miss under ``auto``, a live
  per-file docling crash recovered on this one file, and a corpus that
  deliberately declared ``extractor: legacy`` are three very different
  situations that used to be indistinguishable downstream. :class:`ExtractorLabel`
  (returned in place of the old plain string — issue #81) carries a
  structured ``reason`` — ``"env-missing" | "backend-error" | "declared" |
  None`` — so ``corpus_manifest.json``'s ``version_ingest``,
  ``config.extraction.max_fallback``'s budget, and the ``mine``
  CLI summary can all tell these apart instead of collapsing to one
  ambiguous ``"legacy"``.

  Redline (tracked-changes/commented) DOCX used to be the common case
  landing in that per-file ``"backend-error"`` bucket, since docling 2.x's
  DOCX backend crashes on tracked-changes/comment nodes (``etree.QName`` on
  a comment factory) — exactly what redline drafts contain, the
  highest-value documents in a negotiation corpus. As of issue #84, a DOCX
  docling failure is retried once on a pre-normalized copy (see
  :func:`_retry_docling_on_normalized_docx` and
  :mod:`playbook_engine.docx_normalizer`) before falling back — so most
  redlines now stay on the docling path (``reason=None``, real heading
  structure) instead of degrading to the legacy adapter, which has no
  heading detection at all. Only a DOCX whose docling failure survives even
  the normalized retry still falls all the way through to
  ``reason="backend-error"``.

Markdown → Block parsing (docling path) and citation cleanliness:
  The block ``text`` used for grounding/citation must be the *clean*
  clause text: Markdown decoration is stripped from ``text`` even though
  it is used to *detect* block boundaries. Concretely, per output line:
    - ATX headings (``# Heading``, ``## Heading``, …) become their own
      block; the leading ``#`` markers and following space are stripped.
    - List items (leading ``-``, ``*``, ``+``, or ``N.``) become their own
      block; the leading marker is stripped.
    - Markdown table rows (``| a | b |``) become one block per row with
      leading/trailing pipes stripped and cell separators normalized to
      ``" | "``.
    - Bold/italic decoration (``**text**``, ``*text*``) has its ``*``/``_``
      markers stripped from the block text.
    - Blank lines separate blocks but never become blocks themselves.
  Boundary detection happens on the raw (undecorated) line; stripping is
  applied only to the text stored on the ``Block`` — so heading/list/table
  structure still drives block boundaries even though the punctuation that
  signaled it is gone from the citable text.

canonical_text / char_span contract:
  ``canonical_text`` is every block's ``text`` joined by ``"\\n"``, in
  reading order.  Each block's ``char_span`` is its ``[start, end)`` offset
  into that joined string, so for every block:
  ``block.text == canonical_text[block.char_span[0]:block.char_span[1]]``.
  This is the same joining convention as the ``_stream`` test helper in
  ``tests/test_segmentation_grounding.py`` — grounding depends on it being
  exact, since it is what lets a clause's ``char_span`` be reconstructed
  verbatim from the LLM's block references.

Not in scope (see issue #78): OCR toggling for scanned PDFs (docling's OCR
is enabled in a later slice) and Dockerfile packaging of the ``docling``
binary (separate issue) — this module only shells out to it when present.

``ExtractionCache`` (issue #132): extraction (especially docling OCR over a
scanned PDF) is the single most expensive step in the LLM-segmentation path —
far more expensive than the LLM segmentation call itself, which
``SegmentationVerdictCache`` (see ``llm_segmenter_batch.py``) already caches
independently. Before this, ``extract_blocks`` had no cache of its own, so
any caller that (for good reason — see ``agent_judge.StoreBackedClassificationJudge``
et al.) bypasses the pipeline's L1-L4 ``ArtifactStore``/``JudgmentCache``
(``no_cache=True``, forced by ``playbook judge``'s store-backed judges to
avoid replaying stale ``needs_review`` sentinels — see ``cli.py``'s
``_verdict_store_kwargs``) also silently threw away every prior extraction,
re-running docling/pdfplumber/pandoc from scratch on every judge round over a
real multi-hundred-version corpus. ``ExtractionCache`` is content-addressed
on the source file's bytes and the current extractor environment (docling
vs. legacy — issue #77; extraction has no other judge/config dependency), so
it is safe to keep warm across rounds regardless of whatever ``no_cache``
value the judge wiring forces for the verdict-cache layers, while still
missing cleanly — rather than silently replaying stale output — if the
extractor environment itself changes between rounds.

``refresh`` (issue #78): the flip side of the above — an operator who
suspects bad extraction (e.g. a stale pre-docling-install entry) needs a way
to force a real re-extraction, not just a re-judgment; before this,
``--no-cache`` disabled the L1-L4 ``ArtifactStore``/``JudgmentCache`` but
never touched ``ExtractionCache``, so it silently replayed the same stale
blocks. :func:`extract_blocks`'s ``refresh`` parameter is the fix: when
True, ``cache.get``/``cache.get_failure`` are skipped (this call always
re-extracts from source) but ``cache.put`` still runs at the end (and so
does ``cache.put_failure``, EXCEPT when a cached success already exists
under the identical key — see :func:`extract_blocks`'s ``if not lines``
guard, issue #78 round 2: a same-key failure there is necessarily transient,
since extraction is a pure function of (bytes, extractor environment), so it
must not permanently clobber output already proven extractable), so the
cache is left warm and correct for the *next* call.
``refresh`` is deliberately a plain parameter to ``extract_blocks`` — not
tied to ``no_cache`` — because ``no_cache`` is also forced ``True`` by the
judge wiring described above, and gating extraction refresh on that same
boolean would defeat the judge-warm-cache guarantee this class exists for.
Callers thread their own distinct signal sourced from the operator's actual
``--no-cache`` CLI flag (see ``pipeline.mine_corpus``'s
``refresh_extraction`` parameter) — never from a forced/internal
``no_cache``.
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document

from playbook_engine.agent_judge import VerdictStore
from playbook_engine.artifact_store import _sha256_file
from playbook_engine.docx_ingester import (
    DocxIngesterError,
    TrackedChanges,
    _extract_para_text,
    _iter_body_blocks,
    ingest_docx,
)
from playbook_engine.docx_normalizer import normalize_tracked_docx
from playbook_engine.segmentation_grounding import Block

_log = logging.getLogger(__name__)

#: Bumped whenever ``extract_blocks``'s output shape changes in a way that
#: should invalidate previously cached entries (e.g. a block-parsing bug fix).
#:
#: v2 (issue #81): the returned/cached ``extractor`` (the tuple's third
#: element) changed from a plain "docling"/"legacy" string to an
#: :class:`ExtractorLabel` carrying a structured ``reason`` ("env-missing" |
#: "backend-error" | "declared" | ``None``) and ``fallback_from``. A warm
#: ``extraction_cache.jsonl`` entry from before this fix has neither key, so
#: ``ExtractionCache.get`` would silently reload it as ``reason=None`` — the
#: same "legacy" label a genuine per-file docling fallback carries — making
#: ``config.extraction.max_fallback``, ``corpus_manifest.json``, the review
#: flags, and the CLI's reason breakdown all permanently blind on any corpus
#: ever extracted before this fix. Bumping this forces one clean re-extraction
#: per file, after which the cache is warm and correct again.
#:
#: v3 (issue #84): a redline (tracked-changes/commented) DOCX that previously
#: hit the docling->legacy fallback is now retried once on a normalized copy
#: before falling back (see :func:`_retry_docling_on_normalized_docx` and
#: :mod:`playbook_engine.docx_normalizer`), so for the SAME file bytes under
#: the SAME extractor environment, the correct output now differs from what
#: was cached before this fix: docling block structure and ``reason=None``
#: instead of legacy blocks and ``reason="backend-error"``. Neither
#: ``file_sha256`` nor ``extractor_env`` changes for that file (the KEY
#: records the environment "docling", not the post-fallback adapter
#: "legacy" — see the pipeline.py comment added in d08b895), so without this
#: bump a warm pre-#84 entry would silently keep reloading the stale
#: legacy/backend-error result forever and this fix would never reach any
#: corpus already extracted. Bumping this forces one clean re-extraction per
#: file, after which the cache is warm and correct again.
_EXTRACTION_CACHE_FORMAT_VERSION = "3"

# ---------------------------------------------------------------------------
# Error
# ---------------------------------------------------------------------------


class ExtractionError(Exception):
    """Raised when a document cannot be extracted into a block stream.

    Covers unsupported file extensions, a missing ``pandoc``/``docling``
    binary, and extraction that yields no usable text (e.g. an empty/blank
    source or a failed/empty docling conversion).
    """


# ---------------------------------------------------------------------------
# ExtractorLabel — which extractor produced a document's blocks, and why
# ---------------------------------------------------------------------------


class ExtractorLabel(str):
    """Which extractor produced a document's blocks, and why (issue #81).

    A ``str`` subclass equal to (and interchangeable with) the plain
    ``"docling"``/``"legacy"`` string :func:`extract_blocks` returned before
    this — every pre-existing ``extractor == "legacy"``/f-string/``json.dumps``
    call site keeps working unchanged (``str.__eq__``/``__str__``/the ``json``
    module all operate on the underlying character data, ignoring the extra
    attributes below) — while carrying the structured reason a caller that
    wants one (``corpus_manifest.json``'s ``version_ingest``, the CLI's
    fallback tally, ``mine_corpus``'s ``max_fallback`` budget) can read
    directly, instead of reverse-engineering it from a suppressed log line.

    Attributes:
        extractor: Same value as ``str(self)`` — ``"docling"`` or
            ``"legacy"`` — exposed as a named attribute for readability at
            call sites that already hold a label.
        reason: Why LEGACY ran, or ``None`` when docling ran with no
            degradation at all. One of:
              - ``"env-missing"``: ``auto`` mode, docling was never on PATH.
              - ``"backend-error"``: docling was attempted (on PATH, or
                declared) and raised on THIS file — a live per-file recovery
                fallback.
              - ``"declared"``: config explicitly set
                ``extraction.extractor: legacy`` — a deliberate choice, not
                a degradation.
              - ``None``: ``extractor == "docling"`` — no fallback happened.
        fallback_from: ``"docling"`` when this is a live ``"backend-error"``
            fallback; ``None`` otherwise (including for ``"env-missing"``/
            ``"declared"`` — docling was never attempted in either case, so
            there is nothing it fell back FROM).
        detail: ``str(exc)`` from the docling failure that triggered a
            ``"backend-error"`` fallback, or ``None``. **IN-MEMORY / LOGGING
            ONLY**: it embeds the absolute source path, which embeds the
            counterparty/entity name baked into the corpus folder structure.
            NEVER persist this to ``corpus_manifest.json``, ``review.json``,
            or any other artifact inside the born-safe pseudonymization
            boundary — only ``reason``/``fallback_from`` (closed enums, no
            raw names) are safe to persist (issue #81).
    """

    reason: str | None
    fallback_from: str | None
    detail: str | None

    def __new__(
        cls,
        extractor: str,
        *,
        reason: str | None = None,
        fallback_from: str | None = None,
        detail: str | None = None,
    ) -> ExtractorLabel:
        obj = super().__new__(cls, extractor)
        obj.reason = reason
        obj.fallback_from = fallback_from
        obj.detail = detail
        return obj

    @property
    def extractor(self) -> str:
        return str(self)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def detect_extractor(path: Path) -> str:
    """Return which extractor :func:`extract_blocks` will use for *path*
    under ``extractor="auto"`` (the default — see :func:`_resolve_extractor_env`
    for the declared-override cases).

    ``"docling"`` when ``shutil.which("docling")`` finds the binary,
    ``"legacy"`` otherwise (pdfplumber/python-docx/pandoc, per format) —
    the exact same check ``extract_blocks`` makes internally to choose its
    code path when no extractor is declared. Factored out so callers that
    need to *record* the choice (corpus manifest, ``mine`` CLI
    output) can do so up front — including for a version whose extraction
    subsequently fails — without duplicating or risking drift from
    ``extract_blocks``'s own branch (issue #129: this was previously only
    visible via a suppressed ``logging.info`` line).
    """
    return "docling" if shutil.which("docling") is not None else "legacy"


#: Valid values for ``extraction.extractor`` (config.py) / ``extract_blocks``'s
#: ``extractor`` parameter — issue #80.
_VALID_EXTRACTORS = frozenset({"docling", "legacy", "auto"})


def _resolve_extractor_env(declared: str, path: Path) -> str:
    """Resolve a *declared* extractor ("docling" | "legacy" | "auto") to the
    concrete environment governing *path* for this call — always "docling"
    or "legacy", never "auto" (issue #80).

    This is the single source of truth for both (a) which code path
    :func:`extract_blocks` takes and (b) the extraction cache key's
    ``extractor_env`` component (:func:`_extraction_cache_payload`) — a
    declared "legacy" run must never share a cache key with a same-machine
    "docling"-environment run of the same file just because docling happens
    to be on PATH (mirrors the issue #77 fix that first folded the
    environment into the cache key).

    - ``"legacy"`` always resolves to ``"legacy"``, even when docling IS on
      PATH — this is what lets a config force the legacy adapters for a
      deterministic, container-free run.
    - ``"docling"`` resolves to ``"docling"`` if available, else raises
      immediately (see below) — a declared "docling" requirement is never
      silently downgraded.
    - ``"auto"`` (the default) preserves today's behavior exactly:
      whatever :func:`detect_extractor` reports for *path* right now.

    Raises:
        ExtractionError: *declared* is not one of "docling"/"legacy"/"auto",
            or *declared* is "docling" and the docling binary is not on
            PATH. Checked here — before any file I/O — so a corpus-wide
            misconfiguration (``extraction.extractor: docling`` on a
            docling-less host) fails immediately with ONE clear,
            actionable error instead of silently falling back to the
            legacy adapters (no OCR, no heading detection) for every
            version, as the Jul 14 host-run incident this closes did (see
            module docstring, "docling vs. legacy adapters", and issue
            #80).
    """
    if declared not in _VALID_EXTRACTORS:
        raise ExtractionError(
            f"invalid extractor {declared!r}; expected 'docling', 'legacy', or 'auto'"
        )
    if declared == "docling" and shutil.which("docling") is None:
        raise ExtractionError(
            "config declares extraction.extractor: docling, but the docling "
            "binary was not found on PATH. Install docling, run this corpus "
            "inside the project's container (see Dockerfile), or set "
            "extraction.extractor to 'legacy' or 'auto' (or omit the "
            "extraction: section) to use the legacy adapters instead."
        )
    if declared == "auto":
        return detect_extractor(path)
    return declared


def _resolve_cache_environment(path: Path, environment: str | None) -> str:
    """Resolve the cache-key environment for *path*: *environment* itself if
    given, else today's bare :func:`detect_extractor` PATH check.

    ``None`` is the "no override" default every :class:`ExtractionCache`
    method accepts (issue #80) — it preserves every pre-existing direct
    caller's behavior (tests, and ``pipeline._compute_doc_result``'s
    manifest-reporting lookup) exactly: a fresh, unconditional
    ``detect_extractor(path)`` check, same as before this parameter
    existed. Only :func:`extract_blocks` (via :func:`_resolve_extractor_env`)
    ever passes an explicit, possibly config-declared value.
    """
    return environment if environment is not None else detect_extractor(path)


def _extraction_cache_payload(path: Path, environment: str | None = None) -> dict[str, str]:
    """Cache-key payload for *path*: the file's raw content hash plus the
    extractor environment governing this call.

    Extraction is a pure function of the source bytes for a FIXED extractor
    environment — no judge, no segmentation model, no engine config affects
    it beyond that (issue #132) — but legacy and docling can produce
    materially different output for the SAME bytes (legacy has no OCR and
    can garble columns/scanned text), so the environment is part of the key
    itself, not just the stored value (issue #77). This is the authoritative
    scoping mechanism: a legacy-era entry (success OR failure) simply misses
    once docling becomes available, and vice versa, while a
    same-environment repeat lookup keeps hitting across store
    instances/rounds.

    *environment* defaults to ``None`` — see :func:`_resolve_cache_environment`
    for what that means. A caller (:func:`extract_blocks`) that has a
    config-declared extractor override (``"docling"`` or ``"legacy"``, issue
    #80) MUST pass the resolved value explicitly here, or a declared
    "legacy" run on a docling-equipped host would key under "docling" (the
    bare PATH check) and silently collide with a real docling-environment
    entry for the same file.
    """
    return {
        "file_sha256": _sha256_file(path),
        "format_version": _EXTRACTION_CACHE_FORMAT_VERSION,
        "extractor_env": _resolve_cache_environment(path, environment),
    }


class ExtractionCache:
    """Judge-once, deterministic-replay cache for :func:`extract_blocks` output.

    Wraps a :class:`~playbook_engine.agent_judge.VerdictStore` rather than
    reimplementing content-hash JSONL storage (same pattern as
    :class:`~playbook_engine.llm_segmenter_batch.SegmentationVerdictCache`).

    Cache key: the source file's raw content hash plus the current extractor
    environment (see :func:`_extraction_cache_payload`; issue #77) —
    independent of ``no_cache``, judge identity, or engine config, so a
    repeat ``playbook judge`` round can reuse a prior run's extracted
    blocks/clause trees for every version whose source file AND extractor
    environment are unchanged, even though the L1-L4 ``ArtifactStore``/
    ``JudgmentCache`` stage cache is deliberately bypassed for store-backed
    judge runs (issue #132). A docling install/removal between rounds is
    exactly the case that must NOT keep hitting — the environment component
    turns that into a clean miss instead of silently replaying stale output.
    """

    def __init__(self, cache_path: Path) -> None:
        self._store = VerdictStore(cache_path)

    def get(
        self, path: Path, *, extractor: str | None = None
    ) -> tuple[str, list[Block], ExtractorLabel] | None:
        """Return the cached ``(canonical_text, blocks, ExtractorLabel)``, or ``None`` on a miss.

        Per-block ``text`` is reconstructed from
        ``canonical_text[char_span[0]:char_span[1]]`` (the documented
        canonical_text/char_span invariant — module docstring) rather than
        read back from the stored entry, since :meth:`put` no longer persists
        it (issue #67). A ``b.get("text")`` fallback keeps pre-existing cache
        entries (written before this fix, which still carry a per-block
        ``"text"``) loading unchanged.

        ``extractor``: the concrete environment ("docling" or "legacy") to
        look this entry up under — see :func:`_resolve_cache_environment`
        (issue #80). Defaults to ``None`` (today's bare PATH-check
        behavior, unchanged for every pre-existing caller).

        The returned label's ``reason``/``fallback_from`` come from
        ``cached.get(...)`` (issue #81) — absent on any entry written before
        that field existed, so a stale (pre-#81) on-disk value still loads
        without raising; ``detail`` is never round-tripped through the cache
        (in-memory-only — see :class:`ExtractorLabel`), so a reloaded label's
        ``detail`` is always ``None`` even for an entry whose live extraction
        did carry one.
        """
        cached = self._store.get(_extraction_cache_payload(path, extractor))
        if cached is None or "error" in cached:
            return None
        canonical_text = cached["canonical_text"]
        blocks = [
            Block(
                block_id=b["block_id"],
                page=b["page"],
                char_span=(b["char_span"][0], b["char_span"][1]),
                text=b.get("text", canonical_text[b["char_span"][0] : b["char_span"][1]]),
            )
            for b in cached["blocks"]
        ]
        label = ExtractorLabel(
            cached["extractor"],
            reason=cached.get("reason"),
            fallback_from=cached.get("fallback_from"),
        )
        return canonical_text, blocks, label

    def put(
        self,
        path: Path,
        canonical_text: str,
        blocks: list[Block],
        extractor: str | ExtractorLabel,
        *,
        environment: str | None = None,
    ) -> None:
        """Store *canonical_text*/*blocks*/*extractor* for *path*'s current content.

        Per-block ``text`` is deliberately NOT persisted: it is fully
        determined by ``canonical_text[char_span[0]:char_span[1]]`` (the
        canonical_text/char_span invariant — module docstring), so storing it
        duplicated every document's text in the cache entry. ``get()``
        reconstructs it from the span instead, halving on-disk entry size and
        the in-memory footprint of ``VerdictStore``, which parses the whole
        file at construction (issue #67).

        ``extractor`` is the adapter that actually PRODUCED this content
        ("docling" or "legacy") — stored in the VALUE for reporting
        (corpus_manifest, etc.). When it is an :class:`ExtractorLabel`
        (issue #81), its ``reason``/``fallback_from`` are stored alongside it
        (``getattr`` so a plain ``str`` caller — none remain in this repo,
        but the type stays accepted for the same "richer third element, same
        arity" reason :func:`extract_blocks` does — degrades to ``None`` for
        both, same as a pre-#81 cache entry). ``detail`` is deliberately never
        stored (see :class:`ExtractorLabel` — it embeds the raw source path).
        ``environment`` is the concrete environment governing this entry's
        KEY (issue #80; see :func:`_resolve_cache_environment`); it differs
        from ``extractor`` only when a per-file docling failure fell back to
        legacy under a docling-declared/available environment (see
        :func:`extract_blocks`) — the KEY must stay pinned to the environment
        regardless of that per-file fallback, or a docling-environment run's
        cache would fragment across "docling" and "legacy" keys for the same
        file. Defaults to ``None`` (today's bare PATH-check behavior).
        """
        value: dict[str, Any] = {
            "canonical_text": canonical_text,
            "blocks": [
                {
                    "block_id": b.block_id,
                    "page": b.page,
                    "char_span": list(b.char_span),
                }
                for b in blocks
            ],
            "extractor": str(extractor),
            "reason": getattr(extractor, "reason", None),
            "fallback_from": getattr(extractor, "fallback_from", None),
        }
        self._store.put(_extraction_cache_payload(path, environment), value)

    def get_failure(self, path: Path, *, extractor: str | None = None) -> str | None:
        """Return the cached failure message for *path*, or ``None``.

        A failure entry only counts when it was produced by the SAME extractor
        environment as the current one — a file that failed under the legacy
        adapter must be retried once docling becomes available, but a file
        that already failed under docling should not be re-OCR'd (up to the
        full per-file timeout) on every subsequent pipeline command.

        ``extractor``: the concrete environment to look this entry up under
        (issue #80) — defaults to ``None`` (today's bare PATH-check
        behavior via :func:`_resolve_cache_environment`, unchanged for every
        pre-existing caller).

        Since issue #77 added the extractor environment to the cache KEY
        (:func:`_extraction_cache_payload`), ``self._store.get()`` below can
        no longer return a different-environment entry at all — the key is
        now the authoritative scoping mechanism, and this method's own
        ``cached.get("extractor") != resolved`` check is redundant for any
        entry written through :meth:`put_failure`. It is kept (not removed)
        as a belt-and-braces guard against the stored *value* ever
        disagreeing with the *key* it was filed under — e.g. a future direct
        ``_store.put`` call that bypasses :func:`_extraction_cache_payload`.
        """
        resolved = _resolve_cache_environment(path, extractor)
        cached = self._store.get(_extraction_cache_payload(path, resolved))
        if cached is None or "error" not in cached:
            return None
        if cached.get("extractor") != resolved:
            return None
        return str(cached["error"])

    def put_failure(self, path: Path, message: str, extractor: str) -> None:
        """Store a failure marker for *path*'s current content (see get_failure).

        ``extractor`` here is the environment that FAILED — used for BOTH
        the cache KEY and the stored value's label (issue #80/#77; unlike
        :meth:`put`, a failure has no separate "actual adapter" worth
        preserving distinct from the environment: a docling timeout that
        falls back to legacy internally before yielding nothing must still
        negative-cache under the DOCLING environment, or every subsequent
        docling-environment lookup would miss and re-burn the full OCR
        timeout instead of hitting this negative cache).
        """
        self._store.put(
            _extraction_cache_payload(path, extractor),
            {"error": message, "extractor": extractor},
        )


def extract_blocks(
    path: Path,
    *,
    cache: ExtractionCache | None = None,
    refresh: bool = False,
    extractor: str = "auto",
) -> tuple[str, list[Block], ExtractorLabel]:
    """Extract ``path`` into ``(canonical_text, blocks, extractor)``.

    ``blocks`` are in reading order; ``block_id`` values are ``"b0", "b1",
    …`` in that order.  Every block's ``char_span`` is an offset into
    ``canonical_text`` such that
    ``block.text == canonical_text[slice(*block.char_span)]``.
    The returned ``extractor`` is an :class:`ExtractorLabel` — a ``str``
    equal to ``"docling"`` or ``"legacy"`` (the adapter that actually
    produced this content — see :func:`detect_extractor` and, for a per-file
    docling failure, the fallback below) carrying a structured ``reason``
    for why LEGACY ran, if it did (issue #81).

    Args:
        path:  Path to a ``.docx``, ``.pdf``, or ``.rtf`` file.
        cache: Optional :class:`ExtractionCache`. When given, a hit on
               *path*'s current content skips extraction entirely (no
               docling subprocess, no pdfplumber/python-docx/pandoc call) —
               see issue #132. On a miss, the result is stored before
               returning. Defaults to ``None`` (no caching — every call
               re-extracts).
        refresh: If True, skip *cache*'s ``get``/``get_failure`` reads for
               this call — always re-extract from source — while still
               writing the fresh result via ``put``/``put_failure`` at the
               end: reads bypassed, writes still happen (issue #78). This is
               what lets an operator-invoked ``--no-cache`` force a real
               recompute of a suspect extraction instead of silently
               replaying stale cached blocks. Exception: a refresh attempt
               that yields no text does NOT overwrite a pre-existing cached
               SUCCESS for the same key — see the ``if not lines`` guard
               below; a same-key failure is necessarily transient (issue #78
               round 2), not new evidence the content is bad. Ignored when
               *cache* is None (already always re-extracts). Defaults to
               False.
        extractor: Declared extractor environment — ``"docling"``,
               ``"legacy"``, or ``"auto"`` (default; mirrors
               ``config.extraction.extractor``, issue #80). ``"auto"`` is
               exactly today's behavior (prefer docling when on PATH, else
               legacy). ``"docling"`` forces docling and raises
               :class:`ExtractionError` immediately — before any file I/O —
               if the docling binary is not on PATH, instead of silently
               falling back to the legacy adapters. ``"legacy"`` forces the
               legacy per-format adapters even when docling IS on PATH
               (deterministic, container-free runs). See
               :func:`_resolve_extractor_env`.

    Raises:
        ExtractionError: an invalid *extractor* value; *extractor* is
                          ``"docling"`` but the binary is not on PATH;
                          unsupported extension; missing ``pandoc`` (RTF);
                          or the document yields no non-empty text.
    """
    # Validate + availability-check the DECLARED extractor and resolve it to
    # the concrete environment governing this call — before any file I/O, so
    # a corpus-wide misconfiguration (extraction.extractor: docling on a
    # docling-less host) fails immediately rather than after this file has
    # already been opened/hashed (issue #80).
    environment = _resolve_extractor_env(extractor, path)

    if not path.is_file():
        raise ExtractionError(f"file not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in (".docx", ".pdf", ".rtf"):
        raise ExtractionError(f"unsupported file extension: {suffix!r} ({path})")

    if cache is not None and not refresh:
        cached = cache.get(path, extractor=environment)
        if cached is not None:
            return cached
        cached_failure = cache.get_failure(path, extractor=environment)
        if cached_failure is not None:
            raise ExtractionError(f"{cached_failure} (cached failure — same extractor)")

    resolved = environment
    # Structured reason for why LEGACY ran, if it did (issue #81) — None
    # means extractor=="docling" ran clean, no degradation. Set below in
    # whichever of the two branches actually determines *resolved*.
    reason: str | None = None
    fallback_from: str | None = None
    detail: str | None = None
    if resolved == "docling":
        _log.info("extract_blocks: using docling for %s", path)
        try:
            lines = _extract_docling_lines(path)
        except ExtractionError as exc:
            # docling's per-format backends can fail on inputs the legacy
            # adapters handle fine — notably docling 2.x's DOCX backend raises
            # on tracked-changes/comment nodes (``etree.QName`` on a comment
            # factory), which is exactly what redline drafts contain — the
            # highest-value documents in a negotiation corpus. For DOCX,
            # retry docling once on a pre-normalized copy (insertions
            # accepted, deletions rejected, comment markup stripped — see
            # :func:`playbook_engine.docx_normalizer.normalize_tracked_docx`)
            # before giving up on docling structure: this recovers real
            # heading detection for redlines instead of routing them through
            # the legacy adapter, which has none at all (issue #84). Only
            # when that retry ALSO fails (or the file isn't DOCX) does the
            # original per-file legacy fallback below run. Skipping the
            # document outright would silently drop negotiation versions and
            # corrupt the trail, so a scanned PDF that docling cannot OCR
            # will still yield little here (legacy has no OCR) and then
            # raise below, as before; born-digital docx/pdf are recovered.
            # The fallback is logged and reflected in the returned
            # ``extractor`` label (reason="backend-error") so it is visible
            # in reporting (issues #129, #81) and countable against
            # ``config.extraction.max_fallback`` — a redline recovered via
            # the normalized retry is NOT counted here, since ``resolved``
            # never leaves "docling" for it (see
            # :func:`_retry_docling_on_normalized_docx`). Unchanged by issue
            # #80's declared-extractor option: a declared "docling" run
            # still allows this same PER-FILE fallback (the corpus-wide
            # availability precondition above is a separate, one-time
            # concern — see _resolve_extractor_env).
            retried_lines = _retry_docling_on_normalized_docx(path) if suffix == ".docx" else None
            if retried_lines is not None:
                lines = retried_lines
            else:
                _log.warning(
                    "extract_blocks: docling failed on %s (%s); falling back to legacy adapter",
                    path,
                    exc,
                )
                lines = _extract_legacy_lines(path, suffix)
                resolved = "legacy"
                reason = "backend-error"
                fallback_from = "docling"
                # IN-MEMORY ONLY — never persisted (see ExtractorLabel.detail):
                # str(exc) embeds the absolute source path, which embeds the
                # counterparty/entity name baked into the corpus folder
                # structure (issue #81).
                detail = str(exc)
    else:
        _log.info(
            "extract_blocks: using legacy adapter for %s (declared extractor=%r)", path, extractor
        )
        lines = _extract_legacy_lines(path, suffix)
        # Two distinct reasons legacy was resolved with no docling attempt at
        # all (issue #81): the config explicitly declared it (a deliberate
        # choice, never counted against max_fallback), or "auto" mode found
        # no docling on PATH at all (a real degradation, counted). These are
        # the only two ways `environment` can already be "legacy" here — see
        # _resolve_extractor_env: declared=="docling" either raises above (if
        # unavailable) or resolves to "docling", never reaching this branch.
        reason = "declared" if extractor == "legacy" else "env-missing"

    if not lines:
        message = f"extraction yielded no text: {path}"
        # Negative-cache the full failed attempt (docling OCR can burn its
        # whole per-file timeout) so later pipeline commands fail fast
        # instead of re-attempting per round — see get_failure for the
        # extractor-environment scoping that keeps a docling upgrade able to
        # retry a legacy-era failure. Record the ENVIRONMENT (not the
        # post-fallback adapter label): a docling timeout falls back to
        # legacy before landing here, and storing "legacy" would make every
        # docling-environment lookup miss and re-burn the OCR timeout each
        # round.
        #
        # Exception: under refresh=True, if a cached SUCCESS already exists
        # under this exact key (same file bytes + extractor environment —
        # see _extraction_cache_payload), do NOT overwrite it with this
        # failure (issue #78 round 2). put_failure files under the SAME key
        # as put, so an unguarded write here would let a transient
        # re-extraction failure (e.g. the docling OCR-timeout mode this
        # comment already documents) permanently clobber a known-good entry
        # — recovery would be back to hand-deleting extraction_cache.jsonl,
        # exactly what this issue exists to eliminate. This is sound because
        # extraction is documented as a pure function of (bytes, extractor
        # environment) — the cache key inputs — so a prior success under
        # the IDENTICAL key already proves the content is extractable; a
        # refresh that nonetheless yields nothing is not new evidence to
        # the contrary. A refresh with no pre-existing success (first
        # attempt, or a prior same-key failure) still negative-caches as
        # before.
        if cache is not None and not (
            refresh and cache.get(path, extractor=environment) is not None
        ):
            cache.put_failure(path, message, environment)
        raise ExtractionError(message)

    canonical_text, blocks = _build_stream(lines)

    label = ExtractorLabel(resolved, reason=reason, fallback_from=fallback_from, detail=detail)

    if cache is not None:
        cache.put(path, canonical_text, blocks, label, environment=environment)

    return canonical_text, blocks, label


def _extract_legacy_lines(path: Path, suffix: str) -> list[tuple[str, int]]:
    """Dispatch to the legacy per-format adapter for ``suffix``.

    Shared by the ``docling`` fallback path and the no-docling path in
    :func:`extract_blocks`. ``suffix`` is the already-lowercased extension.
    """
    if suffix == ".docx":
        return _extract_docx_lines(path)
    if suffix == ".pdf":
        return _extract_pdf_lines(path)
    return _extract_rtf_lines(path)


# ---------------------------------------------------------------------------
# Shared stream builder
# ---------------------------------------------------------------------------


def _build_stream(lines: list[tuple[str, int]]) -> tuple[str, list[Block]]:
    """Build ``(canonical_text, blocks)`` from ``(text, page)`` pairs.

    ``canonical_text`` is every text joined by ``"\\n"``; each block's
    ``char_span`` is computed against that joined string.  Mirrors the
    ``_stream`` helper pattern in ``tests/test_segmentation_grounding.py``.
    """
    blocks: list[Block] = []
    offset = 0
    for i, (text, page) in enumerate(lines):
        if i > 0:
            offset += 1  # the "\n" separator
        span = (offset, offset + len(text))
        blocks.append(Block(block_id=f"b{i}", page=page, char_span=span, text=text))
        offset += len(text)
    canonical_text = "\n".join(text for text, _ in lines)
    return canonical_text, blocks


# ---------------------------------------------------------------------------
# docling (preferred: structure-preserving, via CLI subprocess)
# ---------------------------------------------------------------------------


def _extract_docling_lines(path: Path) -> list[tuple[str, int]]:
    """One ``(text, page=0)`` per logical Markdown unit, via ``docling``.

    Runs ``docling convert <path> --to md --image-export-mode placeholder
    --output <tmpdir>`` (a subprocess, never imported as a Python module) and
    parses the resulting ``<stem>.md`` into blocks. The ``convert`` subcommand
    is mandatory in docling >=2.x; ``placeholder`` image mode keeps page
    images out of the Markdown. docling Markdown is not paginated, so ``page``
    is always 0 (mirrors the RTF/DOCX convention).
    """
    with tempfile.TemporaryDirectory(prefix="docling-") as tmpdir:
        markdown = _run_docling(path, Path(tmpdir))
        return _parse_markdown_lines(markdown)


def _retry_docling_on_normalized_docx(path: Path) -> list[tuple[str, int]] | None:
    """Retry docling once on a pre-normalized copy of a DOCX that just failed it.

    docling 2.x's DOCX backend crashes on tracked-changes/comment nodes
    (``etree.QName`` on a comment factory) — exactly what redline drafts
    contain, the highest-value documents in a negotiation corpus. Rather than
    degrading straight to the legacy adapter (no heading detection at all)
    for those, normalize a temp copy — insertions accepted, deletions
    rejected, comment markup stripped, matching the accepted-changes
    semantics :mod:`playbook_engine.docx_ingester` already produces — and
    retry docling on THAT (issue #84).

    Design choice (retry-on-failure vs. proactive normalization): this module
    normalizes only AFTER docling has already raised on the original file,
    rather than detecting tracked changes/comments up front and normalizing
    every DOCX proactively before the first docling attempt. Retry-on-failure
    costs nothing on the overwhelming majority of DOCX that docling already
    parses cleanly on the first try, at the price of one wasted docling
    subprocess round-trip only on the redlines that actually need the retry —
    proactive detection would pay a normalization cost on every DOCX just to
    save that one wasted round-trip on the minority that fail.

    Returns the extracted lines on success, or ``None`` if normalization or
    the retry itself failed for any reason — a plain DOCX with nothing to
    normalize will fail identically on both attempts and correctly fall
    through to ``None`` here, so the caller's existing legacy fallback is
    unchanged for that case. Deliberately broad ``except Exception`` (not
    just :class:`ExtractionError`): normalizing an unparseable/corrupt DOCX
    can raise a python-docx/lxml error that isn't an ``ExtractionError`` at
    all, and any such failure must still fall through to the legacy adapter
    rather than propagate and mask the original docling failure. The temp
    file is always cleaned up before returning, in either outcome.
    """
    normalized_path: Path | None = None
    try:
        normalized_path = normalize_tracked_docx(path)
        return _extract_docling_lines(normalized_path)
    except Exception as exc:  # noqa: BLE001 — any failure here just means "no recovery"
        _log.warning(
            "extract_blocks: docling retry on normalized copy of %s failed (%s); "
            "falling back to legacy adapter",
            path,
            exc,
        )
        return None
    finally:
        if normalized_path is not None:
            normalized_path.unlink(missing_ok=True)


# OCR language passed to ``docling convert --ocr-lang``. English by default;
# docling's own default is Chinese, which corrupts Latin-script scans.
_DOCLING_OCR_LANG = "eng"

# Per-file wall-clock cap on the docling subprocess. docling cold-loads its
# models on every invocation and can hang indefinitely on a pathological
# input; without a cap, one bad file blocks the whole corpus mine run
# forever (issue #98). 10 minutes comfortably covers even large scanned
# PDFs; a genuinely stuck conversion is far more likely than a legitimate
# one still running at that point.
_DOCLING_TIMEOUT_S = 600


def _run_docling(path: Path, outdir: Path) -> str:
    """Invoke the ``docling`` CLI and return the produced Markdown text.

    Isolated in its own helper so the exact invocation is easy to adjust
    after in-container validation (see issue #78 Notes).
    """
    try:
        subprocess.run(
            [
                "docling",
                "convert",
                str(path),
                "--to",
                "md",
                # Emit a `<!-- image -->` placeholder instead of embedding page
                # images as multi-KB base64 data URIs (docling's default),
                # which would otherwise become garbage blocks that wreck token
                # cost and citation text. The parser drops the placeholders.
                "--image-export-mode",
                "placeholder",
                # Pin OCR to English. docling's RapidOCR default is Chinese
                # (`lang=["chinese"]`), which joins/garbles Latin-script words on
                # scanned SIGNED copies. Change per corpus language (e.g. `deu`).
                "--ocr-lang",
                _DOCLING_OCR_LANG,
                "--output",
                str(outdir),
            ],
            capture_output=True,
            text=True,
            check=True,
            timeout=_DOCLING_TIMEOUT_S,
        )
    except subprocess.TimeoutExpired as exc:
        raise ExtractionError(
            f"docling timed out after {_DOCLING_TIMEOUT_S}s converting {path}"
        ) from exc
    except subprocess.CalledProcessError as exc:
        raise ExtractionError(f"docling failed to convert {path}: {exc.stderr.strip()}") from exc
    except OSError as exc:
        raise ExtractionError(f"cannot run docling: {exc}") from exc

    md_path = outdir / f"{path.stem}.md"
    try:
        markdown = md_path.read_text(encoding="utf-8")
    except OSError as exc:
        raise ExtractionError(f"docling did not produce expected output {md_path}: {exc}") from exc

    if not markdown.strip():
        raise ExtractionError(f"docling produced empty output for {path}")

    return markdown


# Markdown line patterns used to detect block boundaries. Detection happens
# on the raw line; the *stored* block text has decoration stripped (see
# module docstring, "Markdown → Block parsing").
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_LIST_ITEM_RE = re.compile(r"^([-*+]|\d+[.)])\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\|(.*)\|$")
_EMPHASIS_RE = re.compile(r"(\*\*|\*|__|_)")
# Markdown image lines (`![alt](...)`, incl. base64 data URIs) and HTML
# comments (docling's `<!-- image -->` placeholder) are not citable clause
# text — they are dropped entirely, never emitted as blocks.
_IMAGE_RE = re.compile(r"^!\[")


def _strip_markdown_decoration(text: str) -> str:
    """Strip bold/italic emphasis markers from already-boundary-parsed text."""
    return _EMPHASIS_RE.sub("", text).strip()


def _parse_markdown_lines(markdown: str) -> list[tuple[str, int]]:
    """Parse docling Markdown into ``(clean_text, page=0)`` blocks.

    One block per heading, paragraph, list item, or table row, in reading
    order. Markdown decoration (``#``, ``**``/``*``/``_``, leading
    ``-``/``*``/``N.``, table pipes) is stripped from the stored text while
    still being used to detect block boundaries.
    """
    lines: list[tuple[str, int]] = []
    for raw_line in markdown.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            continue

        # Drop image lines and HTML comments (docling's image placeholder) —
        # they are never citable clause text.
        if stripped.startswith("<!--") or _IMAGE_RE.match(stripped):
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            text = _strip_markdown_decoration(heading_match.group(2))
            if text:
                lines.append((text, 0))
            continue

        table_match = _TABLE_ROW_RE.match(stripped)
        if table_match:
            cells = [c.strip() for c in table_match.group(1).split("|")]
            # Skip Markdown table separator rows, e.g. "| --- | --- |".
            if all(re.fullmatch(r":?-+:?", c) for c in cells if c):
                continue
            text = _strip_markdown_decoration(" | ".join(c for c in cells if c))
            if text:
                lines.append((text, 0))
            continue

        list_match = _LIST_ITEM_RE.match(stripped)
        if list_match:
            text = _strip_markdown_decoration(list_match.group(2))
            if text:
                lines.append((text, 0))
            continue

        text = _strip_markdown_decoration(stripped)
        if text:
            lines.append((text, 0))

    return lines


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _extract_docx_lines(path: Path) -> list[tuple[str, int]]:
    """One ``(text, page=0)`` per non-empty paragraph or table, in document order.

    Walks the body XML via :func:`playbook_engine.docx_ingester._iter_body_blocks`
    (paragraph elements and flattened ``w:tbl`` tables) instead of
    ``doc.paragraphs`` / ``paragraph.text``, so that:
      - tracked-change insertions (``w:ins``) are included — ``paragraph.text``
        only concatenates runs that are direct children of ``w:p`` and silently
        drops ``w:ins``-nested runs (issue #85);
      - table content is captured at all — ``doc.paragraphs`` skips tables
        entirely.
    Tracked-change deletions are excluded (mirrors ``docx_ingester``): the
    canonical text reflects current/accepted content, not withdrawn language.
    DOCX is not paginated at the model level, so ``page`` is always 0.
    Heading detection (``paragraph.style.name`` starting with "Heading") is
    not captured here — ``Block`` has no heading field; text blocks are
    sufficient for this slice (see module docstring).
    """
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"cannot open DOCX: {exc}") from exc

    lines: list[tuple[str, int]] = []
    for block in _iter_body_blocks(doc):
        if isinstance(block, str):
            # Flattened table text (see docx_ingester._flatten_table).
            text = block.strip()
        else:
            # Paragraph XML element — w:ins-aware, w:del-excluded text.
            para_text, _tracked = _extract_para_text(block)
            text = para_text.strip()
        if text:
            lines.append((text, 0))
    return lines


def extract_tracked_changes(path: Path) -> TrackedChanges | None:
    """Harvest *path*'s DOCX tracked-changes side-channel (issue #85).

    ``extract_blocks`` prefers docling for DOCX whenever it's on ``PATH``
    (see the module docstring), and docling only ever sees Markdown — it has
    no notion of ``w:ins``/``w:del`` at all. ``_extract_docx_lines`` (the
    legacy adapter) *does* walk the raw XML via ``_extract_para_text``, but
    only runs when docling is unavailable, declared off, or fails outright
    (and even then, discards the ``TrackedChange`` list it computes per
    paragraph — see that function). So neither of ``extract_blocks``'s two
    adapters is a seam callers can hang tracked-changes capture off of
    without making it depend on which one happened to run for this file.

    This function is independent of that choice entirely: it re-walks
    *path* itself via :func:`playbook_engine.docx_ingester.ingest_docx` —
    the SAME parse the deterministic (non-LLM) ingestion path already uses
    to populate its own ``tracked_by_vid`` entries (see
    ``pipeline._ingest_file_tracked``) — so an LLM-segmented version gets
    the identical side-channel a deterministically-segmented version of the
    same file would, regardless of whether docling or the legacy adapter
    produced its canonical text/blocks.

    MUST be called on *path* itself, never a normalized copy: docling's
    per-file redline retry (:func:`_retry_docling_on_normalized_docx`, issue
    #84) normalizes a TEMP copy — accepting insertions, rejecting deletions,
    stripping comment markup — before re-attempting docling on it, which
    deletes exactly the ``w:ins``/``w:del`` nodes this function harvests.
    That temp path is entirely internal to ``extract_blocks`` and never
    escapes this module, so any caller that (like the pipeline) only ever
    holds the original corpus source path already satisfies this.

    ``document_id``/``version`` on the returned ``TrackedChanges`` are left
    blank: this module has no notion of corpus document identity (mirrors
    ``extract_blocks(path)``'s own signature, which also takes no document
    id), and no consumer of ``TrackedChanges`` reads those two fields — only
    ``.changes`` (see ``tracked_changes_overlay.enrich_clause_diff``).

    ``clause_path`` on each returned ``TrackedChange`` comes straight from
    ``ingest_docx``'s deterministic clause-numbering (heading style /
    explicit numeric prefix / generated sequential path — see
    ``docx_ingester``'s module docstring). For a document whose clause
    boundaries the LLM segmenter groups the same way (the common case for a
    cleanly-numbered agreement), this lines up with the LLM tree's own
    ``clause_path`` and lets ``enrich_clause_diff`` find candidates; for a
    document where the two disagree, ``enrich_clause_diff`` (issue #112)
    also selects candidates by ``char_span`` interval overlap between this
    function's ``TrackedChange.char_span`` values and the diffed clause
    tree's ``ClauseNode.char_span`` — but that join is only trustworthy
    when both sides are offsets into the SAME normalized text. The
    ``char_span`` this function returns always comes from re-parsing
    *path* via ``ingest_docx`` (the docx-ingester's own paragraph-join
    text). That matches the LLM tree's ``char_span`` only when
    ``extract_blocks`` also produced that tree via the legacy DOCX adapter
    (``_extract_docx_lines`` reuses the same paragraph-join). Under the
    default ``extraction.extractor="auto"``, ``extract_blocks`` prefers
    docling for DOCX, so the LLM tree's ``char_span`` is instead an offset
    into docling's Markdown-derived text — or, for redline DOCX that crash
    docling's DOCX backend, a THIRD normalized-copy text via
    ``_retry_docling_on_normalized_docx`` — a different coordinate system
    from this function's ``char_span``. On that (default) path, span
    overlap degrades to numeric coincidence rather than real co-location,
    same as the clause-path matching it supplements: best-effort, not
    guaranteed (see ``tracked_changes_overlay.enrich_clause_diff``'s module
    comment for the full precondition).

    Returns ``None`` for non-DOCX paths (no tracked-changes concept) —
    mirrors ``pipeline._ingest_file_tracked``'s convention for RTF/PDF. A
    DOCX with no ``w:ins``/``w:del`` elements still returns a
    ``TrackedChanges`` with an empty ``changes`` list (not ``None``) — same
    convention.

    Raises:
        ExtractionError: *path* is not a valid/openable DOCX.
    """
    if path.suffix.lower() != ".docx":
        return None
    try:
        return ingest_docx(path, document_id="", version="").tracked
    except DocxIngesterError as exc:
        raise ExtractionError(f"cannot open DOCX: {exc}") from exc


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _extract_pdf_lines(path: Path) -> list[tuple[str, int]]:
    """One ``(text, page)`` per extracted text line, ``page`` 1-based."""
    try:
        lines: list[tuple[str, int]] = []
        with pdfplumber.open(str(path)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                for line in page.extract_text_lines():
                    text = line["text"].strip()
                    if text:
                        lines.append((text, page_number))
        return lines
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"cannot open PDF: {exc}") from exc


# ---------------------------------------------------------------------------
# RTF (via pandoc subprocess)
# ---------------------------------------------------------------------------


def _extract_rtf_lines(path: Path) -> list[tuple[str, int]]:
    """One ``(text, page=0)`` per paragraph, via a ``pandoc`` subprocess.

    RTF is not paginated at the model level, so ``page`` is always 0.
    ``pandoc`` is a system binary (not a Python package) — when it is not on
    ``PATH`` this raises :class:`ExtractionError` with a clear message
    rather than attempting a degraded fallback.
    """
    if shutil.which("pandoc") is None:
        raise ExtractionError(
            "pandoc is required to extract RTF but was not found on PATH "
            "(install it, e.g. `brew install pandoc`, or `apt-get install pandoc`)"
        )

    try:
        result = subprocess.run(
            ["pandoc", str(path), "-t", "plain", "--wrap=none"],
            capture_output=True,
            text=True,
            check=True,
        )
    except subprocess.CalledProcessError as exc:
        raise ExtractionError(f"pandoc failed to convert RTF: {exc.stderr.strip()}") from exc
    except OSError as exc:
        raise ExtractionError(f"cannot run pandoc: {exc}") from exc

    # pandoc's plain output separates paragraphs by blank lines (with
    # --wrap=none disabling mid-paragraph line wrapping, so each paragraph is
    # exactly one physical output line). Splitting on newlines and dropping
    # empty lines recovers one entry per paragraph.
    lines: list[tuple[str, int]] = []
    for raw_line in result.stdout.splitlines():
        text = raw_line.strip()
        if text:
            lines.append((text, 0))
    return lines
